"""Buyer-facing output checks for the market research focused routes.

Covers the rendered-deliverable defects fixed in the build script and
validators: internal evidence-class tokens, session narration, duplicated
action tables, evidence ID gaps, forced page-one break, and named-firm
coherence.
"""

from __future__ import annotations

import copy
import importlib.util
import json
import subprocess
import sys
import tempfile
import unittest
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
PYTHON = sys.executable
SCRIPTS = ROOT / "skills/market-research-workflow/scripts"


def load_module(path: Path, name: str):
    spec = importlib.util.spec_from_file_location(name, path)
    module = importlib.util.module_from_spec(spec)
    assert spec.loader is not None
    spec.loader.exec_module(module)
    return module


def refresh_record() -> dict:
    return {
        "schema_version": "1.2",
        "skill": "market-research-workflow",
        "workflow_mode": "refresh",
        "question": "What has changed in the nationwide help desk support market since the March 2025 report?",
        "scope": {
            "as_of_date": "2026-08-28",
            "agency": "Example Agency",
            "naics": ["541513"],
            "psc": ["DA01"],
            "period": "FY2023 through FY2026, with FY2026 identified as partial",
        },
        "document_register": [],
        "user_context": [],
        "assumptions": [],
        "web_research": {
            "mode": "no_public_web",
            "approved": True,
            "approved_at": "2026-08-28T14:00:00Z",
            "disclosure_acknowledged": True,
            "planned_providers": [],
            "providers_used": [],
            "fallback_events": [],
        },
        "queries": [
            {
                "id": "Q001",
                "provider": "federal_mcp",
                "operation": "USASpending spending over time",
                "parameters": {"naics": ["541513"], "psc": ["DA01"]},
                "retrieved_at": "2026-08-28T14:10:00Z",
                "count": 4,
                "limitations": "Illustrative example data (not live research)",
            }
        ],
        "evidence": [
            {
                "id": "E001",
                "source_class": "document",
                "title": "2025 Market Research Report",
                "locator": "2025_Market_Research_Report.docx, section 4.1",
                "retrieved_at": "2026-08-28T13:50:00Z",
                "as_of_date": "2025-03-15",
                "fact": "The March 2025 baseline recorded 14 active registrants under the working NAICS filter.",
                "limitations": "Baseline dated March 2025",
            },
            {
                "id": "E002",
                "source_class": "federal_mcp",
                "title": "USASpending obligation trend",
                "locator": "Spending-over-time result, NAICS 541513",
                "retrieved_at": "2026-08-28T14:10:00Z",
                "as_of_date": "2026-08-28",
                "fact": "Obligations grew each complete fiscal year from FY2023 through FY2025.",
                "limitations": "Illustrative example data (not live research)",
            },
            {
                "id": "E003",
                "source_class": "federal_mcp",
                "title": "Active registrant population",
                "locator": "Entity search result, NAICS 541513",
                "retrieved_at": "2026-08-28T14:10:00Z",
                "as_of_date": "2026-08-28",
                "fact": "Twelve concerns hold active registrations under the working NAICS filter.",
                "limitations": "Illustrative example data (not live research)",
            },
            {
                "id": "E004",
                "source_class": "user_statement",
                "title": "Program office requirement confirmation",
                "locator": "Refresh intake record",
                "retrieved_at": "2026-08-28T13:45:00Z",
                "as_of_date": "2026-08-28",
                "fact": "The program office confirmed the nationwide remote-support requirement is unchanged.",
                "limitations": "Program statement; not independently verified",
            },
            {
                "id": "E005",
                "source_class": "official_web",
                "title": "Agency forecast of contracting opportunities",
                "locator": "Agency acquisition forecast page",
                "retrieved_at": "2026-08-28T14:30:00Z",
                "as_of_date": "2026-08-28",
                "fact": "The agency forecast continues to list the help desk recompete for FY2027 award.",
                "limitations": "Forecast entries are non-binding",
            },
        ],
        "findings": [
            {
                "id": "F001",
                "text": "The requirement scope is unchanged, so the 2025 requirement description remains usable for planning.",
                "evidence_ids": ["E001"],
            },
            {
                "id": "F002",
                "text": "The registrant population declined from 14 to 12 concerns, so the competition assumption must be rechecked.",
                "evidence_ids": ["E001", "E003"],
            },
        ],
        "inferences": [],
        "user_decisions": [
            "D001: The Contracting Officer deferred the acquisition-strategy decision pending rechecked competition evidence."
        ],
        "conflicts": [],
        "unresolved_questions": [
            "U001: Does the smaller registrant population still support the prior competition approach? Deferred and carried as a limitation of this refresh."
        ],
        "outputs": [],
        "validation": {
            "findings_approved": True,
            "findings_approved_at": "2026-08-28T15:00:00Z",
            "decisions_approved": True,
            "decisions_approved_at": "2026-08-28T15:01:00Z",
            "unresolved_items_disposition_approved": True,
            "unresolved_items_disposition_approved_at": "2026-08-28T15:01:00Z",
            "executive_summary": "The requirement is unchanged and the registrant population contracted from 14 to 12 concerns; the competition assumption must be rechecked before the FY2027 recompete strategy is approved.",
            "change_assessment": "Obligations grew through FY2025 and the registrant population contracted from 14 to 12 concerns.",
            "methodology": "The refresh compared the approved March 2025 baseline against current results under the same filters.",
            "refresh_comparison": [
                {
                    "decision_area": "Demand trend",
                    "prior_baseline": "Obligations flat through FY2024",
                    "current_evidence": "Obligations grew each complete year through FY2025",
                    "evidence_ids": ["E002"],
                    "delta": "Growth replaced a flat trend",
                    "decision_impact": "Budget assumptions should be revisited",
                },
                {
                    "decision_area": "Forecast timing",
                    "prior_baseline": "FY2027 recompete listed in the 2025 forecast",
                    "current_evidence": "The current agency forecast still lists the FY2027 recompete",
                    "evidence_ids": ["E005"],
                    "delta": "No change",
                    "decision_impact": "Planning milestones remain valid",
                },
            ],
            "vendor_landscape_changes": [
                "The active registrant population contracted from 14 to 12 concerns."
            ],
            "strategy_changes": [
                "Recheck the competition assumption before approving the FY2027 recompete strategy."
            ],
            "remains_usable": ["The 2025 requirement description remains usable."],
            "recheck_items": ["The competition assumption built on the 2025 registrant population."],
            "next_actions": [
                {
                    "owner": "Market research lead",
                    "action": "Recheck the competition assumption against current registrant capability evidence.",
                    "output": "Updated competition evidence before strategy approval",
                },
                {
                    "owner": "Contracting Officer",
                    "action": "Review the refreshed comparison and decide whether the prior approach still applies.",
                    "output": "Documented strategy decision for the FY2027 recompete",
                },
            ],
            "decision_implications": [
                "The smaller registrant population means the competition assumption cannot be carried forward without a recheck."
            ],
        },
    }


def complete_record() -> dict:
    return json.loads(
        (ROOT / "tests/fixtures/market-research-record.json").read_text(encoding="utf-8")
    )


def build(record: dict, directory: str) -> Path:
    record_path = Path(directory) / "record.json"
    output = Path(directory) / "report.docx"
    record_path.write_text(json.dumps(record), encoding="utf-8")
    result = subprocess.run(
        [PYTHON, str(SCRIPTS / "build_market_research_report.py"), str(record_path), str(output)],
        capture_output=True,
        text=True,
    )
    if result.returncode:
        raise AssertionError(result.stdout + result.stderr)
    return output


def document_text(document: Document) -> str:
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


class FocusedRouteRenderingTests(unittest.TestCase):
    @classmethod
    def setUpClass(cls):
        cls.directory = tempfile.TemporaryDirectory()
        cls.output = build(refresh_record(), cls.directory.name)
        cls.document = Document(cls.output)
        cls.text = document_text(cls.document)

    @classmethod
    def tearDownClass(cls):
        cls.directory.cleanup()

    def test_no_internal_evidence_class_tokens_render(self):
        for token in ("federal_mcp", "official_web", "user_statement", "other_web", "source_class"):
            self.assertNotIn(token, self.text)
        # E004 (a customer statement) is uncited, so its row is filtered out of
        # the focused register; the cited classes render as reader labels.
        for label in ("Supplied document", "Federal data service", "Official website"):
            self.assertIn(label, self.text)

    def test_reader_visible_evidence_ids_are_sequential(self):
        # The record cites E001, E002, E003, E005 (E004 is uncited), so the
        # rendered register renumbers to a gap-free E001 through E004.
        self.assertIn("E004", self.text)
        self.assertNotIn("E005", self.text)
        register = self.document.tables[-1]
        rendered_ids = [row.cells[0].text.splitlines()[0] for row in register.rows[1:]]
        self.assertEqual(rendered_ids, [f"E{index:03d}" for index in range(1, len(rendered_ids) + 1)])

    def test_closing_action_table_is_cross_referenced_not_duplicated(self):
        action_headers = [
            table
            for table in self.document.tables
            if [cell.text for cell in table.rows[0].cells] == ["Owner", "Action", "Output or gate"]
        ]
        self.assertEqual(len(action_headers), 1)
        self.assertIn("consolidated in the Next practical actions table", self.text)

    def test_focused_route_has_no_forced_page_break(self):
        breaks = [
            br
            for br in self.document.element.body.iter(qn("w:br"))
            if br.get(qn("w:type")) == "page"
        ]
        self.assertEqual(breaks, [])

    def test_report_validator_passes_and_shared_identity_is_present(self):
        record_path = Path(self.directory.name) / "record.json"
        check = subprocess.run(
            [
                PYTHON,
                str(SCRIPTS / "validate_market_research_report.py"),
                str(self.output),
                "--record",
                str(record_path),
            ],
            capture_output=True,
            text=True,
        )
        self.assertEqual(check.returncode, 0, check.stdout + check.stderr)
        header_text = "\n".join(p.text for p in self.document.sections[0].header.paragraphs)
        self.assertIn("1102tools", header_text)
        self.assertEqual(self.document.styles["Normal"].font.name, "Aptos")

    def test_font_table_declares_sans_serif_fallback(self):
        for part in self.document.part.package.iter_parts():
            if str(part.partname) == "/word/fontTable.xml":
                blob = part.blob
                self.assertIn(b'w:name="Aptos"', blob)
                self.assertIn(b'w:altName w:val="Calibri"', blob)
                return
        self.fail("fontTable.xml part is missing")


class CompleteRouteRenderingTests(unittest.TestCase):
    @classmethod
    def setUpClass(cls):
        cls.directory = tempfile.TemporaryDirectory()
        record = complete_record()
        record["validation"]["next_actions"] = [
            {
                "when": "Within 2 weeks",
                "owner": "Market research lead",
                "action": "Issue the market engagement instrument to the candidate vendor population and collect comparable delivery evidence.",
                "output": "Approved evidence-register updates before the strategy decision",
            },
            {
                "when": "Before acquisition strategy",
                "owner": "Contracting Officer",
                "action": "Review the collected capability evidence and confirm whether the competition assumption still holds.",
                "output": "Documented competition assessment for the strategy decision",
            },
        ]
        cls.record = record
        cls.output = build(record, cls.directory.name)
        cls.document = Document(cls.output)
        cls.text = document_text(cls.document)

    @classmethod
    def tearDownClass(cls):
        cls.directory.cleanup()

    def test_no_stray_page_break_paragraph_after_completion_boundary(self):
        # The forced break used to be an empty paragraph carrying only a
        # page-break run, which could land alone on a page and render an
        # entirely blank page. The body page start is now a paragraph property
        # on the first body heading, so no page-break run exists at all.
        breaks = [
            br
            for br in self.document.element.body.iter(qn("w:br"))
            if br.get(qn("w:type")) == "page"
        ]
        self.assertEqual(breaks, [])
        heading = next(
            p for p in self.document.paragraphs if p.text == "Acquisition and decision frame"
        )
        self.assertTrue(heading.paragraph_format.page_break_before)

    def test_execution_plan_cross_references_lead_actions(self):
        lead_tables = [
            table
            for table in self.document.tables
            if [cell.text for cell in table.rows[0].cells] == ["Owner", "Action", "Output or gate"]
        ]
        self.assertEqual(len(lead_tables), 1)
        plan_tables = [
            table
            for table in self.document.tables
            if [cell.text for cell in table.rows[0].cells] == ["When", "Action"]
        ]
        self.assertEqual(len(plan_tables), 1)
        plan_text = "\n".join(
            cell.text for row in plan_tables[0].rows for cell in row.cells
        )
        for item in self.record["validation"]["next_actions"]:
            self.assertIn(item["when"], plan_text)
            self.assertNotIn(item["action"], plan_text)
            self.assertNotIn(item["output"], plan_text)
        self.assertIn("consolidated in the Next practical actions table", self.text)


class ReportValidatorTokenTests(unittest.TestCase):
    def test_report_validator_rejects_rendered_internal_class_token(self):
        with tempfile.TemporaryDirectory() as directory:
            output = build(refresh_record(), directory)
            document = Document(output)
            document.add_paragraph("federal_mcp")
            document.save(output)
            check = subprocess.run(
                [
                    PYTHON,
                    str(SCRIPTS / "validate_market_research_report.py"),
                    str(output),
                    "--record",
                    str(Path(directory) / "record.json"),
                ],
                capture_output=True,
                text=True,
            )
            self.assertNotEqual(check.returncode, 0)
            self.assertIn("internal evidence-class token", check.stdout + check.stderr)


class RecordLanguageValidationTests(unittest.TestCase):
    @classmethod
    def setUpClass(cls):
        cls.validator = load_module(
            SCRIPTS / "validate_research_record.py",
            "market_record_validator_language",
        )

    def failing(self, record: dict) -> list[str]:
        result = self.validator.validate_record(record)
        self.assertEqual(result["status"], "fail")
        return result["failures"]

    def test_clean_focused_record_passes(self):
        result = self.validator.validate_record(refresh_record())
        self.assertEqual(result["status"], "pass", result["failures"])

    def test_session_narration_in_reader_visible_fields_fails(self):
        record = refresh_record()
        record["unresolved_questions"][0] = (
            "U001: Which strategy applies? Deferred by the user and approved for inclusion as a limitation."
        )
        failures = self.failing(record)
        self.assertTrue(any("session or tool narration" in item and "'the user'" in item for item in failures))
        self.assertTrue(any("Illustrative example data (not live research)" in item for item in failures))

    def test_fixture_vocabulary_in_evidence_fails(self):
        record = refresh_record()
        record["evidence"][1]["limitations"] = "Synthetic offline fixture; not a live result"
        failures = self.failing(record)
        self.assertTrue(any("evidence[1].limitations" in item for item in failures))

    def test_narration_check_does_not_apply_to_internal_query_log(self):
        record = refresh_record()
        record["queries"][0]["limitations"] = "Fixture data; FY2026 is partial"
        result = self.validator.validate_record(record)
        self.assertEqual(result["status"], "pass", result["failures"])

    def test_evidence_id_gap_fails(self):
        record = refresh_record()
        record["evidence"] = [item for item in record["evidence"] if item["id"] != "E004"]
        failures = self.failing(record)
        self.assertTrue(any("contiguous starting at E001" in item for item in failures))

    def test_named_firm_reference_without_vendor_names_fails(self):
        record = refresh_record()
        record["validation"]["next_actions"][0]["action"] = "Validate the named firms against current capability evidence."
        failures = self.failing(record)
        self.assertTrue(any("references named firms or vendors" in item for item in failures))

    def test_named_firm_reference_with_named_candidates_passes(self):
        record = refresh_record()
        record["validation"]["next_actions"][0]["action"] = "Validate the named firms against current capability evidence."
        record["validation"]["small_business_candidates"] = [
            {"name": "Example Concern A", "status_and_vehicles": "Active registration"}
        ]
        result = self.validator.validate_record(record)
        self.assertEqual(result["status"], "pass", result["failures"])

    def test_complete_report_fixture_language_rules_do_not_regress(self):
        fixture = json.loads((ROOT / "tests/fixtures/market-research-record.json").read_text(encoding="utf-8"))
        result = self.validator.validate_record(fixture)
        self.assertEqual(result["status"], "pass", result["failures"])
        gapped = copy.deepcopy(fixture)
        gapped["evidence"][3]["id"] = "E009"
        for finding in gapped["findings"] + gapped["inferences"]:
            finding["evidence_ids"] = [
                "E009" if value == "E004" else value for value in finding["evidence_ids"]
            ]
        gap_result = self.validator.validate_record(gapped)
        self.assertEqual(gap_result["status"], "fail")
        self.assertTrue(any("contiguous starting at E001" in item for item in gap_result["failures"]))


class DanglingCitationValidationTests(unittest.TestCase):
    @classmethod
    def setUpClass(cls):
        cls.validator = load_module(
            SCRIPTS / "validate_research_record.py",
            "market_record_validator_citations",
        )

    def test_record_with_resolving_citations_passes(self):
        record = refresh_record()
        record["validation"]["change_assessment"] += " Registrant contraction is documented [E003]."
        result = self.validator.validate_record(record)
        self.assertEqual(result["status"], "pass", result["failures"])

    def test_dangling_text_citation_fails_with_named_identifier(self):
        record = refresh_record()
        record["validation"]["change_assessment"] += " Registrant contraction is documented [E026]."
        result = self.validator.validate_record(record)
        self.assertEqual(result["status"], "fail")
        self.assertTrue(
            any("dangling evidence citations" in item and "E026" in item for item in result["failures"]),
            result["failures"],
        )

    def test_dangling_validation_evidence_ids_entry_fails(self):
        record = refresh_record()
        record["validation"]["refresh_comparison"][0]["evidence_ids"] = ["E026"]
        result = self.validator.validate_record(record)
        self.assertEqual(result["status"], "fail")
        self.assertTrue(
            any("dangling evidence citations" in item and "E026" in item for item in result["failures"]),
            result["failures"],
        )

    def test_report_validator_rejects_dangling_rendered_citation(self):
        with tempfile.TemporaryDirectory() as directory:
            output = build(refresh_record(), directory)
            document = Document(output)
            document.add_paragraph("Registrant contraction is documented [E026].")
            document.save(output)
            check = subprocess.run(
                [
                    PYTHON,
                    str(SCRIPTS / "validate_market_research_report.py"),
                    str(output),
                    "--record",
                    str(Path(directory) / "record.json"),
                ],
                capture_output=True,
                text=True,
            )
            self.assertNotEqual(check.returncode, 0)
            self.assertIn("dangling evidence citations", check.stdout + check.stderr)
            self.assertIn("E026", check.stdout + check.stderr)


class ActionNameTests(unittest.TestCase):
    """The execution-plan Action cell must always read as intelligible English.

    The shipped defect split an action on the period inside a citation and cut
    siblings mid-word, rendering cells such as "Complete the FAR 19",
    "as a...", "every re...", and "the vehicle-acces...".
    """

    @classmethod
    def setUpClass(cls):
        cls.builder = load_module(
            SCRIPTS / "build_market_research_report.py", "market_research_builder"
        )

    def shorten(self, text: str) -> str:
        return self.builder.action_name({"action": text})

    def assertIntelligible(self, text: str) -> str:
        name = self.shorten(text)
        self.assertLessEqual(len(name), self.builder.ACTION_NAME_LIMIT)
        self.assertFalse(name.endswith("...."), name)
        self.assertNotRegex(name, r"\.\.\.\.", name)
        body = name[:-3] if name.endswith("...") else name
        # Never end mid-citation or on a stranded citation lead-in.
        self.assertNotRegex(
            body,
            r"(?i)\b(?:FAR|DFARS|CFR|C\.F\.R\.|U\.S\.C\.|Part|Section|NAICS|PSC)$",
            name,
        )
        # Never end mid-word: the retained text must be a whole-word prefix.
        self.assertTrue(
            " ".join(text.split()).startswith(body) or body in " ".join(text.split()),
            name,
        )
        if body and body != " ".join(text.split()):
            remainder = " ".join(text.split())[len(body) :]
            if remainder:
                self.assertFalse(remainder[0].isalnum(), name)
        return name

    def test_citation_period_is_not_a_sentence_boundary(self):
        action = (
            "Complete the FAR 19.502-2 Rule of Two inquiry using the response evidence "
            "on price, quality, and delivery, and expressly address whether the pool "
            "sustains small business status across the three-year period of performance."
        )
        name = self.assertIntelligible(action)
        self.assertNotEqual(name, "Complete the FAR 19")
        self.assertIn("19.502-2", name)
        self.assertTrue(name.endswith("..."), name)

    def test_shipped_execution_plan_rows_are_not_fragments(self):
        cases = [
            (
                "Issue the ten-theme market-engagement instrument in this report as a "
                "sources-sought notice under NAICS 541512 and PSC DA01, with a 15-day "
                "response period, requesting size status against the 541512 standard "
                "and vehicle holdings as mandatory response elements.",
                "as a",
            ),
            (
                "Re-verify SAM registration currency and exclusion status for every "
                "responding concern, giving specific attention to the Bellese "
                "Technologies registration expiring 2026-12-19.",
                "every re",
            ),
            (
                "Analyze responses against the capability hypotheses and the packaging "
                "hypotheses in this report, and record for each hypothesis whether it "
                "survived, failed, or remains untested.",
                "the packaging",
            ),
            (
                "Determine commerciality on this requirement's facts, and determine "
                "contract type with steady-state operations and quarterly modernization "
                "increments treated as distinct pricing problems.",
                "determine co",
            ),
            (
                "Select the vehicle and competition strategy against the vehicle-access "
                "responses, and complete the consolidation and bundling analysis on this "
                "requirement's own facts.",
                "the vehicle-acces",
            ),
        ]
        for action, old_stub in cases:
            with self.subTest(action=action[:40]):
                name = self.assertIntelligible(action)
                self.assertTrue(name.endswith("..."), name)
                self.assertFalse(name[:-3].endswith(old_stub), name)

    def test_every_truncated_name_carries_one_ellipsis(self):
        action = (
            "Send the same instrument directly to the six verified concerns, and log "
            "each transmittal in the contract file so the record shows the outreach."
        )
        name = self.shorten(action)
        self.assertEqual(name.count("..."), 1)
        self.assertTrue(name.endswith("..."))

    def test_decimal_number_is_not_a_sentence_boundary(self):
        name = self.shorten(
            "Confirm the 7.5 million dollar ceiling with the program office before "
            "the acquisition strategy meeting is scheduled for the quarter."
        )
        self.assertNotEqual(name, "Confirm the 7")
        self.assertIn("7.5 million", name)

    def test_abbreviations_are_not_sentence_boundaries(self):
        for action, fragment in (
            ("Cite U.S. Government contracting data in the record.", "U.S. Government"),
            ("Use No. 5 pricing tables from the schedule.", "No. 5 pricing"),
            ("Confirm the award with Acme, Inc. before the gate.", "Inc."),
            ("Screen e.g. incumbent vendors before the gate.", "e.g. incumbent"),
            ("Apply 13 C.F.R. 121.402 to the size determination.", "13 C.F.R. 121.402"),
        ):
            with self.subTest(action=action):
                self.assertIn(fragment, self.shorten(action))

    def test_short_action_is_returned_whole_without_ellipsis(self):
        name = self.shorten("Issue the sources-sought notice.")
        self.assertEqual(name, "Issue the sources-sought notice")

    def test_first_sentence_is_preferred_when_it_ends_and_fits(self):
        name = self.shorten(
            "Issue the sources-sought notice. Then analyze every response received "
            "against the capability hypotheses recorded in this report."
        )
        self.assertEqual(name, "Issue the sources-sought notice")

    def test_rendered_execution_plan_cells_are_intelligible(self):
        record = complete_record()
        record["validation"]["next_actions"] = [
            {
                "when": "Before the set-aside decision",
                "owner": "Contracting Officer",
                "action": (
                    "Complete the FAR 19.502-2 Rule of Two inquiry using the response "
                    "evidence on price, quality, and delivery, and expressly address "
                    "whether the pool sustains small business status."
                ),
                "output": "Documented Rule of Two determination",
            },
            {
                "when": "Within 5 business days",
                "owner": "Market research lead",
                "action": "Issue the sources-sought notice.",
                "output": "Posted notice",
            },
        ]
        with tempfile.TemporaryDirectory() as directory:
            document = Document(build(record, directory))
            plan = next(
                table
                for table in document.tables
                if [cell.text for cell in table.rows[0].cells] == ["When", "Action"]
            )
            cells = [row.cells[1].text for row in plan.rows[1:]]
            self.assertEqual(len(cells), 2)
            self.assertNotIn("Complete the FAR 19", cells)
            self.assertIn("Issue the sources-sought notice", cells)
            for cell in cells:
                self.assertLessEqual(len(cell), self.builder.ACTION_NAME_LIMIT)
                self.assertFalse(cell.endswith(("FAR", "the", "a", "as a")), cell)


class PlaceholderTimestampTests(unittest.TestCase):
    """Retrieval stamps are per-call facts, not one synthesized batch value."""

    @classmethod
    def setUpClass(cls):
        cls.validator = load_module(
            SCRIPTS / "validate_market_research_report.py", "market_report_validator_stamps"
        )

    def record_with(self, stamps: list[str]) -> dict:
        return {
            "queries": [{"retrieved_at": value} for value in stamps],
            "evidence": [],
        }

    def test_clean_varied_timestamps_pass(self):
        failures = self.validator.placeholder_timestamp_failures(
            self.record_with(
                [
                    "2026-08-28T21:40:00Z",
                    "2026-08-28T21:42:13Z",
                    "2026-08-28T21:47:51Z",
                    "2026-08-28T22:01:09Z",
                    "2026-08-28T22:05:44Z",
                ]
            )
        )
        self.assertEqual(failures, [])

    def test_three_midnight_exact_timestamps_fail(self):
        failures = self.validator.placeholder_timestamp_failures(
            self.record_with(
                ["2026-08-26T00:00:00Z", "2026-08-27T00:00:00Z", "2026-08-28T00:00Z"]
            )
        )
        self.assertTrue(any("midnight-exact placeholders" in item for item in failures), failures)
        self.assertTrue(any("actual retrieval time" in item for item in failures), failures)

    def test_two_midnight_exact_timestamps_pass(self):
        failures = self.validator.placeholder_timestamp_failures(
            self.record_with(["2026-08-26T00:00:00Z", "2026-08-27T00:00:00Z"])
        )
        self.assertEqual(failures, [])

    def test_five_identical_timestamps_fail(self):
        failures = self.validator.placeholder_timestamp_failures(
            self.record_with(["2026-08-28T21:40:00Z"] * 5)
        )
        self.assertTrue(any("are identical" in item for item in failures), failures)
        self.assertTrue(any("actual retrieval time" in item for item in failures), failures)

    def test_four_identical_timestamps_pass(self):
        failures = self.validator.placeholder_timestamp_failures(
            self.record_with(["2026-08-28T21:40:00Z"] * 4)
        )
        self.assertEqual(failures, [])

    def test_gate_counts_evidence_and_query_stamps_together(self):
        record = self.record_with(["2026-08-28T21:40:00Z"] * 3)
        record["evidence"] = [{"retrieved_at": "2026-08-28T21:40:00Z"} for _ in range(2)]
        failures = self.validator.placeholder_timestamp_failures(record)
        self.assertTrue(any("are identical" in item for item in failures), failures)

    def test_report_validator_rejects_identical_batch_stamps(self):
        record = refresh_record()
        stamp = "2026-08-28T21:40:00Z"
        for query in record.get("queries", []):
            query["retrieved_at"] = stamp
        for item in record.get("evidence", []):
            if item.get("retrieved_at"):
                item["retrieved_at"] = stamp
        with tempfile.TemporaryDirectory() as directory:
            output = build(record, directory)
            check = subprocess.run(
                [
                    PYTHON,
                    str(SCRIPTS / "validate_market_research_report.py"),
                    str(output),
                    "--record",
                    str(Path(directory) / "record.json"),
                ],
                capture_output=True,
                text=True,
            )
            self.assertNotEqual(check.returncode, 0)
            self.assertIn("retrieval time", check.stdout + check.stderr)


if __name__ == "__main__":
    unittest.main()
