from __future__ import annotations

import copy
import importlib.util
import json
import shutil
import subprocess
import sys
import tempfile
import unittest
from pathlib import Path


ROOT = Path(__file__).resolve().parents[1]
SKILL = ROOT / "skills" / "acquisition-policy-workflow"
FIXTURE = ROOT / "tests" / "fixtures" / "acquisition-policy-record.json"
PYTHON = sys.executable


def load_module(path: Path, name: str):
    spec = importlib.util.spec_from_file_location(name, path)
    module = importlib.util.module_from_spec(spec)
    assert spec.loader is not None
    spec.loader.exec_module(module)
    return module


class PolicySkillStaticTests(unittest.TestCase):
    def test_menu_and_direct_routing_are_present(self):
        core = (SKILL / "SKILL.md").read_text(encoding="utf-8")
        launch = (SKILL / "references" / "launch-menu-and-framing.md").read_text(encoding="utf-8")
        self.assertIn("An unambiguous request enters its matching mode directly", core)
        self.assertIn("explicit policy-status assertion or instruction", core)
        self.assertIn("routes directly to the relevant status boundary", launch)
        for number in range(1, 11):
            self.assertRegex(core, rf"(?m)^{number}\. ")
            self.assertRegex(launch, rf"(?m)^{number}\. ")
        self.assertIn(
            "Which option would you like? You can reply with the number, label, or your own wording.",
            core,
        )
        self.assertIn(
            "Which option would you like? You can reply with the number, label, or your own wording.",
            launch,
        )

    def test_boundary_language_is_front_loaded(self):
        core = (SKILL / "SKILL.md").read_text(encoding="utf-8")
        self.assertIn("FAR Council model deviation text is not operative for an agency", core)
        self.assertIn("Never describe a proposed rule", core)
        self.assertIn("Public comments are stakeholder evidence", core)
        self.assertIn("Do not substitute direct HTTP", core)
        self.assertIn("ask for the relevant solicitation, award, modification, option, or performance date", core)
        self.assertIn("Do not ask for an audience lens for a status-only answer", core)

    def test_regulations_readiness_is_first_and_discloses_demo_limit(self):
        core = (SKILL / "SKILL.md").read_text(encoding="utf-8")
        self.assertLess(core.index("## Startup data-access readiness"), core.index("## Permanent gates"))
        self.assertIn("regulations-gov", core)
        self.assertIn("get_access_status", core)
        self.assertIn("limited_fallback", core)
        self.assertIn("REGULATIONS_GOV_API_KEY is not configured", core)
        self.assertIn("approximately 10 requests per hour", core)
        self.assertIn("https://1102tools.com/setup#credentials", core)
        self.assertIn("sole pre-approval exception", core)

    def test_route_products_help_and_preview_contract_are_explicit(self):
        core = (SKILL / "SKILL.md").read_text(encoding="utf-8")
        launch = (SKILL / "references" / "launch-menu-and-framing.md").read_text(encoding="utf-8")
        products = [
            "Current Rule Explanation in chat",
            "Documented Agency Policy Status Matrix in chat",
            "Three-Layer Policy Comparison in chat",
            "Regulatory Change Comparison in chat",
            "Rulemaking Timeline in chat",
            "Open Rulemaking Watchlist in chat",
            "Public Comment Position Analysis in chat",
            "Validated Acquisition Policy Impact Brief `.docx`",
            "Refreshed Policy Analysis with a change log",
        ]
        for product in products:
            self.assertIn(product, launch)
        positions = [launch.index(label) for label in ("Recommended outcome:", "Includes:", "Boundary/default:", "Next:")]
        self.assertEqual(positions, sorted(positions))
        self.assertIn("routes directly to the relevant status boundary", launch)
        self.assertIn("no more than these three", launch)
        self.assertIn("recommend exactly one", launch)
        self.assertIn("Never reprint or paraphrase the full menu", launch)
        self.assertIn("Do you want me to proceed with option N using these defaults?", launch)
        self.assertIn("the entire response is exactly the numbered questions below", launch)
        self.assertIn("Do not preface them with a question, repeat a question", launch)
        self.assertIn("first non-whitespace characters must be `Recommended outcome:`", core)
        self.assertIn("routing narration, or code fence", core)
        self.assertIn("before any skill or tool invocation", core)
        self.assertIn("never say `I'm routing` or `I’m routing` before the preview", core)
        self.assertIn("silently read", core)
        self.assertIn("Reuse the retained", core)
        self.assertIn("without reading or loading it again", core)


class PolicyRecordTests(unittest.TestCase):
    @classmethod
    def setUpClass(cls):
        cls.validator = load_module(
            SKILL / "scripts" / "validate_policy_research_record.py",
            "policy_record_validator_tests",
        )

    def fixture(self):
        return json.loads(FIXTURE.read_text(encoding="utf-8"))

    def assert_fails_with(self, record, phrase: str):
        result = self.validator.validate_record(record)
        self.assertEqual(result["status"], "fail")
        self.assertTrue(any(phrase in failure for failure in result["failures"]), result["failures"])

    def test_valid_record(self):
        result = self.validator.validate_record(self.fixture())
        self.assertEqual(result["status"], "pass", result["failures"])

    def test_model_deviation_cannot_be_operative(self):
        record = self.fixture()
        record["policy_items"][1]["operative_for_agency"] = True
        self.assert_fails_with(record, "cannot mark model_deviation operative")

    def test_proposed_rule_cannot_be_operative(self):
        record = self.fixture()
        record["policy_items"][3]["operative_for_agency"] = True
        self.assert_fails_with(record, "cannot mark proposed_rule operative")

    def test_operative_deviation_requires_deviation_evidence(self):
        record = self.fixture()
        record["policy_items"][2]["evidence_ids"] = ["E002"]
        self.assert_fails_with(record, "must cite agency_deviation evidence")

    def test_unknown_evidence_reference_fails(self):
        record = self.fixture()
        record["findings"][0]["evidence_ids"] = ["E999"]
        self.assert_fails_with(record, "unknown evidence IDs")

    def test_sensitive_query_key_fails(self):
        record = self.fixture()
        record["queries"][0]["parameters"]["document_text"] = "nonpublic acquisition text"
        self.assert_fails_with(record, "unsafe parameter keys")

    def test_stakeholder_sample_requires_method_and_limits(self):
        record = self.fixture()
        record["stakeholder_positions"] = [{
            "id": "S001",
            "position": "Supports the proposal",
            "submitter_type": "Association",
            "sample_method": "",
            "reviewed_count": 1,
            "returned_count": 2,
            "evidence_ids": ["E005"],
            "limitations": "",
        }]
        self.assert_fails_with(record, "sample_method must be a non-empty string")
        self.assert_fails_with(record, "limitations must be a non-empty string")

    def conflicting_threshold_record(self):
        record = self.fixture()
        record["conflicts"] = [{
            "id": "C001",
            "issue": "The model text states a $9 million threshold while the agency deviation summary table states $7.5 million.",
            "evidence_ids": ["E002", "E003"],
            "status": "unresolved",
            "resolution": "",
            "resolved_by": "",
            "resolved_at": "",
        }]
        record["findings"][1]["status_resolution"] = "documented_conflict"
        record["findings"][1]["conflict_ids"] = ["C001"]
        record["findings"][1]["text"] = "The cited model text and agency deviation summary state different thresholds; published sources do not resolve the discrepancy."
        record["limitations"].append("An authorized official must resolve the threshold conflict for procurement-specific use.")
        return record

    def test_unresolved_policy_conflict_is_valid_when_reserved(self):
        record = self.conflicting_threshold_record()
        result = self.validator.validate_record(record)
        self.assertEqual(result["status"], "pass", result["failures"])

    def test_unresolved_policy_conflict_rejects_controlling_value_claim(self):
        record = self.conflicting_threshold_record()
        record["findings"][1]["text"] = "The $9 million model-text threshold controls."
        self.assert_fails_with(record, "controlling-value claim")

    def test_unresolved_policy_conflict_cannot_be_marked_authorized_resolution(self):
        record = self.conflicting_threshold_record()
        record["findings"][1]["status_resolution"] = "authorized_resolution"
        self.assert_fails_with(record, "requires official resolution")

    def test_scope_requires_customer_organization(self):
        record = self.fixture()
        del record["scope"]["customer_organization"]
        self.assert_fails_with(record, "scope.customer_organization")
        record["scope"]["customer_organization"] = "   "
        self.assert_fails_with(record, "scope.customer_organization")

    def test_scope_requires_decision_date(self):
        record = self.fixture()
        del record["scope"]["decision_date"]
        self.assert_fails_with(record, "scope.decision_date")
        record["scope"]["decision_date"] = "next month"
        self.assert_fails_with(record, "scope.decision_date")

    def test_refresh_requires_prior_analysis_identity(self):
        record = self.fixture()
        record["workflow_mode"] = "refresh"
        self.assert_fails_with(record, "scope.prior_analysis")
        record["scope"]["prior_analysis"] = {"title": "", "date": "2026-05-01"}
        self.assert_fails_with(record, "prior_analysis.title")
        record["scope"]["prior_analysis"] = {"title": "FAR Part 10 status analysis", "date": "not a date"}
        self.assert_fails_with(record, "prior_analysis.date")
        record["scope"]["prior_analysis"] = {"title": "FAR Part 10 status analysis", "date": "2026-05-01"}
        result = self.validator.validate_record(record)
        self.assertEqual(result["status"], "pass", result["failures"])

    def test_reader_visible_test_harness_vocabulary_is_rejected(self):
        for term in ("synthetic test record", "test fixture", "a synthetic summary", "the fixture represents"):
            record = self.fixture()
            record["evidence"][0]["fact"] = f"This value comes from {term}."
            self.assert_fails_with(record, "test-harness vocabulary")
        record = self.fixture()
        record["timeline"][0]["event"] = "The test fixture represents GSA issuance of the deviation."
        self.assert_fails_with(record, "test-harness vocabulary")
        record = self.fixture()
        record["limitations"].append("Synthetic data only.")
        self.assert_fails_with(record, "test-harness vocabulary")

    def test_non_reader_visible_query_limitations_may_mention_fixtures(self):
        record = self.fixture()
        record["queries"][0]["limitations"] = "Offline test fixture; no live call was made."
        result = self.validator.validate_record(record)
        self.assertEqual(result["status"], "pass", result["failures"])

    def test_illustrative_label_passes(self):
        record = self.fixture()
        record["limitations"].append("Illustrative example (not live data).")
        result = self.validator.validate_record(record)
        self.assertEqual(result["status"], "pass", result["failures"])


class PolicyRenderingTests(unittest.TestCase):
    """Regression coverage for rendered-product defects: duplicated tables, narrow
    gate columns, scope-header propagation, and documented-status vocabulary."""

    @classmethod
    def setUpClass(cls):
        cls.docx_validator = load_module(
            SKILL / "scripts" / "validate_acquisition_policy_brief.py",
            "policy_docx_validator_tests",
        )

    def build_docx(self, record: dict, directory: str) -> Path:
        record_path = Path(directory) / "record.json"
        record_path.write_text(json.dumps(record), encoding="utf-8")
        output = Path(directory) / "product.docx"
        build = subprocess.run(
            [PYTHON, str(SKILL / "scripts" / "build_acquisition_policy_brief.py"), str(record_path), str(output)],
            capture_output=True,
            text=True,
        )
        self.assertEqual(build.returncode, 0, build.stdout + build.stderr)
        return output

    def test_focused_product_does_not_duplicate_gate_rows(self):
        from docx import Document
        from docx.oxml.ns import qn

        record = json.loads(FIXTURE.read_text(encoding="utf-8"))
        record["workflow_mode"] = "current_rule"
        with tempfile.TemporaryDirectory() as directory:
            output = self.build_docx(record, directory)
            document = Document(str(output))
            gate_tables = [
                table
                for table in document.tables
                if table.rows and "Decision-ready evidence" in [cell.text.strip() for cell in table.rows[0].cells]
            ]
            self.assertEqual(len(gate_tables), 1, "gate/action rows must render exactly once")
            text = "\n".join(paragraph.text for paragraph in document.paragraphs)
            self.assertIn("stated once in the Owners and Decision Gates table above", text)
            widths = [int(col.get(qn("w:w"))) for col in gate_tables[0]._tbl.tblGrid]
            self.assertGreaterEqual(widths[0], 1700, "gate label column must not force mid-word breaks")

    def test_scope_header_carries_customer_organization_and_decision_date(self):
        from docx import Document

        record = json.loads(FIXTURE.read_text(encoding="utf-8"))
        record["workflow_mode"] = "three_layer"
        with tempfile.TemporaryDirectory() as directory:
            output = self.build_docx(record, directory)
            document = Document(str(output))
            text = "\n".join(paragraph.text for paragraph in document.paragraphs)
            self.assertIn(record["scope"]["customer_organization"], text)
            self.assertIn(record["scope"]["decision_date"], text)

    def test_refresh_scope_header_identifies_prior_analysis(self):
        from docx import Document

        record = json.loads(FIXTURE.read_text(encoding="utf-8"))
        record["workflow_mode"] = "refresh"
        record["scope"]["prior_analysis"] = {"title": "FAR Part 10 status analysis", "date": "2026-05-01"}
        record["validation"]["refresh_changes"] = [
            {
                "issue": "GSA deviation status",
                "prior_conclusion": "Deviation posted and current",
                "current_evidence": "Deviation remains posted",
                "delta": "Unchanged",
                "planning_consequence": "Continue from the documented baseline",
                "evidence_ids": ["E003"],
            }
        ]
        with tempfile.TemporaryDirectory() as directory:
            output = self.build_docx(record, directory)
            document = Document(str(output))
            text = "\n".join(paragraph.text for paragraph in document.paragraphs)
            self.assertIn("FAR Part 10 status analysis", text)
            self.assertIn("2026-05-01", text)

    def test_three_layer_status_cells_use_allowed_vocabulary(self):
        from docx import Document

        record = json.loads(FIXTURE.read_text(encoding="utf-8"))
        record["workflow_mode"] = "three_layer"
        with tempfile.TemporaryDirectory() as directory:
            output = self.build_docx(record, directory)
            document = Document(str(output))
            self.assertEqual(self.docx_validator.status_vocabulary_failures(document), [])

    def test_agency_status_rows_never_label_rulemaking_as_baseline(self):
        from docx import Document

        record = json.loads(FIXTURE.read_text(encoding="utf-8"))
        record["workflow_mode"] = "agency_status"
        with tempfile.TemporaryDirectory() as directory:
            output = self.build_docx(record, directory)
            document = Document(str(output))
            self.assertEqual(self.docx_validator.status_vocabulary_failures(document), [])
            matrix = next(
                table
                for table in document.tables
                if table.rows and "Status for this question" in [cell.text.strip() for cell in table.rows[0].cells]
            )
            by_layer = {row.cells[0].text.strip(): row.cells[2].text.strip() for row in matrix.rows[1:]}
            self.assertEqual(by_layer["Proposed Rule"], "Pending rulemaking; not current policy")
            self.assertEqual(by_layer["Model Deviation"], "Published model text; not agency-operative")
            self.assertEqual(by_layer["Codified Current"], "Government-wide baseline")

    def test_status_vocabulary_validator_rejects_bad_cells(self):
        from docx import Document

        document = Document()
        table = document.add_table(rows=3, cols=2)
        table.rows[0].cells[0].text = "Layer"
        table.rows[0].cells[1].text = "Documented status"
        table.rows[1].cells[0].text = "Proposed rule"
        table.rows[1].cells[1].text = "Government-wide baseline"
        table.rows[2].cells[0].text = "Agency deviation"
        table.rows[2].cells[1].text = "General Services Administration"
        failures = self.docx_validator.status_vocabulary_failures(document)
        self.assertTrue(any("codified-baseline status" in failure for failure in failures), failures)
        self.assertTrue(any("outside the documented-status vocabulary" in failure for failure in failures), failures)


class PolicyFocusedRouteValidationTests(unittest.TestCase):
    """Regression coverage for per-route validator structure: a focused-route
    product must validate against its own route-native headings, not the full
    Impact Brief section list, and the impact_brief route must keep requiring
    the full twelve-section structure."""

    ROUTE_HEADINGS = {
        "current_rule": ["Current Rule Card"],
        "agency_status": ["Agency Adoption Status"],
        "three_layer": ["Three-Layer Comparison and Adoption Test"],
        "rulemaking": ["Rulemaking Milestones and Next Trigger"],
        "refresh": ["Refresh Change Register"],
    }

    @classmethod
    def setUpClass(cls):
        cls.docx_validator = load_module(
            SKILL / "scripts" / "validate_acquisition_policy_brief.py",
            "policy_docx_validator_route_tests",
        )

    def route_record(self, mode: str) -> dict:
        record = json.loads(FIXTURE.read_text(encoding="utf-8"))
        record["workflow_mode"] = mode
        if mode == "refresh":
            record["scope"]["prior_analysis"] = {"title": "FAR Part 10 status analysis", "date": "2026-05-01"}
            record["validation"]["refresh_changes"] = [
                {
                    "issue": "GSA deviation status",
                    "prior_conclusion": "Deviation posted and current",
                    "current_evidence": "Deviation remains posted",
                    "delta": "Unchanged",
                    "planning_consequence": "Continue from the documented baseline",
                    "evidence_ids": ["E003"],
                }
            ]
        return record

    def build_docx(self, record: dict, directory: str) -> tuple[Path, Path]:
        record_path = Path(directory) / "record.json"
        record_path.write_text(json.dumps(record), encoding="utf-8")
        output = Path(directory) / "product.docx"
        build = subprocess.run(
            [PYTHON, str(SKILL / "scripts" / "build_acquisition_policy_brief.py"), str(record_path), str(output)],
            capture_output=True,
            text=True,
        )
        self.assertEqual(build.returncode, 0, build.stdout + build.stderr)
        return output, record_path

    def test_focused_routes_validate_against_route_native_structure(self):
        from docx import Document

        for mode, payload_headings in self.ROUTE_HEADINGS.items():
            with self.subTest(mode=mode), tempfile.TemporaryDirectory() as directory:
                output, record_path = self.build_docx(self.route_record(mode), directory)
                result = self.docx_validator.validate(output, record_path)
                self.assertEqual(result["status"], "pass", (mode, result["failures"]))
                headings = [
                    paragraph.text.strip()
                    for paragraph in Document(str(output)).paragraphs
                    if getattr(paragraph.style, "name", "") == "Heading 1"
                ]
                for heading in payload_headings:
                    self.assertIn(heading, headings, mode)
                self.assertNotIn("Question and Scope", headings, mode)
                self.assertNotIn("Planning Scenarios", headings, mode)

    def test_focused_route_cli_passes_current_rule_product(self):
        with tempfile.TemporaryDirectory() as directory:
            output, record_path = self.build_docx(self.route_record("current_rule"), directory)
            check = subprocess.run(
                [
                    PYTHON,
                    str(SKILL / "scripts" / "validate_acquisition_policy_brief.py"),
                    str(output),
                    "--record",
                    str(record_path),
                ],
                capture_output=True,
                text=True,
            )
            self.assertEqual(check.returncode, 0, check.stdout + check.stderr)

    def test_impact_brief_route_still_requires_full_structure(self):
        with tempfile.TemporaryDirectory() as directory:
            output, _ = self.build_docx(self.route_record("current_rule"), directory)
            brief_record = json.loads(FIXTURE.read_text(encoding="utf-8"))
            self.assertEqual(brief_record["workflow_mode"], "impact_brief")
            brief_record_path = Path(directory) / "brief-record.json"
            brief_record_path.write_text(json.dumps(brief_record), encoding="utf-8")
            result = self.docx_validator.validate(output, brief_record_path)
            self.assertEqual(result["status"], "fail")
            self.assertTrue(
                any("missing Heading 1 section: Question and Scope" in failure for failure in result["failures"]),
                result["failures"],
            )


class PolicyArtifactTests(unittest.TestCase):
    def test_build_and_validate_brief(self):
        with tempfile.TemporaryDirectory() as directory:
            output = Path(directory) / "acquisition-policy-impact-brief.docx"
            build = subprocess.run(
                [PYTHON, str(SKILL / "scripts" / "build_acquisition_policy_brief.py"), str(FIXTURE), str(output)],
                capture_output=True,
                text=True,
            )
            self.assertEqual(build.returncode, 0, build.stdout + build.stderr)
            check = subprocess.run(
                [
                    PYTHON,
                    str(SKILL / "scripts" / "validate_acquisition_policy_brief.py"),
                    str(output),
                    "--record",
                    str(FIXTURE),
                ],
                capture_output=True,
                text=True,
            )
            self.assertEqual(check.returncode, 0, check.stdout + check.stderr)
            self.assertGreater(output.stat().st_size, 20_000)
            soffice = shutil.which("soffice")
            if soffice:
                converted = subprocess.run(
                    [soffice, "--headless", "--convert-to", "pdf", "--outdir", directory, str(output)],
                    capture_output=True,
                    text=True,
                    timeout=90,
                )
                self.assertEqual(converted.returncode, 0, converted.stdout + converted.stderr)
                self.assertTrue((Path(directory) / f"{output.stem}.pdf").is_file())


if __name__ == "__main__":
    unittest.main()
