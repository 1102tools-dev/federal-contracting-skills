from __future__ import annotations

import importlib.util
import json
import copy
import shutil
import subprocess
import sys
import tempfile
import unittest
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
PYTHON = sys.executable


def load_module(path: Path, name: str):
    spec = importlib.util.spec_from_file_location(name, path)
    module = importlib.util.module_from_spec(spec)
    assert spec.loader is not None
    spec.loader.exec_module(module)
    return module


class SkillStaticTests(unittest.TestCase):
    def test_catalog_and_shared_files(self):
        result = subprocess.run([PYTHON, str(ROOT / "scripts" / "validate_skills.py")], capture_output=True, text=True)
        self.assertEqual(result.returncode, 0, result.stdout + result.stderr)

    def test_mandatory_launch_sequences_are_front_loaded(self):
        market = (ROOT / "skills/market-research-workflow/SKILL.md").read_text(encoding="utf-8")
        growth = (ROOT / "skills/govcon-growth-workflow/SKILL.md").read_text(encoding="utf-8")
        canonical_market_menu = """What would you like to do?

1. Conduct quick market research and show the findings in chat
2. Build a complete FAR Part 10 market research report
3. Refresh or revise an existing market research report
4. Analyze one acquisition question or decision area
5. Prepare market-research findings for the Pre-Award Agent
6. Help me choose"""
        self.assertLess(market.index("## Mandatory first response"), market.index("## Purpose"))
        self.assertGreaterEqual(market.count(canonical_market_menu), 2)
        self.assertIn("Do not summarize it, rename options, omit an option", market)
        self.assertIn("summarized, renamed, reordered, condensed, or incomplete menu is invalid", market)
        self.assertLess(market.index("## Stage 1: launch menu"), market.index("## Stage 2: outcome preview and mandatory document intake"))
        self.assertLess(market.index("## Stage 2: outcome preview and mandatory document intake"), market.index("## Stage 6: capability preflight"))
        self.assertIn("local presence-only SAM.gov `get_access_status` call", market)
        self.assertIn("No upstream research, file generation, capability preflight, web-research request, or other MCP tool invocation occurs first", market)
        self.assertIn("those restrictions never suppress activation", market)
        self.assertIn("Restrictions do not suppress activation", market)
        self.assertIn("never disables this skill or permits a generic answer", market)
        self.assertIn("local presence-only SAM.gov `get_access_status` call", growth)
        self.assertIn("No upstream research, file generation, capability preflight, web-research request, or other MCP tool invocation occurs first", growth)
        self.assertIn("specific opportunity, bid screen, attached notice, company, or desired analysis", growth)
        self.assertIn("Do not replace the menu with intake questions", growth)
        for text in (market, growth):
            self.assertIn("missing_required", text)
            self.assertIn("SAM_API_KEY is not configured", text)
            self.assertIn("https://1102tools.com/setup#credentials", text)
            self.assertIn("do not retry", text.lower())
            self.assertIn("Native web only", text)
            self.assertIn("Native web with Tavily fallback", text)
            self.assertIn("Tavily only", text)
            self.assertIn("No public web", text)
            self.assertGreaterEqual(
                text.count("Which option would you like? You can reply with the number, label, or your own wording."),
                2,
            )
        for number in range(1, 7):
            self.assertRegex(market, rf"(?m)^{number}\. ")
        for number in range(1, 10):
            self.assertRegex(growth, rf"(?m)^{number}\. ")

    def test_route_products_and_post_selection_contract_are_explicit(self):
        cases = {
            "market-research-workflow/references/launch-menu-and-question-blocks.md": [
                "Sourced Market Research Findings in chat",
                "Validated FAR Part 10 Market Research Report `.docx`",
                "Refreshed Market Research Package with a change log",
                "Focused Acquisition Question Analysis in chat",
                "Structured Pre-Award Market Research Handoff in chat",
            ],
            "govcon-growth-workflow/references/launch-menu-and-question-blocks.md": [
                "Federal Opportunity Shortlist in chat",
                "Opportunity Evidence Screen in chat",
                "Competitor/Incumbent Intelligence Profile in chat",
                "Recompete Pipeline in chat",
                "Partner Shortlist or Due-Diligence Profile in chat",
                "Agency/Market Intelligence Snapshot in chat",
                "Labor-Rate/Pricing Context Table in chat",
                "Refreshed Prior Research with a change log",
            ],
        }
        labels = ("Recommended outcome:", "Includes:", "Boundary/default:", "Next:")
        for relative_path, products in cases.items():
            text = (ROOT / "skills" / relative_path).read_text(encoding="utf-8")
            for product in products:
                self.assertIn(product, text)
            positions = [text.index(label) for label in labels]
            self.assertEqual(positions, sorted(positions))
            self.assertIn("no more than these three", text)
            self.assertIn("recommend exactly one", text)
            self.assertIn("Never reprint the full menu", text)
            self.assertIn("Do you want me to proceed with option N using these defaults?", text)
            self.assertIn("the entire response is exactly the numbered questions below", text)
            self.assertIn("Do not preface them with a question, repeat a question", text)

    def test_artifact_preflight_is_phase_adaptive(self):
        component_paths = [
            "sow-pws-builder/SKILL.md",
            "igce-builder-ffp/SKILL.md",
            "igce-builder-lh-tm/SKILL.md",
            "igce-builder-cr/SKILL.md",
            "ot-project-description-builder/SKILL.md",
            "ot-cost-analysis/SKILL.md",
        ]
        for relative_path in component_paths:
            text = (ROOT / "skills" / relative_path).read_text(encoding="utf-8")
            self.assertIn("read-only or artifact-limited session may still", text)
            self.assertIn("Preserve", text)
            self.assertTrue(
                "Before promising or beginning" in text or "before the first dependent MCP call" in text
                or "before its first dependent MCP call" in text,
                relative_path,
            )

    def test_component_routing_fallback_precedes_intake(self):
        component_products = {
            "sow-pws-builder/SKILL.md": "Validated SOW/PWS `.docx` plus two chat-only handoffs",
            "igce-builder-ffp/SKILL.md": "Routed IGCE `.xlsx`, separated by confirmed pricing method or hybrid CLIN",
            "igce-builder-lh-tm/SKILL.md": "Routed IGCE `.xlsx`, separated by confirmed pricing method or hybrid CLIN",
            "igce-builder-cr/SKILL.md": "Routed IGCE `.xlsx`, separated by confirmed pricing method or hybrid CLIN",
            "ot-project-description-builder/SKILL.md": "Validated OT Project Description `.docx` plus chat-only milestone handoff",
            "ot-cost-analysis/SKILL.md": "Milestone-based OT Cost Analysis `.xlsx`",
        }
        labels = ("Recommended outcome:", "Includes:", "Boundary/default:", "Next:")
        for relative_path, product in component_products.items():
            text = (ROOT / "skills" / relative_path).read_text(encoding="utf-8")
            with self.subTest(relative_path=relative_path):
                self.assertIn("routing fallback, not a second preview", text)
                self.assertIn(product, text)
                positions = [text.index(label) for label in labels]
                self.assertEqual(positions, sorted(positions))
                fallback = text.index("When this skill is entered immediately after a numbered")
                self.assertLess(fallback, text.index("begin", fallback))

    def test_post_selection_preview_is_first_visible_text(self):
        for relative_path in (
            "market-research-workflow/SKILL.md",
            "govcon-growth-workflow/SKILL.md",
        ):
            text = (ROOT / "skills" / relative_path).read_text(encoding="utf-8")
            with self.subTest(relative_path=relative_path):
                self.assertIn("first non-whitespace characters must be `Recommended outcome:`", text)
                self.assertIn("routing narration, or code fence", text)
                self.assertIn("before any skill or tool invocation", text)
                self.assertIn("silently read", text)
                self.assertIn("reuse the retained launch reference without reading or loading it again", text)
        for relative_path in (
            "sow-pws-builder/SKILL.md",
            "igce-builder-ffp/SKILL.md",
            "igce-builder-lh-tm/SKILL.md",
            "igce-builder-cr/SKILL.md",
            "ot-project-description-builder/SKILL.md",
            "ot-cost-analysis/SKILL.md",
        ):
            text = (ROOT / "skills" / relative_path).read_text(encoding="utf-8")
            with self.subTest(relative_path=relative_path):
                self.assertIn("Begin line 1 with `Recommended outcome:`", text)
                self.assertIn("routing narration, or code fence", text)

    def test_behavioral_regression_gates_are_explicit(self):
        market = (ROOT / "skills/market-research-workflow/SKILL.md").read_text(encoding="utf-8")
        market_web = (ROOT / "skills/market-research-workflow/references/web-research-method.md").read_text(encoding="utf-8")
        policy = (ROOT / "skills/acquisition-policy-workflow/SKILL.md").read_text(encoding="utf-8")
        ot = (ROOT / "skills/ot-project-description-builder/SKILL.md").read_text(encoding="utf-8")
        self.assertIn("A bare `Approved` does not approve multiple reserved decisions", market)
        for text in (market, market_web):
            self.assertIn("newly discovered URL", text)
            self.assertIn("explicit updated approval", text)
            self.assertIn("Provider fallback", text)
        self.assertLess(market.index("**Exact-URL approval:**"), market.index("## Stage 7: evidence gathering"))
        self.assertIn("record a structured conflict and report `documented_conflict`", policy)
        self.assertIn("Ctrl+A (Cmd+A on Mac), then F9 (or Fn+F9)", ot)

    def test_native_web_is_recommended_and_tavily_requires_explicit_selection(self):
        skill_texts = [
            (ROOT / "skills/market-research-workflow/SKILL.md").read_text(encoding="utf-8"),
            (ROOT / "skills/govcon-growth-workflow/SKILL.md").read_text(encoding="utf-8"),
        ]
        for skill in skill_texts:
            self.assertIn("**Provider-selection hard gate:**", skill)
            self.assertIn("`OK`, `go ahead`, `native`", skill)
            self.assertIn("preserve approved federal MCP", skill)
            self.assertIn("without asking the user to create an account or pay", skill)
            self.assertIn("without paraphrasing", skill)
            self.assertIn("zero, thin, or inconclusive results do not", skill)
        policies = [
            (ROOT / "skills/market-research-workflow/references/web-provider-policy.md").read_text(encoding="utf-8"),
            (ROOT / "skills/govcon-growth-workflow/references/web-provider-policy.md").read_text(encoding="utf-8"),
        ]
        for policy in policies:
            choices = (
                "1. **Native web only (Recommended):**",
                "2. **Native web with Tavily fallback:**",
                "3. **Tavily only:**",
                "4. **No public web:**",
            )
            positions = [policy.index(choice) for choice in choices]
            self.assertEqual(positions, sorted(positions))
            self.assertNotIn("Tavily with native fallback (Recommended)", policy)
            self.assertIn("Never request payment, create an account, or switch providers for the user", policy)
            self.assertIn("stop and obtain new approval rather than proceeding", policy)
            self.assertIn("an ambiguous response such as `native` does not select a mode", policy)
            self.assertIn("Zero results, thin or inconclusive results", policy)
            self.assertIn("not local or private browsing", policy)
            self.assertIn("Approved federal MCP calls and supplied-document analysis remain permitted", policy)

    def test_spreadsheet_host_precedence_and_fallback_are_explicit(self):
        paths = [
            ROOT / "skills/igce-builder-cr/references/runtime-adaptation.md",
            ROOT / "skills/igce-builder-ffp/references/runtime-adaptation.md",
            ROOT / "skills/igce-builder-lh-tm/references/runtime-adaptation.md",
            ROOT / "skills/ot-cost-analysis/references/runtime-adaptation.md",
        ]
        for path in paths:
            text = path.read_text(encoding="utf-8")
            with self.subTest(path=path.name):
                self.assertIn("host", text.lower())
                self.assertIn("hard stop", text.lower())
                self.assertIn("structured JSON", text)
                self.assertIn("Markdown or CSV", text)
                self.assertIn("Do not label the fallback as a completed workbook", text)

    def test_pricing_skills_disclose_keyless_limits_before_data_calls(self):
        paths = [
            ROOT / "skills/igce-builder-cr/SKILL.md",
            ROOT / "skills/igce-builder-ffp/SKILL.md",
            ROOT / "skills/igce-builder-lh-tm/SKILL.md",
            ROOT / "skills/ot-cost-analysis/SKILL.md",
        ]
        for path in paths:
            text = path.read_text(encoding="utf-8")
            with self.subTest(path=path.name):
                self.assertIn("bls-oews.get_access_status", text)
                self.assertIn("BLS_API_KEY", text)
                self.assertIn("25 requests per day", text)
                self.assertIn("10 years per query", text)
                self.assertIn("gsa-perdiem.get_access_status", text)
                self.assertIn("PERDIEM_API_KEY", text)
                self.assertIn("approximately 10 requests per hour", text)
                self.assertIn("configured_unverified", text)
                self.assertIn("https://1102tools.com/setup#credentials", text)


class RecordValidationTests(unittest.TestCase):
    @classmethod
    def setUpClass(cls):
        cls.validator = load_module(
            ROOT / "skills/market-research-workflow/scripts/validate_research_record.py",
            "research_record_validator",
        )

    def fixture(self, name: str):
        return json.loads((ROOT / "tests/fixtures" / name).read_text(encoding="utf-8"))

    def test_valid_market_and_growth_records(self):
        for name in ("market-research-record.json", "govcon-growth-record.json"):
            result = self.validator.validate_record(self.fixture(name))
            self.assertEqual(result["status"], "pass", result["failures"])

    def test_legacy_market_skill_identifier_remains_valid_during_rc_transition(self):
        record = self.fixture("market-research-record.json")
        record["skill"] = "market-research-builder"
        result = self.validator.validate_record(record)
        self.assertEqual(result["status"], "pass", result["failures"])

    def test_schema_11_market_record_is_readable_but_not_artifact_ready(self):
        record = self.fixture("market-research-record.json")
        record["schema_version"] = "1.1"
        read_result = self.validator.validate_record(record, purpose="read")
        self.assertEqual(read_result["status"], "pass", read_result["failures"])
        artifact_result = self.validator.validate_record(record)
        self.assertEqual(artifact_result["status"], "fail")
        self.assertTrue(any("schema_version must be 1.2" in item for item in artifact_result["failures"]))

    def test_complete_report_requires_explicit_decision_and_unresolved_disposition_approvals(self):
        record = self.fixture("market-research-record.json")
        record["validation"]["decisions_approved"] = False
        record["validation"]["unresolved_items_disposition_approved"] = False
        result = self.validator.validate_record(record)
        self.assertEqual(result["status"], "fail")
        self.assertTrue(any("decisions_approved must be true" in item for item in result["failures"]))
        self.assertTrue(any("unresolved_items_disposition_approved must be true" in item for item in result["failures"]))

    def test_complete_report_requires_stable_decision_and_unresolved_ids(self):
        record = self.fixture("market-research-record.json")
        record["user_decisions"] = ["Approved"]
        record["unresolved_questions"] = ["Which strategy applies?"]
        result = self.validator.validate_record(record)
        self.assertEqual(result["status"], "fail")
        self.assertTrue(any("stable D###" in item for item in result["failures"]))
        self.assertTrue(any("stable U###" in item for item in result["failures"]))

    def test_unknown_evidence_reference_fails(self):
        record = self.fixture("market-research-record.json")
        record["findings"][0]["evidence_ids"] = ["E999"]
        result = self.validator.validate_record(record)
        self.assertEqual(result["status"], "fail")
        self.assertTrue(any("unknown evidence" in item for item in result["failures"]))

    def test_sensitive_query_key_fails(self):
        record = self.fixture("market-research-record.json")
        record["queries"][0]["parameters"]["document_text"] = "private text"
        result = self.validator.validate_record(record)
        self.assertEqual(result["status"], "fail")
        self.assertTrue(any("unsafe parameter" in item for item in result["failures"]))

    def test_credential_pattern_fails(self):
        record = self.fixture("govcon-growth-record.json")
        record["assumptions"].append("token=abcdefghijklmnopqrstuvwxyz123456")
        result = self.validator.validate_record(record)
        self.assertEqual(result["status"], "fail")
        self.assertTrue(any("credential" in item for item in result["failures"]))

    def web_fixture(self, mode: str, planned: list[str], used: list[str] | None = None):
        record = self.fixture("market-research-record.json")
        record["web_research"] = {
            "mode": mode,
            "approved": True,
            "approved_at": "2026-08-21T19:00:00Z",
            "disclosure_acknowledged": True,
            "planned_providers": planned,
            "providers_used": list(used or []),
            "fallback_events": [],
        }
        return record

    def test_all_provider_modes_validate(self):
        modes = {
            "native_only": ["native_web"],
            "native_with_tavily_fallback": ["native_web", "tavily"],
            "tavily_only": ["tavily"],
            "no_public_web": [],
        }
        for mode, providers in modes.items():
            with self.subTest(mode=mode):
                record = self.web_fixture(mode, providers)
                result = self.validator.validate_record(record)
                self.assertEqual(result["status"], "pass", result["failures"])

    def test_provider_mode_and_order_are_exact(self):
        reversed_combined = self.web_fixture(
            "native_with_tavily_fallback",
            ["tavily", "native_web"],
        )
        reversed_result = self.validator.validate_record(reversed_combined)
        self.assertEqual(reversed_result["status"], "fail")
        self.assertTrue(any("provider order" in item for item in reversed_result["failures"]))

        unknown = self.web_fixture("automatic", ["native_web"])
        unknown_result = self.validator.validate_record(unknown)
        self.assertEqual(unknown_result["status"], "fail")
        self.assertTrue(any("mode is not approved" in item for item in unknown_result["failures"]))

    def test_combined_mode_records_approved_fallback(self):
        record = self.web_fixture(
            "native_with_tavily_fallback",
            ["native_web", "tavily"],
            ["native_web", "tavily"],
        )
        record["web_research"]["fallback_events"] = [{
            "timestamp": "2026-08-21T19:05:00Z",
            "failed_provider": "native_web",
            "replacement_provider": "tavily",
            "reason": "rate_limited",
        }]
        record["queries"].append({
            "id": "Q002",
            "provider": "tavily",
            "operation": "tavily_search",
            "parameters": {"query": "official federal market research guidance"},
            "retrieved_at": "2026-08-21T19:06:00Z",
            "count": 3,
            "limitations": "Synthetic fallback fixture",
        })
        result = self.validator.validate_record(record)
        self.assertEqual(result["status"], "pass", result["failures"])

    def test_combined_mode_rejects_tavily_to_native_fallback_direction(self):
        record = self.web_fixture(
            "native_with_tavily_fallback",
            ["native_web", "tavily"],
            ["native_web", "tavily"],
        )
        record["web_research"]["fallback_events"] = [{
            "timestamp": "2026-08-21T19:05:00Z",
            "failed_provider": "tavily",
            "replacement_provider": "native_web",
            "reason": "Synthetic reversed fallback",
        }]
        result = self.validator.validate_record(record)
        self.assertEqual(result["status"], "fail")
        self.assertTrue(any("native_web-to-tavily" in item for item in result["failures"]))

    def test_combined_mode_rejects_nonfailure_fallback_reasons(self):
        for reason in ("zero_results", "thin_results", "user_declined_permission", "content_refusal"):
            with self.subTest(reason=reason):
                record = self.web_fixture(
                    "native_with_tavily_fallback",
                    ["native_web", "tavily"],
                    ["native_web", "tavily"],
                )
                record["web_research"]["fallback_events"] = [{
                    "timestamp": "2026-08-21T19:05:00Z",
                    "failed_provider": "native_web",
                    "replacement_provider": "tavily",
                    "reason": reason,
                }]
                result = self.validator.validate_record(record)
                self.assertEqual(result["status"], "fail")
                self.assertTrue(any("approved native failure class" in item for item in result["failures"]))

    def test_legacy_tavily_first_combined_mode_is_readable_but_not_artifact_ready(self):
        record = self.web_fixture(
            "tavily_with_native_fallback",
            ["tavily", "native_web"],
            ["tavily"],
        )
        read_result = self.validator.validate_record(record, purpose="read")
        self.assertEqual(read_result["status"], "pass", read_result["failures"])
        artifact_result = self.validator.validate_record(record)
        self.assertEqual(artifact_result["status"], "fail")
        self.assertTrue(any("mode is not approved" in item for item in artifact_result["failures"]))

    def test_unapproved_provider_fails(self):
        record = self.web_fixture("native_only", ["native_web"], ["native_web"])
        record["queries"].append({
            "id": "Q002",
            "provider": "tavily",
            "operation": "tavily_search",
            "parameters": {"query": "public query"},
            "retrieved_at": "2026-08-21T19:06:00Z",
            "count": 1,
            "limitations": "Synthetic fixture",
        })
        result = self.validator.validate_record(record)
        self.assertEqual(result["status"], "fail")
        self.assertTrue(any("not approved" in item for item in result["failures"]))

    def test_fallback_event_requires_combined_mode_and_provider_switch(self):
        record = self.web_fixture("tavily_only", ["tavily"], ["tavily"])
        record["web_research"]["fallback_events"] = [{
            "timestamp": "2026-08-21T19:05:00Z",
            "failed_provider": "tavily",
            "replacement_provider": "tavily",
            "reason": "Synthetic invalid fallback",
        }]
        result = self.validator.validate_record(record)
        self.assertEqual(result["status"], "fail")
        self.assertTrue(any("allowed only" in item for item in result["failures"]))
        self.assertTrue(any("must switch providers" in item for item in result["failures"]))

    def test_unapproved_or_unacknowledged_plan_fails(self):
        record = self.web_fixture("tavily_only", ["tavily"])
        record["web_research"]["approved"] = False
        record["web_research"]["disclosure_acknowledged"] = False
        result = self.validator.validate_record(record)
        self.assertEqual(result["status"], "fail")
        self.assertTrue(any("approved must be true" in item for item in result["failures"]))
        self.assertTrue(any("disclosure_acknowledged" in item for item in result["failures"]))

    def test_private_and_signed_urls_fail(self):
        base = self.web_fixture("tavily_only", ["tavily"], ["tavily"])
        bad_urls = (
            "file:///Users/example/private.pdf",
            "http://127.0.0.1/report",
            "https://intranet.internal/report",
            "https://example.com/report?token=secret-value",
        )
        for url in bad_urls:
            with self.subTest(url=url):
                record = copy.deepcopy(base)
                record["queries"].append({
                    "id": "Q002",
                    "provider": "tavily",
                    "operation": "tavily_extract",
                    "parameters": {"urls": [url]},
                    "retrieved_at": "2026-08-21T19:06:00Z",
                    "count": 0,
                    "limitations": "Synthetic safety fixture",
                })
                result = self.validator.validate_record(record)
                self.assertEqual(result["status"], "fail")
                self.assertTrue(any("parameters.urls" in item for item in result["failures"]))

    def test_public_extraction_url_passes(self):
        record = self.web_fixture("tavily_only", ["tavily"], ["tavily"])
        record["queries"].append({
            "id": "Q002",
            "provider": "tavily",
            "operation": "tavily_extract",
            "parameters": {"urls": ["https://www.acquisition.gov/far/part-10"]},
            "retrieved_at": "2026-08-21T19:06:00Z",
            "count": 1,
            "limitations": "Synthetic safety fixture",
        })
        result = self.validator.validate_record(record)
        self.assertEqual(result["status"], "pass", result["failures"])

    def test_nonapproved_tavily_tool_fails(self):
        record = self.web_fixture("tavily_only", ["tavily"], ["tavily"])
        record["queries"].append({
            "id": "Q002",
            "provider": "tavily",
            "operation": "tavily_research",
            "parameters": {"query": "public query"},
            "retrieved_at": "2026-08-21T19:06:00Z",
            "count": 0,
            "limitations": "Synthetic prohibited-tool fixture",
        })
        result = self.validator.validate_record(record)
        self.assertEqual(result["status"], "fail")
        self.assertTrue(any("prohibited Tavily operation" in item for item in result["failures"]))

    def test_growth_source_timestamp_must_match_linked_source_call(self):
        growth_validator = load_module(
            ROOT / "skills/govcon-growth-workflow/scripts/validate_research_record.py",
            "growth_record_validator_timestamp",
        )
        record = self.fixture("govcon-growth-record.json")
        record["evidence"][-1]["retrieved_at"] = "2026-08-21T19:00:00Z"
        result = growth_validator.validate_record(record)
        self.assertEqual(result["status"], "fail")
        self.assertTrue(any("linked source-call timestamp" in item for item in result["failures"]))


class ArtifactTests(unittest.TestCase):
    def build_and_validate(self, skill: str, record_name: str, builder: str, validator: str, output_name: str):
        with tempfile.TemporaryDirectory() as directory:
            output = Path(directory) / output_name
            record = ROOT / "tests/fixtures" / record_name
            build = subprocess.run(
                [PYTHON, str(ROOT / "skills" / skill / "scripts" / builder), str(record), str(output)],
                capture_output=True,
                text=True,
            )
            self.assertEqual(build.returncode, 0, build.stdout + build.stderr)
            check = subprocess.run(
                [PYTHON, str(ROOT / "skills" / skill / "scripts" / validator), str(output), "--record", str(record)],
                capture_output=True,
                text=True,
            )
            self.assertEqual(check.returncode, 0, check.stdout + check.stderr)
            self.assertGreater(output.stat().st_size, 10_000)
            soffice = shutil.which("soffice")
            if soffice:
                converted = subprocess.run(
                    [soffice, "--headless", "--convert-to", "pdf", "--outdir", directory, str(output)],
                    capture_output=True,
                    text=True,
                    timeout=90,
                )
                self.assertEqual(converted.returncode, 0, converted.stdout + converted.stderr)
                self.assertTrue((Path(directory) / f"{output.stem}.pdf").is_file())

    def test_market_report(self):
        self.build_and_validate(
            "market-research-workflow",
            "market-research-record.json",
            "build_market_research_report.py",
            "validate_market_research_report.py",
            "market-research.docx",
        )

    def test_market_report_numeric_check_must_cite_linked_evidence_id(self):
        with tempfile.TemporaryDirectory() as directory:
            output = Path(directory) / "market-research.docx"
            record = ROOT / "tests/fixtures/market-research-record.json"
            builder = ROOT / "skills/market-research-workflow/scripts/build_market_research_report.py"
            validator_path = ROOT / "skills/market-research-workflow/scripts/validate_market_research_report.py"
            build = subprocess.run([PYTHON, str(builder), str(record), str(output)], capture_output=True, text=True)
            self.assertEqual(build.returncode, 0, build.stdout + build.stderr)

            document = Document(output)
            calculation_lines = [
                paragraph
                for paragraph in document.paragraphs
                if "Complete-year fixture obligations" in paragraph.text
            ]
            self.assertEqual(len(calculation_lines), 1)
            self.assertIn("[E004]", calculation_lines[0].text)
            for run in calculation_lines[0].runs:
                if "E004" in run.text:
                    run.text = run.text.replace("E004", "calculation")
            document.save(output)

            check = subprocess.run(
                [PYTHON, str(validator_path), str(output), "--record", str(record)],
                capture_output=True,
                text=True,
            )
            self.assertNotEqual(check.returncode, 0)
            self.assertIn("does not cite its calculation evidence ID: E004", check.stdout + check.stderr)

    def test_market_report_builder_rejects_unlinked_numeric_check(self):
        with tempfile.TemporaryDirectory() as directory:
            record = json.loads((ROOT / "tests/fixtures/market-research-record.json").read_text(encoding="utf-8"))
            record["evidence"][3]["locator"] = "validation.numeric_checks[99]"
            record_path = Path(directory) / "unlinked-record.json"
            record_path.write_text(json.dumps(record), encoding="utf-8")
            output = Path(directory) / "market-research.docx"
            builder = ROOT / "skills/market-research-workflow/scripts/build_market_research_report.py"
            build = subprocess.run(
                [PYTHON, str(builder), str(record_path), str(output)],
                capture_output=True,
                text=True,
            )
            self.assertNotEqual(build.returncode, 0)
            self.assertIn("requires exactly one calculation evidence item", build.stdout + build.stderr)
            self.assertFalse(output.exists())

    def test_growth_brief(self):
        self.build_and_validate(
            "govcon-growth-workflow",
            "govcon-growth-record.json",
            "build_growth_brief.py",
            "validate_growth_brief.py",
            "growth-brief.docx",
        )

    def test_growth_brief_labels_no_public_web_and_formats_currency_scope(self):
        with tempfile.TemporaryDirectory() as directory:
            record = json.loads((ROOT / "tests/fixtures/govcon-growth-record.json").read_text(encoding="utf-8"))
            record["scope"]["estimated_total_value_usd"] = 18_000_000
            record_path = Path(directory) / "growth-record.json"
            output = Path(directory) / "growth-brief.docx"
            record_path.write_text(json.dumps(record), encoding="utf-8")
            builder = ROOT / "skills/govcon-growth-workflow/scripts/build_growth_brief.py"
            build = subprocess.run([PYTHON, str(builder), str(record_path), str(output)], capture_output=True, text=True)
            self.assertEqual(build.returncode, 0, build.stdout + build.stderr)
            text = "\n".join(paragraph.text for paragraph in Document(output).paragraphs)
            # The fixture logs one federal source call, so declining public web
            # research must not render a supplied-only claim.
            self.assertIn(
                "Live federal data research with supplied company context | No public web research performed",
                text,
            )
            self.assertNotIn("Supplied evidence only | No public research performed", text)
            values = [cell.text for table in Document(output).tables for row in table.rows for cell in row.cells]
            self.assertIn("$18,000,000", values)
            self.assertNotIn("Estimated Total Value Usd", values)

    def test_growth_brief_numeric_check_must_cite_linked_evidence_id(self):
        with tempfile.TemporaryDirectory() as directory:
            output = Path(directory) / "growth-brief.docx"
            record = ROOT / "tests/fixtures/govcon-growth-record.json"
            builder = ROOT / "skills/govcon-growth-workflow/scripts/build_growth_brief.py"
            validator_path = ROOT / "skills/govcon-growth-workflow/scripts/validate_growth_brief.py"
            build = subprocess.run([PYTHON, str(builder), str(record), str(output)], capture_output=True, text=True)
            self.assertEqual(build.returncode, 0, build.stdout + build.stderr)

            document = Document(output)
            calculation_lines = [
                paragraph
                for paragraph in document.paragraphs
                if "Sample pipeline value" in paragraph.text
            ]
            self.assertEqual(len(calculation_lines), 1)
            self.assertIn("[E003]", calculation_lines[0].text)
            for run in calculation_lines[0].runs:
                if "E003" in run.text:
                    run.text = run.text.replace("E003", "calculation")
            document.save(output)

            check = subprocess.run(
                [PYTHON, str(validator_path), str(output), "--record", str(record)],
                capture_output=True,
                text=True,
            )
            self.assertNotEqual(check.returncode, 0)
            self.assertIn("does not cite its calculation evidence ID: E003", check.stdout + check.stderr)

    def test_growth_brief_builder_rejects_unlinked_numeric_check(self):
        with tempfile.TemporaryDirectory() as directory:
            record = json.loads((ROOT / "tests/fixtures/govcon-growth-record.json").read_text(encoding="utf-8"))
            record["evidence"][2]["locator"] = "validation.numeric_checks[99]"
            record_path = Path(directory) / "unlinked-record.json"
            record_path.write_text(json.dumps(record), encoding="utf-8")
            output = Path(directory) / "growth-brief.docx"
            builder = ROOT / "skills/govcon-growth-workflow/scripts/build_growth_brief.py"
            build = subprocess.run(
                [PYTHON, str(builder), str(record_path), str(output)],
                capture_output=True,
                text=True,
            )
            self.assertNotEqual(build.returncode, 0)
            self.assertIn("requires exactly one calculation evidence item", build.stdout + build.stderr)
            self.assertFalse(output.exists())

    def growth_paths(self):
        return (
            ROOT / "skills/govcon-growth-workflow/scripts/build_growth_brief.py",
            ROOT / "skills/govcon-growth-workflow/scripts/validate_growth_brief.py",
        )

    def build_growth_variant(self, directory: str, mutate) -> tuple[Path, Path]:
        record = json.loads((ROOT / "tests/fixtures/govcon-growth-record.json").read_text(encoding="utf-8"))
        mutate(record)
        record_path = Path(directory) / "growth-record.json"
        record_path.write_text(json.dumps(record), encoding="utf-8")
        output = Path(directory) / "growth-brief.docx"
        builder, _ = self.growth_paths()
        build = subprocess.run([PYTHON, str(builder), str(record_path), str(output)], capture_output=True, text=True)
        self.assertEqual(build.returncode, 0, build.stdout + build.stderr)
        return record_path, output

    def test_growth_brief_rejects_reader_visible_harness_vocabulary(self):
        _, validator_path = self.growth_paths()
        for phrase in ("Fictional internal company context", "Archived bounded-sample record"):
            with self.subTest(phrase=phrase), tempfile.TemporaryDirectory() as directory:
                record_path, output = self.build_growth_variant(
                    directory, lambda record: record["assumptions"].append(phrase)
                )
                check = subprocess.run(
                    [PYTHON, str(validator_path), str(output), "--record", str(record_path)],
                    capture_output=True,
                    text=True,
                )
                self.assertNotEqual(check.returncode, 0)
                self.assertIn("reader-visible harness vocabulary is prohibited", check.stdout + check.stderr)

    def test_growth_brief_rejects_federal_evidence_without_locator(self):
        _, validator_path = self.growth_paths()
        with tempfile.TemporaryDirectory() as directory:
            def mutate(record):
                self.assertEqual(record["evidence"][3]["source_class"], "federal_mcp")
                record["evidence"][3]["locator"] = " "
            record_path, output = self.build_growth_variant(directory, mutate)
            check = subprocess.run(
                [PYTHON, str(validator_path), str(output), "--record", str(record_path)],
                capture_output=True,
                text=True,
            )
            self.assertNotEqual(check.returncode, 0)
            self.assertIn("federal evidence E004 has no checkable locator", check.stdout + check.stderr)

    def validate_growth(self, output: Path, record_path: Path) -> subprocess.CompletedProcess:
        _, validator_path = self.growth_paths()
        return subprocess.run(
            [PYTHON, str(validator_path), str(output), "--record", str(record_path)],
            capture_output=True,
            text=True,
        )

    def mutate_growth_record(self, record_path: Path, mutate) -> None:
        # Apply a post-build mutation so the brief validator sees a record the
        # builder's own record validation would have rejected.
        record = json.loads(record_path.read_text(encoding="utf-8"))
        mutate(record)
        record_path.write_text(json.dumps(record), encoding="utf-8")

    def rewrite_growth_basis_line(self, output: Path, new_text: str) -> None:
        document = Document(output)
        basis = [
            paragraph
            for paragraph in document.paragraphs
            if paragraph.text.startswith("As of ") and " | " in paragraph.text
        ]
        self.assertEqual(len(basis), 1)
        for index, run in enumerate(basis[0].runs):
            run.text = new_text if index == 0 else ""
        document.save(output)

    def test_growth_brief_rejects_supplied_only_claim_with_live_calls(self):
        with tempfile.TemporaryDirectory() as directory:
            record_path, output = self.build_growth_variant(directory, lambda record: None)
            self.rewrite_growth_basis_line(
                output, "As of 2026-08-21 | Supplied evidence only | No public research performed"
            )
            check = self.validate_growth(output, record_path)
            self.assertNotEqual(check.returncode, 0)
            self.assertIn(
                "evidence-basis line claims 'Supplied evidence only'"
                " but the research record logs 1 live source call(s)",
                check.stdout + check.stderr,
            )

    def test_growth_brief_rejects_live_research_claim_with_empty_call_log(self):
        with tempfile.TemporaryDirectory() as directory:
            record_path, output = self.build_growth_variant(directory, lambda record: None)
            self.mutate_growth_record(record_path, lambda record: record.update(queries=[]))
            self.rewrite_growth_basis_line(
                output, "As of 2026-08-21 | Public-source research with supplied company context"
            )
            check = self.validate_growth(output, record_path)
            self.assertNotEqual(check.returncode, 0)
            self.assertIn(
                "evidence-basis line claims live research but the research record logs no source calls",
                check.stdout + check.stderr,
            )

    def test_growth_brief_accepts_consistent_evidence_basis_lines(self):
        with self.subTest(description="live calls with a live-research line"), \
                tempfile.TemporaryDirectory() as directory:
            record_path, output = self.build_growth_variant(directory, lambda record: None)
            check = self.validate_growth(output, record_path)
            self.assertEqual(check.returncode, 0, check.stdout + check.stderr)
        with self.subTest(description="no calls with a supplied-only line"), \
                tempfile.TemporaryDirectory() as directory:
            record_path, output = self.build_growth_variant(directory, lambda record: None)
            self.mutate_growth_record(record_path, lambda record: record.update(queries=[]))
            self.rewrite_growth_basis_line(
                output, "As of 2026-08-21 | Supplied evidence only | No public research performed"
            )
            check = self.validate_growth(output, record_path)
            self.assertEqual(check.returncode, 0, check.stdout + check.stderr)

    def test_growth_brief_rejects_all_midnight_retrieval_timestamps(self):
        with tempfile.TemporaryDirectory() as directory:
            record_path, output = self.build_growth_variant(directory, lambda record: None)

            def mutate(record):
                record["queries"][0]["retrieved_at"] = "2026-08-21T00:00:00Z"
                for item in record["evidence"]:
                    item["retrieved_at"] = "2026-08-21T00:00:00Z"
            self.mutate_growth_record(record_path, mutate)
            check = self.validate_growth(output, record_path)
            self.assertNotEqual(check.returncode, 0)
            self.assertIn(
                "all 5 retrieval timestamps are midnight-exact placeholders;"
                " record actual retrieval times for each source call",
                check.stdout + check.stderr,
            )

    def test_growth_brief_accepts_single_midnight_timestamp_among_real_times(self):
        with tempfile.TemporaryDirectory() as directory:
            record_path, output = self.build_growth_variant(directory, lambda record: None)
            self.mutate_growth_record(
                record_path,
                lambda record: record["queries"][0].update(retrieved_at="2026-08-21T00:00:00Z"),
            )
            check = self.validate_growth(output, record_path)
            self.assertEqual(check.returncode, 0, check.stdout + check.stderr)

    def test_growth_brief_accepts_two_midnight_timestamps_below_row_threshold(self):
        with tempfile.TemporaryDirectory() as directory:
            record_path, output = self.build_growth_variant(directory, lambda record: None)

            def mutate(record):
                record["queries"][0]["retrieved_at"] = "2026-08-21T00:00:00Z"
                for item in record["evidence"]:
                    if item["source_class"] == "federal_mcp":
                        item["retrieved_at"] = "2026-08-21T00:00:00Z"
                    else:
                        item["retrieved_at"] = ""
            self.mutate_growth_record(record_path, mutate)
            check = self.validate_growth(output, record_path)
            self.assertEqual(check.returncode, 0, check.stdout + check.stderr)

    def test_growth_brief_later_sections_do_not_restate_page_one_lists(self):
        with tempfile.TemporaryDirectory() as directory:
            _, output = self.build_growth_variant(directory, lambda record: None)
            paragraphs = [paragraph.text for paragraph in Document(output).paragraphs]
            record = json.loads((ROOT / "tests/fixtures/govcon-growth-record.json").read_text(encoding="utf-8"))
            first_action = record["validation"]["next_actions"][0]
            first_unknown = record["validation"]["missing_bid_context"][0]
            self.assertEqual(sum(1 for text in paragraphs if first_action in text), 1)
            self.assertEqual(sum(1 for text in paragraphs if first_unknown in text), 1)

    def test_growth_brief_long_evidence_table_does_not_repeat_header(self):
        with tempfile.TemporaryDirectory() as directory:
            output = Path(directory) / "growth-brief.docx"
            record = ROOT / "tests/fixtures/govcon-growth-record.json"
            builder = ROOT / "skills/govcon-growth-workflow/scripts/build_growth_brief.py"
            build = subprocess.run([PYTHON, str(builder), str(record), str(output)], capture_output=True, text=True)
            self.assertEqual(build.returncode, 0, build.stdout + build.stderr)

            document = Document(output)
            evidence_table = document.tables[-1]
            header_properties = evidence_table.rows[0]._tr.get_or_add_trPr()
            self.assertEqual(len(header_properties.findall(qn("w:tblHeader"))), 0)


if __name__ == "__main__":
    unittest.main()
