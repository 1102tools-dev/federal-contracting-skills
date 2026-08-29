#!/usr/bin/env python3
"""Validate a generated GovCon Growth Brief and its numeric evidence."""

from __future__ import annotations

import argparse
import json
import re
import sys
import zipfile
from pathlib import Path

from docx import Document


ROUTE_STRUCTURE = {
    "opportunity": ("Pipeline decision", "48-hour shortlist moves", "Priority opportunities", "Pipeline prioritization", "Qualification gaps", "Capture action plan", "PORTFOLIO POSTURE"),
    "bid_screen": ("Pursuit posture", "48-hour decision gates", "Executive scorecard", "Pursuit logic", "Material unknowns", "Conditions before commitment", "MANAGEMENT POSTURE"),
    "competitor": ("Competitive posture", "Immediate positioning moves", "Positioning snapshot", "Competitive implications", "Claims and assumptions to validate", "Engagement plan", "POSITIONING POSTURE"),
    "recompete": ("Pipeline timing posture", "Near-term validation moves", "Recompete radar", "Timing and validation thesis", "Dates and triggers to validate", "Validation calendar", "PIPELINE POSTURE"),
    "teaming": ("Partner posture", "Next partner moves", "Partner-fit scorecard", "Partner-fit decision", "Diligence gaps", "Diligence and engagement plan", "PARTNER POSTURE"),
    "market": ("Account posture", "Next account moves", "Market thesis", "Account implications", "Account unknowns", "90-day account plan", "ACCOUNT POSTURE"),
    "pricing": ("Pricing posture", "Immediate pricing moves", "Rate-position dashboard", "Rate-position interpretation", "Pricing unknowns", "Proposal guardrails", "PRICING POSTURE"),
    "refresh": ("Updated posture", "Immediate update moves", "What changed", "Decision impact of the delta", "Unresolved deltas", "Updated action plan", "DELTA POSTURE"),
}
DEFAULT_STRUCTURE = ("Executive posture", "Immediate moves", "Decision-relevant analysis", "Commercial implications", "Operational unknowns", "Action plan", "MANAGEMENT POSTURE")
FORBIDDEN = [
    re.compile(r"\bguaranteed\s+(?:award|opportunity|recompete|win)\b", re.I),
    re.compile(r"\bmcp__|/mnt/|/Users/|[A-Za-z]:\\", re.I),
    re.compile(r"\bghp_[A-Za-z0-9]{20,}\b|\b(?:sk|cfat|SAM)-[A-Za-z0-9_-]{16,}\b", re.I),
]
# Internal harness vocabulary must never reach a reader-visible field. Honest
# illustrative labeling uses the standard limitations line from the brief
# specification instead.
BANNED_HARNESS_VOCABULARY = [
    re.compile(r"\bfictional\b", re.I),
    re.compile(r"\bsynthetic\b", re.I),
    re.compile(r"\bfixture\b", re.I),
    re.compile(r"\btest\b", re.I),
    re.compile(r"\bbounded[- ]sample\b", re.I),
    re.compile(r"\barchived record\b", re.I),
]
FEDERAL_SOURCE_CLASSES = {"federal_mcp", "official_web", "other_web"}
# Internal evidence-class vocabulary must never render into a reader-visible
# document; the builder maps these tokens to reader labels.
INTERNAL_CLASS_TOKENS = ("federal_mcp", "official_web", "other_web", "user_statement", "source_class")
MONEY_TERMS = ("value", "obligation", "cost", "price", "funding", "fee", "amount")
# Reader-facing evidence-basis claims that assert no live research happened.
SUPPLIED_ONLY_CLAIMS = (
    "Supplied evidence only",
    "No public research performed",
    "No external query recorded",
)
# Reader-facing evidence-basis claims that assert live research happened.
LIVE_RESEARCH_CLAIM = re.compile(r"\b(?:public-source|live)\b.*\bresearch\b", re.I)
MIDNIGHT_TIMESTAMP = re.compile(r"T00:00(?::00(?:\.0+)?)?(?:Z|[+-]00:00)?$")


def rendered_totals(label: str, total: float) -> set[str]:
    """Return the total renderings the builder may emit for a numeric check."""
    renderings = {f"{total:,.2f}"}
    if any(term in label.lower() for term in MONEY_TERMS):
        renderings.add(f"${total:,.0f}")
    return renderings


def all_text(document: Document) -> str:
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    for section in document.sections:
        for part in (section.header, section.footer):
            parts.extend(p.text for p in part.paragraphs)
    return "\n".join(parts)


def validate(document_path: Path, record_path: Path) -> dict:
    failures: list[str] = []
    if not zipfile.is_zipfile(document_path):
        return {"status": "fail", "failures": ["file is not a valid DOCX ZIP"]}
    document = Document(document_path)
    record = json.loads(record_path.read_text(encoding="utf-8"))
    text = all_text(document)
    headings = [p.text.strip() for p in document.paragraphs if getattr(p.style, "name", "") == "Heading 1"]
    posture_heading, immediate_heading, analysis_heading, assessment_heading, unknowns_heading, action_heading, decision_label = ROUTE_STRUCTURE.get(
        record.get("workflow_mode", ""), DEFAULT_STRUCTURE
    )
    required_headings = [
        posture_heading, immediate_heading, analysis_heading, assessment_heading,
        "Business question and scope", "Company context and assumptions", unknowns_heading,
        "Risks, contrary evidence, and limitations", action_heading,
        "Research record", "Evidence appendix",
    ]
    for heading in required_headings:
        if heading not in headings:
            failures.append(f"missing Heading 1 section: {heading}")
    if [h for h in headings if h in required_headings] != required_headings:
        failures.append("required Heading 1 sections are out of order")
    if decision_label not in text:
        failures.append(f"first-page decision label is missing: {decision_label}")
    for label in (posture_heading, immediate_heading, analysis_heading):
        if label not in text:
            failures.append(f"route-native paid-value content is missing: {label}")
    if record.get("findings") and "Decision signal" not in text:
        failures.append("first-page decision dashboard is missing")
    unknown_items = list(record.get("validation", {}).get("missing_bid_context", [])) + list(record.get("unresolved_questions", []))
    for item in unknown_items:
        if isinstance(item, dict):
            value = item.get("text") or item.get("decision") or item.get("question")
        else:
            value = str(item)
        if value and value not in text:
            failures.append(f"operational unknown is missing from brief: {value}")
    for pattern in FORBIDDEN:
        if pattern.search(text):
            failures.append(f"forbidden content matched: {pattern.pattern}")
    for pattern in BANNED_HARNESS_VOCABULARY:
        match = pattern.search(text)
        if match:
            failures.append(f"reader-visible harness vocabulary is prohibited: {match.group(0)}")
    for token in INTERNAL_CLASS_TOKENS:
        if token in text:
            failures.append(f"internal evidence-class token rendered in the document: {token}")
    for item in record.get("findings", []):
        for evidence_id in item.get("evidence_ids", []):
            if evidence_id not in text:
                failures.append(f"finding evidence ID not present in brief: {evidence_id}")
    if not record.get("validation", {}).get("bid_context_complete"):
        if "conditional pursuit posture" not in text:
            failures.append("incomplete internal context is not labeled as a conditional pursuit posture")
        if re.search(r"\b(?:recommend(?:ation)?|decision)\s*:\s*(?:bid|no[- ]bid)\b", text, re.I):
            failures.append("brief makes a bid decision without complete internal context")
    logged_calls = [item for item in record.get("queries", []) if isinstance(item, dict)]
    basis_lines = [
        paragraph.text
        for paragraph in document.paragraphs
        if paragraph.text.startswith("As of ") and " | " in paragraph.text
    ]
    if not basis_lines:
        failures.append("evidence-basis line is missing (expected an 'As of <date> | <basis>' line)")
    for line in basis_lines:
        supplied_only_claims = [claim for claim in SUPPLIED_ONLY_CLAIMS if claim in line]
        if logged_calls and supplied_only_claims:
            failures.append(
                f"evidence-basis line claims '{supplied_only_claims[0]}'"
                f" but the research record logs {len(logged_calls)} live source call(s)"
            )
        if not logged_calls and LIVE_RESEARCH_CLAIM.search(line):
            failures.append(
                "evidence-basis line claims live research but the research record logs no source calls"
            )
    retrieval_times = [
        str(item.get("retrieved_at") or "").strip()
        for item in logged_calls + [item for item in record.get("evidence", []) if isinstance(item, dict)]
        if str(item.get("retrieved_at") or "").strip()
    ]
    if len(retrieval_times) >= 3 and all(MIDNIGHT_TIMESTAMP.search(value) for value in retrieval_times):
        failures.append(
            f"all {len(retrieval_times)} retrieval timestamps are midnight-exact placeholders;"
            " record actual retrieval times for each source call"
        )
    evidence = [item for item in record.get("evidence", []) if isinstance(item, dict)]
    for item in evidence:
        if item.get("source_class") not in FEDERAL_SOURCE_CLASSES:
            continue
        if not str(item.get("locator") or "").strip():
            failures.append(
                f"federal evidence {item.get('id', '(no id)')} has no checkable locator"
                " (notice ID, PIID, UEI, docket, or URL)"
            )
        if not str(item.get("retrieved_at") or "").strip():
            failures.append(f"federal evidence {item.get('id', '(no id)')} has no retrieval date")
    for index, check in enumerate(record.get("validation", {}).get("numeric_checks", [])):
        expected = sum(float(value) for value in check.get("components", []))
        reported = float(check.get("reported_total", expected))
        label = check.get("label", "numeric check")
        if abs(expected - reported) > 0.005:
            failures.append(f"independent recomputation failed for {label}")
        accepted_totals = rendered_totals(label, expected)
        calculation_lines = [
            paragraph.text
            for paragraph in document.paragraphs
            if label in paragraph.text and any(total in paragraph.text for total in accepted_totals)
        ]
        if not calculation_lines:
            failures.append(
                "recomputed total is missing from brief: " + " or ".join(sorted(accepted_totals))
            )
            continue
        locator = f"validation.numeric_checks[{index}]"
        calculation_ids = [
            item.get("id")
            for item in evidence
            if item.get("source_class") == "calculation" and item.get("locator") == locator
        ]
        if len(calculation_ids) != 1:
            failures.append(
                f"numeric check {label} does not have exactly one calculation evidence item for {locator}"
            )
            continue
        for evidence_id in calculation_ids:
            if not any(f"[{evidence_id}]" in line for line in calculation_lines):
                failures.append(
                    f"numeric check {label} does not cite its calculation evidence ID: {evidence_id}"
                )
    return {"status": "pass" if not failures else "fail", "heading_count": len(headings), "failures": failures}


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("document", type=Path)
    parser.add_argument("--record", required=True, type=Path)
    parser.add_argument("--json", action="store_true")
    args = parser.parse_args()
    try:
        result = validate(args.document, args.record)
    except (OSError, ValueError, json.JSONDecodeError, zipfile.BadZipFile) as exc:
        print(f"ERROR: {exc}", file=sys.stderr)
        return 2
    if args.json:
        print(json.dumps(result, indent=2, sort_keys=True))
    elif result["status"] == "pass":
        print("GovCon Growth DOCX validation passed.")
    else:
        print("VALIDATION FAILED")
        for failure in result["failures"]:
            print(f"- {failure}")
    return 0 if result["status"] == "pass" else 1


if __name__ == "__main__":
    raise SystemExit(main())
