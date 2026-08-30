#!/usr/bin/env python3
"""Build a structured GovCon Growth Brief from a validated research record."""

from __future__ import annotations

import argparse
import json
import subprocess
import sys
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Inches, Pt, RGBColor


NAVY = "17365D"
GREEN = "167D5A"
GRAY = "5B6573"

WORKFLOW_PRODUCTS = {
    "opportunity": {
        "title": "Federal Opportunity Shortlist",
        "posture": "Pipeline decision",
        "decision_label": "PORTFOLIO POSTURE",
        "analysis": "Priority opportunities",
        "assessment": "Pipeline prioritization",
        "immediate": "48-hour shortlist moves",
        "actions": "Capture action plan",
        "unknowns": "Qualification gaps",
    },
    "bid_screen": {
        "title": "Opportunity Evidence Screen",
        "posture": "Pursuit posture",
        "decision_label": "MANAGEMENT POSTURE",
        "analysis": "Executive scorecard",
        "assessment": "Pursuit logic",
        "immediate": "48-hour decision gates",
        "actions": "Conditions before commitment",
        "unknowns": "Material unknowns",
    },
    "competitor": {
        "title": "Competitor Landscape",
        "posture": "Competitive posture",
        "decision_label": "POSITIONING POSTURE",
        "analysis": "Positioning snapshot",
        "assessment": "Competitive implications",
        "immediate": "Immediate positioning moves",
        "actions": "Engagement plan",
        "unknowns": "Claims and assumptions to validate",
    },
    "recompete": {
        "title": "Recompete and Follow-on Pipeline",
        "posture": "Pipeline timing posture",
        "decision_label": "PIPELINE POSTURE",
        "analysis": "Recompete radar",
        "assessment": "Timing and validation thesis",
        "immediate": "Near-term validation moves",
        "actions": "Validation calendar",
        "unknowns": "Dates and triggers to validate",
    },
    "teaming": {
        "title": "Teaming Partner Decision Card",
        "posture": "Partner posture",
        "decision_label": "PARTNER POSTURE",
        "analysis": "Partner-fit scorecard",
        "assessment": "Partner-fit decision",
        "immediate": "Next partner moves",
        "actions": "Diligence and engagement plan",
        "unknowns": "Diligence gaps",
    },
    "market": {
        "title": "Agency or Market Account Plan",
        "posture": "Account posture",
        "decision_label": "ACCOUNT POSTURE",
        "analysis": "Market thesis",
        "assessment": "Account implications",
        "immediate": "Next account moves",
        "actions": "90-day account plan",
        "unknowns": "Account unknowns",
    },
    "pricing": {
        "title": "Labor-Rate and Pricing Context",
        "posture": "Pricing posture",
        "decision_label": "PRICING POSTURE",
        "analysis": "Rate-position dashboard",
        "assessment": "Rate-position interpretation",
        "immediate": "Immediate pricing moves",
        "actions": "Proposal guardrails",
        "unknowns": "Pricing unknowns",
    },
    "refresh": {
        "title": "Prior-Research Delta Audit",
        "posture": "Updated posture",
        "decision_label": "DELTA POSTURE",
        "analysis": "What changed",
        "assessment": "Decision impact of the delta",
        "immediate": "Immediate update moves",
        "actions": "Updated action plan",
        "unknowns": "Unresolved deltas",
    },
}

DEFAULT_PRODUCT = {
    "title": "GovCon Growth Analysis",
    "posture": "Executive posture",
    "decision_label": "MANAGEMENT POSTURE",
    "analysis": "Decision-relevant analysis",
    "assessment": "Commercial implications",
    "immediate": "Immediate moves",
    "actions": "Action plan",
    "unknowns": "Operational unknowns",
}


def shade(cell, fill: str) -> None:
    props = cell._tc.get_or_add_tcPr()
    element = OxmlElement("w:shd")
    element.set(qn("w:fill"), fill)
    props.append(element)


def set_cell_text(cell, value: object, *, bold: bool = False, color: str | None = None) -> None:
    cell.text = ""
    run = cell.paragraphs[0].add_run(str(value if value not in (None, "") else "Not provided"))
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def configure(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)
    normal = document.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08
    for name, size, color in (("Title", 28, NAVY), ("Heading 1", 17, NAVY), ("Heading 2", 12.5, GREEN)):
        style = document.styles[name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(5)
    if "Source Citation" not in document.styles:
        style = document.styles.add_style("Source Citation", WD_STYLE_TYPE.CHARACTER)
        style.font.name = "Aptos"
        style.font.size = Pt(8.5)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(GREEN)
    header = section.header.paragraphs[0]
    header.text = "1102tools  |  GovCon Growth"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.runs[0].font.size = Pt(8)
    header.runs[0].font.color.rgb = RGBColor.from_string(GRAY)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("Prepared as of the date shown  |  Page ").font.size = Pt(8)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    footer._p.append(field)


def add_table(
    document: Document,
    headers: list[str],
    rows: list[list[object]],
    widths: list[float],
    *,
    repeat_header: bool = True,
):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for index, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[index], header, bold=True, color="FFFFFF")
        shade(table.rows[0].cells[index], NAVY)
        table.rows[0].cells[index].width = Inches(widths[index])
    for row_index, row in enumerate(rows):
        cells = table.add_row().cells
        for index, value in enumerate(row):
            set_cell_text(cells[index], value)
            if row_index % 2:
                shade(cells[index], "F4F6F8")
            cells[index].width = Inches(widths[index])
    for row in table.rows:
        row._tr.get_or_add_trPr().append(OxmlElement("w:cantSplit"))
    if repeat_header:
        table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    return table


def add_hyperlink(paragraph, text: str, url: str) -> None:
    relationship = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.extend((color, underline))
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.extend((properties, text_node))
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_bullets(document: Document, items: list[object], empty: str) -> None:
    if not items:
        document.add_paragraph(empty)
        return
    for item in items:
        if isinstance(item, dict):
            text = item.get("text") or item.get("decision") or item.get("question") or json.dumps(item, sort_keys=True)
        else:
            text = str(item)
        paragraph = document.add_paragraph(text, style="List Bullet")
        paragraph.paragraph_format.keep_with_next = False
        paragraph.paragraph_format.keep_together = True
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.paragraph_format.line_spacing = 1.05


def item_text(item: object) -> str:
    """Return the reader-facing text from a record item."""
    if isinstance(item, dict):
        return str(
            item.get("text")
            or item.get("decision")
            or item.get("question")
            or json.dumps(item, sort_keys=True)
        )
    return str(item)


def operational_unknowns(record: dict) -> list[object]:
    """Combine decision-blocking unknowns without converting absence into fact."""
    items: list[object] = []
    items.extend(record.get("validation", {}).get("missing_bid_context", []))
    items.extend(record.get("unresolved_questions", []))
    seen: set[str] = set()
    result: list[object] = []
    for item in items:
        text = item_text(item).strip()
        if text and text not in seen:
            seen.add(text)
            result.append(item)
    return result


def source_map(record: dict) -> tuple[dict[str, str], list[dict]]:
    """Map internal evidence IDs to reader-facing sources ordered by first use."""
    evidence = [item for item in record.get("evidence", []) if isinstance(item, dict)]
    by_id = {str(item.get("id")): item for item in evidence if item.get("id")}
    first_use: list[str] = []
    for collection in (record.get("findings", []), record.get("inferences", [])):
        for item in collection:
            if isinstance(item, dict):
                first_use.extend(str(value) for value in item.get("evidence_ids", []))
    first_use.extend(str(item.get("id", "")) for item in evidence)

    id_to_source: dict[str, str] = {}
    key_to_source: dict[tuple[str, str, str], str] = {}
    entries: list[dict] = []
    for evidence_id in first_use:
        if not evidence_id or evidence_id in id_to_source or evidence_id not in by_id:
            continue
        item = by_id[evidence_id]
        key = (
            str(item.get("source_class", "")),
            str(item.get("locator", "")).strip(),
            str(item.get("title", "")).strip(),
        )
        source_id = key_to_source.get(key)
        if source_id is None:
            source_id = f"S{len(entries) + 1}"
            key_to_source[key] = source_id
            entries.append({**item, "source_id": source_id, "facts": [str(item.get("fact", "")).strip()]})
        else:
            entry = next(value for value in entries if value["source_id"] == source_id)
            fact = str(item.get("fact", "")).strip()
            if fact and fact not in entry["facts"]:
                entry["facts"].append(fact)
        id_to_source[evidence_id] = source_id
    return id_to_source, entries


def source_markers(ids: list[str], mapping: dict[str, str]) -> list[str]:
    return list(dict.fromkeys(mapping[value] for value in ids if value in mapping))


def add_page_one_signals(document: Document, findings: list[dict], mapping: dict[str, str]) -> None:
    """Add a compact decision dashboard from approved findings."""
    rows: list[list[object]] = []
    for finding in findings[:3]:
        markers = source_markers(finding.get("evidence_ids", []), mapping)
        rows.append([
            finding.get("text", "No finding text was recorded."),
            f"[{', '.join(markers)}]" if markers else "Uncited",
        ])
    if rows:
        add_table(document, ["Decision signal", "Evidence"], rows, [5.55, 1.35])


def cite_ids(paragraph, ids: list[str], mapping: dict[str, str]) -> None:
    markers = source_markers(ids, mapping)
    if markers:
        paragraph.add_run(" ")
        paragraph.add_run(f"[{', '.join(markers)}]", style="Source Citation")


def research_basis(record: dict) -> str:
    """Return a reader-facing evidence basis without overstating research performed."""
    web_mode = record.get("web_research", {}).get("mode")
    queries = record.get("queries", [])
    if queries:
        # Logged source calls are live research even when public web research
        # was declined; a supplied-only claim would contradict the record.
        if web_mode == "no_public_web":
            return "Live federal data research with supplied company context | No public web research performed"
        return "Public-source research with supplied company context"
    if web_mode == "no_public_web":
        return "Supplied evidence only | No public research performed"
    return "Supplied company and planning evidence | No external query recorded"


def pretty_label(key: object) -> str:
    """Format metadata labels for a reader instead of exposing field names."""
    label = str(key).replace("_", " ").title()
    replacements = {
        "As Of Date": "As of date",
        "As Of": "As of",
        "Naics": "NAICS",
        "Psc": "PSC",
        "Id": "ID",
    }
    for source, target in replacements.items():
        label = label.replace(source, target)
    return label


# Internal evidence-class vocabulary is plumbing, not capture language. Records
# and validators keep the internal tokens; every reader-visible rendering uses
# these labels, which match the market research product word for word.
SOURCE_CLASS_LABELS = {
    "document": "Supplied document",
    "federal_mcp": "Federal data service",
    "official_web": "Official website",
    "other_web": "Public web source",
    "user_statement": "Customer statement",
    "calculation": "Recorded calculation",
}


def source_class_label(value: object) -> str:
    token = str(value or "")
    return SOURCE_CLASS_LABELS.get(token, pretty_label(token)) if token else ""


def display_value(value: object) -> str:
    """Render structured scope values as concise reader-facing text."""
    if value in (None, "", [], {}):
        return "Not stated"
    if isinstance(value, list):
        return ", ".join(display_value(item) for item in value)
    if isinstance(value, dict):
        return "; ".join(f"{pretty_label(key)}: {display_value(item)}" for key, item in value.items())
    return str(value)


def format_calculated_total(label: str, total: float) -> str:
    """Use currency formatting for money-like checks without changing other totals."""
    money_terms = ("value", "obligation", "cost", "price", "funding", "fee", "amount")
    if any(term in label.lower() for term in money_terms):
        return f"{label}: ${total:,.0f}"
    return f"{label}: {total:,.2f}"


def reader_summary(record: dict, fallback: str) -> str:
    """Replace fixture-language with a concise customer-facing status statement."""
    validation = record.get("validation", {})
    summary = str(validation.get("executive_summary") or "").strip()
    if summary and not any(marker in summary.lower() for marker in ("offline fixture", "test fixture", "demonstrates evidence")):
        return summary
    return validation.get("assessment") or fallback


def scope_rows(scope: dict) -> list[list[object]]:
    """Format scope metadata as executive-readable labels and values."""
    rows: list[list[object]] = []
    for key, value in scope.items():
        label = pretty_label(key)
        if key.endswith("_usd"):
            label = label.removesuffix(" Usd")
            if isinstance(value, (int, float)):
                value = f"${value:,.0f}"
        rows.append([label, display_value(value)])
    return rows


def build(record: dict, output: Path) -> None:
    document = Document()
    configure(document)
    as_of = record.get("scope", {}).get("as_of_date", "Not stated")
    has_bid_decision = bool(record.get("validation", {}).get("bid_context_complete"))
    workflow_mode = record.get("workflow_mode", "")
    product = WORKFLOW_PRODUCTS.get(workflow_mode, DEFAULT_PRODUCT)
    label = record.get("validation", {}).get("report_title", product["title"])
    validation = record.get("validation", {})
    findings = record.get("findings", [])
    unknowns = operational_unknowns(record)
    next_actions = validation.get("next_actions", [])
    id_to_source, source_entries = source_map(record)

    brand = document.add_paragraph("GOVCON GROWTH | DECISION PRODUCT")
    brand.runs[0].bold = True
    brand.runs[0].font.size = Pt(9)
    brand.runs[0].font.color.rgb = RGBColor.from_string(GREEN)
    document.add_paragraph(label, style="Title")
    subtitle = document.add_paragraph(record.get("question", "GovCon growth research"))
    subtitle.runs[0].font.size = Pt(13)
    subtitle.runs[0].font.color.rgb = RGBColor.from_string(GRAY)
    document.add_paragraph(f"As of {as_of} | {research_basis(record)}")
    opening_insight = findings[0].get("text", "No approved finding was recorded.") if findings else "No approved finding was recorded."
    posture = validation.get("assessment") or validation.get("executive_summary") or opening_insight
    callout = document.add_table(rows=1, cols=1)
    callout.style = "Table Grid"
    shade(callout.cell(0, 0), "E8EEF5")
    set_cell_text(callout.cell(0, 0), f"{product['decision_label']}\n{posture}", bold=True, color=NAVY)

    document.add_heading(product["posture"], level=1)
    document.add_paragraph(reader_summary(record, opening_insight))
    add_page_one_signals(document, findings, id_to_source)
    if not has_bid_decision:
        quote = document.add_paragraph(style="Intense Quote")
        quote.add_run("Decision boundary: ").bold = True
        quote.add_run("Internal company context is incomplete. This product is a conditional pursuit posture, not a bid or no-bid recommendation.")

    document.add_heading(product["immediate"], level=1)
    add_bullets(document, next_actions[:3], "No approved immediate move was recorded.")
    document.add_heading(product["unknowns"], level=2)
    add_bullets(document, unknowns[:3], "No decision-blocking unknown was recorded.")
    decision_rule = validation.get("decision_rule")
    if decision_rule:
        quote = document.add_paragraph(style="Intense Quote")
        quote.add_run("Decision rule: ").bold = True
        quote.add_run(str(decision_rule))

    analysis_heading = document.add_heading(product["analysis"], level=1)
    analysis_heading.paragraph_format.page_break_before = True
    for finding in findings:
        p = document.add_paragraph(finding.get("text", ""), style="List Bullet")
        cite_ids(p, finding.get("evidence_ids", []), id_to_source)
    if not findings:
        document.add_paragraph("No approved finding was recorded.")
    for index, check in enumerate(record.get("validation", {}).get("numeric_checks", [])):
        total = sum(float(value) for value in check.get("components", []))
        locator = f"validation.numeric_checks[{index}]"
        calculation_ids = [
            item.get("id")
            for item in record.get("evidence", [])
            if isinstance(item, dict)
            and item.get("source_class") == "calculation"
            and item.get("locator") == locator
        ]
        if len(calculation_ids) != 1:
            raise ValueError(
                f"numeric check {index} requires exactly one calculation evidence item whose locator is {locator}"
            )
        paragraph = document.add_paragraph(format_calculated_total(check.get("label", "Calculated total"), total))
        cite_ids(paragraph, calculation_ids, id_to_source)

    document.add_heading(product["assessment"], level=1)
    document.add_paragraph(validation.get("assessment", "No final assessment was approved."))
    pipeline = validation.get("pipeline", [])
    if pipeline:
        add_table(document, ["Candidate", "Signal", "Timing", "Confidence", "Next validation"], [[p.get("candidate", ""), p.get("signal", ""), p.get("timing", ""), p.get("confidence", ""), p.get("next_validation", "")] for p in pipeline], [1.5, 1.8, 1.1, 0.8, 1.7])

    context = record.get("user_context", [])
    assumptions = record.get("assumptions", [])
    scope = record.get("scope", {})
    context_scope = {key: value for key, value in scope.items() if key != "as_of_date" and value not in (None, "", [], {})}
    if context_scope or context or assumptions:
        document.add_heading("Key context", level=1)
        if context_scope:
            add_table(document, ["Context", "Working value"], scope_rows(context_scope), [2.1, 4.8])
        if context:
            add_bullets(document, context, "")
        if assumptions:
            add_bullets(document, assumptions, "")

    if unknowns[3:]:
        document.add_heading("Additional open items", level=1)
        add_bullets(document, unknowns[3:], "")

    conflicts = record.get("conflicts", [])
    inferences = record.get("inferences", [])
    if conflicts or inferences:
        document.add_heading("Risks and evidence limits", level=1)
        add_bullets(document, conflicts, "")
    for inference in inferences:
        p = document.add_paragraph("Inference: " + inference.get("text", inference.get("reasoning", "")), style="List Bullet")
        cite_ids(p, inference.get("evidence_ids", []), id_to_source)

    user_decisions = record.get("user_decisions", [])
    if user_decisions or next_actions[3:]:
        document.add_heading(product["actions"], level=1)
        add_bullets(document, user_decisions, "")
        document.add_paragraph("The immediate moves on page one lead this plan.")
        add_bullets(document, next_actions[3:], "")

    document.add_heading("Source Register", level=1)
    for entry in source_entries:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(3)
        paragraph.add_run(f"[{entry.get('source_id', '')}] ").bold = True
        paragraph.add_run(f"{source_class_label(entry.get('source_class'))}. {entry.get('title', '')}. ")
        locator = str(entry.get("locator", ""))
        url = str(entry.get("canonical_url") or (locator if locator.startswith("http") else ""))
        if url:
            add_hyperlink(paragraph, url, url)
        elif locator:
            paragraph.add_run(locator)
        date = entry.get("as_of_date") or entry.get("retrieved_at")
        if date:
            paragraph.add_run(f". {date}")

    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("record", type=Path)
    parser.add_argument("output", type=Path)
    args = parser.parse_args()
    validator = Path(__file__).with_name("validate_research_record.py")
    result = subprocess.run([sys.executable, str(validator), str(args.record)], check=False)
    if result.returncode:
        return result.returncode
    build(json.loads(args.record.read_text(encoding="utf-8")), args.output)
    print(args.output)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
