#!/usr/bin/env python3
"""Validate a generated Market Research DOCX and its numeric evidence."""

from __future__ import annotations

import argparse
import json
import re
import sys
import zipfile
from pathlib import Path

from docx import Document

ROUTE_OUTCOMES = {
    "complete_report": [
        "Acquisition and decision frame",
        "What the evidence establishes",
        "Market capability and packaging",
        "Market engagement instrument",
        "Evidence-to-decision gates",
    ],
    "refresh": [
        "Refresh assessment",
        "Prior-to-current evidence comparison",
        "Vendor and market-structure changes",
        "Strategy changes to make now",
        "What remains usable and what must be rechecked",
        "Refresh action plan",
    ],
    "one_question": [
        "Bounded answer",
        "Evidence for and against",
        "Decision implications",
        "Further research options",
    ],
    "pre_award_handoff": [
        "Handoff summary",
        "Approved market observations",
        "Market findings translated into acquisition inputs",
        "Pricing inputs and boundaries",
        "Pre-Award risk register",
        "Pre-Award intake and next actions",
    ],
}
FORBIDDEN = [
    re.compile(r"\b(?:automatically|therefore)\s+(?:recommend|requires?|proves?)\b", re.I),
    re.compile(r"\b(?:set[- ]aside|commerciality|contract type|price reasonableness)\s+is\s+automatically\b", re.I),
    re.compile(r"\bmcp__|/mnt/|/Users/|[A-Za-z]:\\", re.I),
    re.compile(r"\bghp_[A-Za-z0-9]{20,}\b|\b(?:sk|cfat|SAM)-[A-Za-z0-9_-]{16,}\b", re.I),
]
# Internal evidence-class vocabulary must never render into a reader-visible
# document; the builder maps these tokens to reader labels.
INTERNAL_CLASS_TOKENS = ("federal_mcp", "official_web", "other_web", "user_statement", "source_class")
MIDNIGHT_TIMESTAMP = re.compile(r"T00:00(?::00(?:\.0+)?)?(?:Z|[+-]00:00)?$")
# A retrieval stamp is a per-call fact. Three or more midnight-exact stamps, or
# five or more byte-identical stamps, are a synthesized batch value rather than
# the recorded time of each source call.
MIDNIGHT_PLACEHOLDER_LIMIT = 3
IDENTICAL_PLACEHOLDER_LIMIT = 5


def evidence_id_map(record: dict) -> dict[str, str]:
    """Mirror the builder's internal-evidence to reader-source mapping."""
    ordered: list[str] = []

    def visit(value: object) -> None:
        if isinstance(value, dict):
            for key, item in value.items():
                if key == "evidence_ids" and isinstance(item, list):
                    ordered.extend(str(candidate) for candidate in item)
                else:
                    visit(item)
        elif isinstance(value, list):
            for item in value:
                visit(item)

    visit(record.get("validation", {}))
    route = record.get("workflow_mode")
    if route in {"complete_report", "one_question", "pre_award_handoff"} or not record.get("validation", {}).get("decision_implications"):
        visit(record.get("findings", []))
    if route == "complete_report":
        ordered.extend(str(item.get("id", "")) for item in record.get("evidence", []) if isinstance(item, dict))
    by_id = {
        str(item.get("id")): item
        for item in record.get("evidence", [])
        if isinstance(item, dict) and item.get("id")
    }
    mapping: dict[str, str] = {}
    keys: dict[tuple[str, str, str], str] = {}
    for evidence_id in ordered:
        if evidence_id in mapping or evidence_id not in by_id:
            continue
        item = by_id[evidence_id]
        key = (str(item.get("source_class", "")), str(item.get("locator", "")).strip(), str(item.get("title", "")).strip())
        source_id = keys.setdefault(key, f"S{len(keys) + 1}")
        mapping[evidence_id] = source_id
    return mapping


def retrieval_timestamps(record: dict) -> list[str]:
    """Every recorded retrieval stamp across the source-call and evidence logs."""
    items = [item for item in record.get("queries", []) if isinstance(item, dict)]
    items += [item for item in record.get("evidence", []) if isinstance(item, dict)]
    return [
        str(item.get("retrieved_at") or "").strip()
        for item in items
        if str(item.get("retrieved_at") or "").strip()
    ]


def placeholder_timestamp_failures(record: dict) -> list[str]:
    failures: list[str] = []
    stamps = retrieval_timestamps(record)
    if len(stamps) >= MIDNIGHT_PLACEHOLDER_LIMIT and all(
        MIDNIGHT_TIMESTAMP.search(value) for value in stamps
    ):
        failures.append(
            f"all {len(stamps)} retrieval timestamps are midnight-exact placeholders;"
            " record the actual retrieval time of each source call"
        )
    if len(stamps) >= IDENTICAL_PLACEHOLDER_LIMIT and len(set(stamps)) == 1:
        failures.append(
            f"all {len(stamps)} retrieval timestamps are identical ({stamps[0]});"
            " record the actual retrieval time of each source call"
        )
    return failures


def all_text(document: Document) -> str:
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


def _looks_synthetic(item: dict) -> bool:
    text = " ".join(
        str(item.get(field, ""))
        for field in ("title", "locator", "fact", "limitations")
    ).lower()
    return any(marker in text for marker in ("fixture", "synthetic", "no live", "test data only"))


def evidence_supports_complete_label(record: dict) -> bool:
    evidence = [item for item in record.get("evidence", []) if isinstance(item, dict)]
    live_federal = any(
        item.get("source_class") == "federal_mcp" and not _looks_synthetic(item)
        for item in evidence
    )
    public_web = any(item.get("source_class") in {"official_web", "other_web"} for item in evidence)
    commercial_approved = record.get("validation", {}).get("commercial_evidence_complete") is True
    return live_federal and public_web and commercial_approved


def validate(document_path: Path, record_path: Path) -> dict:
    failures: list[str] = []
    if not zipfile.is_zipfile(document_path):
        return {"status": "fail", "failures": ["file is not a valid DOCX ZIP"]}
    document = Document(document_path)
    record = json.loads(record_path.read_text(encoding="utf-8"))
    text = all_text(document)
    cited_source_ids = {
        source_id
        for group in re.findall(r"\[([^\[\]]+)\]", text)
        for source_id in re.findall(r"\bS\d+\b", group)
    }
    headings = [p.text.strip() for p in document.paragraphs if getattr(p.style, "name", "") == "Heading 1"]
    route = record.get("workflow_mode")
    required_outcomes = list(ROUTE_OUTCOMES.get(route, []))
    if route == "one_question" and record.get("validation", {}).get("analysis_focus") == "small_business":
        required_outcomes = [
            "Bounded answer",
            "Candidate small-business market",
            "Rule of Two evidence assessment",
            "Evidence supporting and cutting against a small-business strategy",
            "Targeted outreach plan",
            "Decision implications",
            "Further research options",
        ]
    if not required_outcomes:
        failures.append(f"unsupported workflow_mode for report validation: {route}")
    for outcome in required_outcomes:
        if outcome not in text:
            failures.append(f"route-native outcome is missing: {outcome}")
    for label in ("BOTTOM LINE", "Decision implications", "Next practical actions"):
        if label not in text:
            failures.append(f"first-page decision product element is missing: {label}")
    for pattern in FORBIDDEN:
        if pattern.search(text):
            failures.append(f"forbidden content matched: {pattern.pattern}")
    for token in INTERNAL_CLASS_TOKENS:
        if token in text:
            failures.append(f"internal evidence-class token rendered in the document: {token}")
    if re.search(r"\bE\d{3,}\b", text):
        failures.append("internal E-style evidence identifiers are reader-visible")
    if "Source Register" not in headings:
        failures.append("reader-facing Source Register is missing")
    id_map = evidence_id_map(record)
    findings_to_check = record.get("findings", []) if route == "complete_report" else []
    for item in findings_to_check:
        for evidence_id in item.get("evidence_ids", []):
            source_id = id_map.get(evidence_id)
            if not source_id or source_id not in cited_source_ids:
                failures.append(f"finding source marker not present in report for internal evidence: {evidence_id}")
    complete = evidence_supports_complete_label(record)
    if route == "complete_report" and not complete and "Federal-Data Desk-Research Draft" not in text:
        failures.append("incomplete commercial evidence is not labeled as a desk-research draft")
    if route == "complete_report" and not complete and "FAR Part 10 Market Research Report" in text:
        failures.append("incomplete evidence is mislabeled as a FAR Part 10 Market Research Report")
    failures.extend(placeholder_timestamp_failures(record))
    evidence = [item for item in record.get("evidence", []) if isinstance(item, dict)]
    # Every reader-facing [S#] marker must resolve to the Source Register.
    rendered_register_ids = set(id_map.values())
    cited_in_document = {
        cited
        for group in re.findall(r"\[([^\[\]]+)\]", text)
        for cited in re.findall(r"\bS\d+\b", group)
    }
    dangling = sorted(cited_in_document - rendered_register_ids)
    if dangling:
        failures.append(
            "dangling source citations in the document: "
            + ", ".join(dangling)
            + " do not appear in the Source Register"
        )
    numeric_checks = record.get("validation", {}).get("numeric_checks", []) if route == "complete_report" else []
    for index, check in enumerate(numeric_checks):
        expected = sum(float(value) for value in check.get("components", []))
        reported = float(check.get("reported_total", expected))
        label = check.get("label", "numeric check")
        if abs(expected - reported) > 0.005:
            failures.append(f"independent recomputation failed for {label}")
        calculated_total = f"{expected:,.2f}"
        calculation_lines = [
            paragraph.text
            for paragraph in document.paragraphs
            if label in paragraph.text and calculated_total in paragraph.text
        ]
        if not calculation_lines:
            failures.append(f"recomputed total is missing from report: {expected:,.2f}")
            continue
        locator = f"validation.numeric_checks[{index}]"
        calculation_ids = [
            item.get("id")
            for item in evidence
            if item.get("source_class") == "calculation" and item.get("locator") == locator
        ]
        if len(calculation_ids) != 1:
            failures.append(
                f"numeric check {label} does not have exactly one calculation evidence item for {locator}"
            )
            continue
        for evidence_id in calculation_ids:
            source_id = id_map.get(evidence_id)
            if not source_id or not any(f"[{source_id}]" in line for line in calculation_lines):
                failures.append(
                    f"numeric check {label} does not cite its source marker for internal evidence: {evidence_id}"
                )
    return {"status": "pass" if not failures else "fail", "heading_count": len(headings), "failures": failures}


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("document", type=Path)
    parser.add_argument("--record", required=True, type=Path)
    parser.add_argument("--json", action="store_true")
    args = parser.parse_args()
    try:
        result = validate(args.document, args.record)
    except (OSError, ValueError, json.JSONDecodeError, zipfile.BadZipFile) as exc:
        print(f"ERROR: {exc}", file=sys.stderr)
        return 2
    if args.json:
        print(json.dumps(result, indent=2, sort_keys=True))
    elif result["status"] == "pass":
        print("Market research DOCX validation passed.")
    else:
        print("VALIDATION FAILED")
        for failure in result["failures"]:
            print(f"- {failure}")
    return 0 if result["status"] == "pass" else 1


if __name__ == "__main__":
    raise SystemExit(main())
