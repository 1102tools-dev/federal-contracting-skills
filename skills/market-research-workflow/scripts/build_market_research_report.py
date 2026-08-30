#!/usr/bin/env python3
"""Build a structured Market Research DOCX from a validated research record."""

from __future__ import annotations

import argparse
import json
import re
import subprocess
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Inches, Pt, RGBColor


NAVY = "17365D"
GREEN = "167D5A"
PALE_GREEN = "EAF4EF"
GRAY = "5B6573"


ROUTE_TITLES = {
    "complete_report": "FAR Part 10 Market Research Report",
    "refresh": "Market Research Refresh",
    "one_question": "Market Evidence Analysis",
    "pre_award_handoff": "Pre-Award Market Research Handoff",
}

# Reader-facing labels for internal evidence-class tokens. The record and the
# validators keep the internal contract vocabulary; only the rendered DOCX uses
# these labels.
SOURCE_CLASS_LABELS = {
    "document": "Supplied document",
    "federal_mcp": "Federal data service",
    "official_web": "Official website",
    "other_web": "Public web source",
    "user_statement": "Customer statement",
    "calculation": "Recorded calculation",
}


def source_class_label(value: object) -> str:
    token = str(value or "")
    return SOURCE_CLASS_LABELS.get(token, pretty_label(token)) if token else ""


def _looks_synthetic(item: dict) -> bool:
    text = " ".join(
        str(item.get(field, ""))
        for field in ("title", "locator", "fact", "limitations")
    ).lower()
    return any(marker in text for marker in ("fixture", "synthetic", "no live", "test data only"))


def completion_state(record: dict) -> tuple[bool, list[str]]:
    """Return an honest full-report label based on evidence, not one flag alone."""
    validation = record.get("validation", {})
    evidence = [item for item in record.get("evidence", []) if isinstance(item, dict)]
    live_federal = any(
        item.get("source_class") == "federal_mcp" and not _looks_synthetic(item)
        for item in evidence
    )
    web_evidence = any(
        item.get("source_class") in {"official_web", "other_web"}
        for item in evidence
    )
    commercial_approved = validation.get("commercial_evidence_complete") is True
    missing = []
    if not live_federal:
        missing.append("live federal award/entity evidence")
    if not web_evidence:
        missing.append("approved public-web evidence")
    if not commercial_approved:
        missing.append("approved commercial-market evidence")
    return not missing, missing


def item_text(item: object) -> str:
    if not isinstance(item, dict):
        return str(item)
    return str(
        item.get("text")
        or item.get("action")
        or item.get("decision")
        or item.get("question")
        or item.get("title")
        or json.dumps(item, sort_keys=True)
    )


def pretty_label(key: object) -> str:
    """Format metadata labels for a reader instead of exposing field names."""
    label = str(key).replace("_", " ").title()
    replacements = {
        "As Of Date": "As of date",
        "As Of": "As of",
        "Naics": "NAICS",
        "Psc": "PSC",
        "Id": "ID",
    }
    for source, target in replacements.items():
        label = label.replace(source, target)
    return label


def display_value(value: object) -> str:
    """Render structured scope values as concise reader-facing text."""
    if value in (None, "", [], {}):
        return "Not stated"
    if isinstance(value, list):
        return ", ".join(display_value(item) for item in value)
    if isinstance(value, dict):
        return "; ".join(f"{pretty_label(key)}: {display_value(item)}" for key, item in value.items())
    return str(value)


def format_calculated_total(label: str, total: float) -> str:
    """Use currency formatting for money-like checks without changing other totals."""
    money_terms = ("value", "obligation", "cost", "price", "funding", "fee", "amount")
    if any(term in label.lower() for term in money_terms):
        return f"{label}: ${total:,.2f}"
    return f"{label}: {total:,.2f}"


def complete_report_fallbacks() -> dict[str, list[dict[str, str]]]:
    """Give an incomplete record a useful, explicitly prospective research plan."""
    return {
        "capability_model": [
            {
                "title": "Comparable service delivery",
                "evidence_to_request": "Comparable help-desk scale, transition approach, coverage model, service-level results, accessibility, and security controls.",
                "failure_signal": "No comparable evidence or no accountable transition plan.",
            }
        ],
        "market_engagement_instrument": [
            {
                "theme": "Comparable delivery",
                "prompt": "Request comparable service examples, transition approach, coverage model, service metrics, accessibility, and security practices.",
                "decision_use": "Assess capability and execution risk",
            }
        ],
        "decision_gates": [
            {
                "gate": "Capability evidence",
                "owner": "Market research lead",
                "exit_condition": "At least two comparable examples and a documented coverage/transition model are recorded.",
                "evidence": "Approved responses and supporting records.",
            }
        ],
        "next_actions": [
            {
                "when": "Before acquisition strategy",
                "owner": "Market research lead + CO",
                "action": "Approve the evidence request, collect comparable responses, and refresh federal data for the current scope.",
                "output": "Evidence register and decision-ready market findings.",
            }
        ],
    }


def reader_summary(record: dict, complete: bool) -> str:
    """Replace fixture-language with a concise customer-facing status statement."""
    validation = record.get("validation", {})
    summary = str(validation.get("executive_summary") or "").strip()
    if summary and not any(marker in summary.lower() for marker in ("offline fixture", "test fixture", "demonstrates evidence")):
        return summary
    question = str(record.get("question", "the stated requirement")).rstrip(" ?")
    for prefix in ("What evidence informs acquisition of ", "What evidence informs the acquisition of "):
        if question.lower().startswith(prefix.lower()):
            question = question[len(prefix) :]
            break
    if record.get("workflow_mode") == "complete_report" and not complete:
        return (
            f"The available record supports a preliminary market frame for {question}, "
            "but live federal, public-web, and commercial evidence are still needed before "
            "a complete report or acquisition-strategy decision."
        )
    return validation.get("assessment") or validation.get("executive_summary") or "No approved executive summary was supplied."


def structured_rows(items: list[object], fields: list[tuple[str, str]], default: str = "Not recorded") -> list[list[str]]:
    rows = []
    for item in items:
        if isinstance(item, dict):
            rows.append([str(item.get(key) or fallback) for key, fallback in fields])
        else:
            rows.append([str(item)] + [fallback for _, fallback in fields[1:]])
    return rows or [[default] + [fallback for _, fallback in fields[1:]]]


def route_payload(record: dict, field: str) -> list[dict]:
    value = record.get("validation", {}).get(field, [])
    return [item for item in value if isinstance(item, dict)] if isinstance(value, list) else []


def require_route_value(record: dict, field: str, purpose: str) -> list[dict]:
    rows = route_payload(record, field)
    if not rows:
        raise ValueError(
            f"{record.get('workflow_mode')} cannot produce a paid-value DOCX without {purpose}; "
            "return a concise evidence-acquisition note instead"
        )
    return rows


ACTION_NAME_LIMIT = 72
ACTION_NAME_ELLIPSIS = "..."
# Words that read as debris when a truncation lands on them, so a word-boundary
# cut drops them before the ellipsis is appended.
DANGLING_WORDS = {
    "a", "an", "and", "against", "as", "at", "by", "for", "from", "in", "into",
    "its", "no", "of", "on", "or", "over", "per", "that", "the", "their",
    "these", "this", "those", "to", "under", "using", "via", "with", "within",
    "each", "every", "both", "such", "so",
}
DETERMINERS = {"a", "an", "the", "its", "their", "this", "these", "those", "each", "every"}
# Citation lead-ins: a label left stranded at the end of a truncation would read
# as a broken reference, so it is dropped with the citation it introduces.
CITATION_LEAD_INS = {
    "far", "dfars", "cfr", "c.f.r.", "u.s.c.", "usc", "part", "subpart",
    "section", "clause", "naics", "psc", "sin", "fasa", "pub.", "no.",
}
# Abbreviations whose trailing period never ends a sentence.
ABBREVIATIONS = {
    "u.s", "e.g", "i.e", "etc", "no", "inc", "corp", "co", "llc", "ltd", "jr",
    "sr", "dr", "mr", "mrs", "ms", "st", "vs", "fig", "approx", "est", "dept",
    "govt", "cf", "al",
}


def _is_sentence_end(text: str, index: int) -> bool:
    """Report whether the period at ``index`` genuinely ends a sentence.

    A period inside a citation (FAR 19.502-2, 48 CFR 10.001, 13 C.F.R.
    121.402), a decimal number, or an abbreviation (U.S., e.g., No., Inc.) is
    not a sentence boundary.
    """
    after = text[index + 1 :]
    # A decimal or a citation continues with a digit or a word character.
    if after[:1] and not after[:1].isspace():
        return False
    # A real sentence break is followed by end-of-string or a capitalized word.
    remainder = after.lstrip()
    if remainder and not remainder[0].isupper():
        return False
    preceding = re.search(r"([A-Za-z.]+)$", text[:index])
    if preceding:
        token = preceding.group(1).lower()
        # A single letter or a dotted initialism before the period is an
        # abbreviation (U.S., e.g., C.F.R.), not a sentence end.
        if len(token) == 1 or "." in token or token in ABBREVIATIONS:
            return False
    return True


def first_sentence(text: str) -> str:
    """Return the first genuine sentence of ``text`` without its terminator."""
    for index, char in enumerate(text):
        if char in ";:":
            return text[:index].strip()
        if char == "." and _is_sentence_end(text, index):
            return text[:index].strip()
    return text.strip()


def _trim_dangling(words: list[str]) -> list[str]:
    while words:
        bare = words[-1].strip(",;:.()").lower()
        if bare in DANGLING_WORDS or bare in CITATION_LEAD_INS or not bare:
            words.pop()
            continue
        break
    return words


def shorten_action(text: str, limit: int = ACTION_NAME_LIMIT) -> str:
    """Return an intelligible short identifier for an action.

    The whole first sentence is used when one genuinely ends and fits. Anything
    longer is cut at a clause or word boundary, never inside a citation, a
    decimal, an abbreviation, or a word, and always carries a single trailing
    ellipsis.
    """
    text = " ".join(str(text).split())
    if not text:
        return text
    sentence = first_sentence(text)
    if sentence and len(sentence) <= limit:
        return sentence
    source = sentence or text
    budget = limit - len(ACTION_NAME_ELLIPSIS)
    # Prefer a comma boundary when it keeps a substantial clause.
    clause = ""
    for match in re.finditer(r",", source):
        candidate = source[: match.start()].rstrip()
        if len(candidate) <= budget and len(candidate) >= budget // 2:
            clause = candidate
    if clause:
        return _finish(clause)
    words = source[: budget + 1].split(" ")
    if len(words) > 1 and len(source) > budget:
        words = words[:-1]
    words = _trim_dangling(words)
    return _finish(_drop_trailing_conjunct(words, budget))


def _drop_trailing_conjunct(words: list[str], budget: int) -> str:
    """Drop a stranded "and the <noun>" tail when enough text remains.

    "…the capability hypotheses and the packaging" keeps only the first
    conjunct, while "…currency and exclusion status" and "Select the vehicle
    and competition strategy" are left whole: their second conjunct still
    carries meaning, or dropping it would leave too little to identify the row.
    """
    for index in range(len(words) - 1, 0, -1):
        if words[index].strip(",;:").lower() not in {"and", "or", "plus"}:
            continue
        tail = words[index + 1 :]
        if len(tail) > 3 or tail[0].strip(",;:").lower() not in DETERMINERS:
            break
        prefix = " ".join(words[:index]).rstrip(" ,;:")
        if len(prefix) >= budget // 2:
            return prefix
        break
    return " ".join(words)


def _finish(fragment: str) -> str:
    fragment = fragment.rstrip(" ,;:.-")
    return f"{fragment}{ACTION_NAME_ELLIPSIS}" if fragment else fragment


def action_name(item: object) -> str:
    """Short reader label for an action row so a closing table can point back to
    the lead Next practical actions table instead of repeating its Action and
    Output text verbatim."""
    if isinstance(item, dict):
        text = str(item.get("action") or "No approved action was recorded.")
    else:
        text = str(item)
    return shorten_action(text) or text


def with_evidence(item: dict, field: str, id_map: dict[str, str] | None = None) -> str:
    text = str(item.get(field, ""))
    ids = [map_evidence_id(value, id_map) for value in item.get("evidence_ids", [])]
    if ids:
        text += " [" + ", ".join(ids) + "]"
    return text


def map_evidence_id(value: object, id_map: dict[str, str] | None) -> str:
    token = str(value)
    return (id_map or {}).get(token, token)


def evidence_id_map(record: dict) -> dict[str, str]:
    """Map internal evidence rows to reader-facing source IDs by first use.

    Multiple facts from the same source share one S number. Internal E-style
    identifiers remain available to validators and sidecars only.
    """
    ordered: list[str] = []

    def visit(value: object) -> None:
        if isinstance(value, dict):
            for key, item in value.items():
                if key == "evidence_ids" and isinstance(item, list):
                    ordered.extend(str(candidate) for candidate in item)
                else:
                    visit(item)
        elif isinstance(value, list):
            for item in value:
                visit(item)

    visit(record.get("validation", {}))
    route = record.get("workflow_mode")
    if route in {"complete_report", "one_question", "pre_award_handoff"} or not record.get("validation", {}).get("decision_implications"):
        visit(record.get("findings", []))
    if record.get("workflow_mode") == "complete_report":
        ordered.extend(str(item.get("id", "")) for item in record.get("evidence", []) if isinstance(item, dict))

    by_id = {
        str(item.get("id")): item
        for item in record.get("evidence", [])
        if isinstance(item, dict) and item.get("id")
    }
    mapping: dict[str, str] = {}
    key_to_source: dict[tuple[str, str, str], str] = {}
    for evidence_id in ordered:
        if evidence_id in mapping or evidence_id not in by_id:
            continue
        item = by_id[evidence_id]
        key = (
            str(item.get("source_class", "")),
            str(item.get("locator", "")).strip(),
            str(item.get("title", "")).strip(),
        )
        source_id = key_to_source.get(key)
        if source_id is None:
            source_id = f"S{len(key_to_source) + 1}"
            key_to_source[key] = source_id
        mapping[evidence_id] = source_id
    return mapping


def source_register_rows(record: dict, id_map: dict[str, str]) -> list[list[str]]:
    rows: dict[str, list[str]] = {}
    for item in record.get("evidence", []):
        if not isinstance(item, dict) or item.get("id") not in id_map:
            continue
        source_id = id_map[str(item["id"])]
        row = rows.setdefault(
            source_id,
            [
                source_id,
                source_class_label(item.get("source_class")),
                f"{item.get('title', '')}\n{item.get('locator', '')}",
                "",
            ],
        )
        fact = str(item.get("fact", "")).strip()
        if fact and fact not in row[3].split(" | "):
            row[3] = " | ".join(value for value in (row[3], fact) if value)
    return [rows[key] for key in sorted(rows, key=lambda value: int(value[1:]))]


def enforce_route_content(record: dict) -> None:
    route = record.get("workflow_mode")
    validation = record.get("validation", {})
    if route == "refresh":
        require_route_value(record, "refresh_comparison", "a dated prior/current comparison")
    elif route == "one_question" and validation.get("analysis_focus") == "small_business":
        require_route_value(record, "small_business_candidates", "a named candidate or documented search-result population")
        if not str(validation.get("rule_of_two_assessment", "")).strip():
            raise ValueError("small-business analysis requires a bounded Rule of Two assessment")
    elif route == "pre_award_handoff":
        for field, purpose in (
            ("scope_implications", "scope implications"),
            ("packaging_implications", "packaging implications"),
            ("performance_implications", "performance implications"),
            ("pricing_inputs", "pricing inputs and boundaries"),
            ("handoff_risks", "a risk register"),
        ):
            require_route_value(record, field, purpose)


def collect_evidence_ids(value: object) -> set[str]:
    ids: set[str] = set()
    if isinstance(value, dict):
        for key, item in value.items():
            if key == "evidence_ids" and isinstance(item, list):
                ids.update(str(candidate) for candidate in item)
            else:
                ids.update(collect_evidence_ids(item))
    elif isinstance(value, list):
        for item in value:
            ids.update(collect_evidence_ids(item))
    return ids


def shade(cell, fill: str) -> None:
    props = cell._tc.get_or_add_tcPr()
    element = OxmlElement("w:shd")
    element.set(qn("w:fill"), fill)
    props.append(element)


def set_cell_text(cell, value: object, bold: bool = False, color: str | None = None) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(str(value if value not in (None, "") else "Not provided"))
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def register_font_fallbacks(document: Document) -> None:
    """Declare sans-serif fallbacks so every route keeps the same modern look.

    Renderers without Aptos (for example LibreOffice) otherwise substitute a
    default serif face for the body while headings resolve to a sans face,
    which splits the product family into two visual designs."""
    for part in document.part.package.iter_parts():
        if str(part.partname) != "/word/fontTable.xml":
            continue
        blob = part.blob
        additions = b""
        for name in (b"Aptos", b"Aptos Display"):
            if b'w:name="' + name + b'"' in blob:
                continue
            additions += (
                b'<w:font w:name="' + name + b'">'
                b'<w:altName w:val="Calibri"/>'
                b'<w:family w:val="swiss"/>'
                b'<w:pitch w:val="variable"/>'
                b"</w:font>"
            )
        if additions:
            part._blob = blob.replace(b"</w:fonts>", additions + b"</w:fonts>")
        break


def configure(document: Document) -> None:
    register_font_fallbacks(document)
    section = document.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    normal = document.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08
    for name, size, color in (("Title", 28, NAVY), ("Heading 1", 17, NAVY), ("Heading 2", 12.5, GREEN)):
        style = document.styles[name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(5)

    if "Source Citation" not in document.styles:
        style = document.styles.add_style("Source Citation", WD_STYLE_TYPE.CHARACTER)
        style.font.name = "Aptos"
        style.font.size = Pt(8.5)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(GREEN)

    header = section.header.paragraphs[0]
    header.text = "1102tools  |  Market Research"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.runs[0].font.size = Pt(8)
    header.runs[0].font.color.rgb = RGBColor.from_string(GRAY)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Prepared as of the date shown  |  Page ")
    run.font.size = Pt(8)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    footer._p.append(field)


def add_table(document: Document, headers: list[str], rows: list[list[object]], widths: list[float] | None = None):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for index, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[index], header, bold=True, color="FFFFFF")
        shade(table.rows[0].cells[index], NAVY)
        if widths:
            table.rows[0].cells[index].width = Inches(widths[index])
    for row_index, row in enumerate(rows):
        cells = table.add_row().cells
        for index, value in enumerate(row):
            set_cell_text(cells[index], value)
            if row_index % 2:
                shade(cells[index], "F4F6F8")
            if widths:
                cells[index].width = Inches(widths[index])
    for row in table.rows:
        row._tr.get_or_add_trPr().append(OxmlElement("w:cantSplit"))
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    return table


def add_hyperlink(paragraph, text: str, url: str) -> None:
    relationship = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.extend((color, underline))
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.extend((properties, text_node))
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_bullets(document: Document, items: list[object], empty: str = "None recorded") -> None:
    if not items:
        document.add_paragraph(empty)
        return
    for item in items:
        if isinstance(item, dict):
            text = item.get("text") or item.get("decision") or item.get("question") or item.get("action") or json.dumps(item, sort_keys=True)
        else:
            text = str(item)
        paragraph = document.add_paragraph(text, style="List Bullet")
        paragraph.paragraph_format.keep_with_next = False
        paragraph.paragraph_format.keep_together = True
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.paragraph_format.line_spacing = 1.05


def cite_ids(paragraph, ids: list[str], id_map: dict[str, str] | None = None) -> None:
    if not ids:
        return
    paragraph.add_run(" [")
    markers = list(dict.fromkeys(map_evidence_id(value, id_map) for value in ids))
    run = paragraph.add_run(", ".join(markers), style="Source Citation")
    run.bold = True
    paragraph.add_run("]")


def build(record: dict, output: Path) -> None:
    validation = record.get("validation", {})
    if record.get("schema_version") != "1.2":
        raise ValueError("market research records must be migrated to schema 1.2 before report generation")
    for field in ("findings_approved", "decisions_approved", "unresolved_items_disposition_approved"):
        if validation.get(field) is not True:
            raise ValueError(f"{field} must be true before report generation")
    enforce_route_content(record)
    document = Document()
    configure(document)
    route = record.get("workflow_mode")
    complete, missing_classes = completion_state(record)
    fallbacks = complete_report_fallbacks() if route == "complete_report" else {}
    if route == "complete_report" and not complete:
        title = "Federal-Data Desk-Research Draft"
    else:
        title = validation.get("report_title") or ROUTE_TITLES.get(route, "Market Research Analysis")
    subtitle = record.get("question", "Market research")
    as_of = record.get("scope", {}).get("as_of_date", "Not stated")
    document.core_properties.title = title
    document.core_properties.subject = subtitle
    document.core_properties.author = "1102tools"

    paragraph = document.add_paragraph("MARKET RESEARCH | DECISION PRODUCT")
    run = paragraph.runs[0]
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string(GREEN)
    title_p = document.add_paragraph(title, style="Title")
    sub = document.add_paragraph(subtitle)
    sub.runs[0].font.size = Pt(13)
    document.add_paragraph(f"As of {as_of} | {title}")
    lead = document.add_table(rows=1, cols=1)
    lead.style = "Table Grid"
    lead.cell(0, 0).text = "BOTTOM LINE\n" + reader_summary(record, complete)
    shade(lead.cell(0, 0), "E8EEF5")
    evidence = {item["id"]: item for item in record.get("evidence", []) if isinstance(item, dict) and "id" in item}
    id_map = evidence_id_map(record)
    findings = record.get("findings", [])
    first_page_findings = validation.get("decision_implications") or findings[:3]
    document.add_heading("Decision implications", level=2)
    for finding in first_page_findings:
        if isinstance(finding, dict) and finding.get("evidence_ids"):
            p = document.add_paragraph(item_text(finding), style="List Bullet")
            cite_ids(p, finding.get("evidence_ids", []), id_map)
        else:
            document.add_paragraph(item_text(finding), style="List Bullet")
    document.add_heading("Next practical actions", level=2)
    lead_action_rows = structured_rows(
        validation.get("next_actions", []) or fallbacks.get("next_actions", []),
        [("owner", "Acquisition team"), ("action", "No approved next action was recorded."), ("output", "Before the related decision")],
    )
    add_table(document, ["Owner", "Action", "Output or gate"], lead_action_rows, [1.45, 3.85, 1.6])

    def add_closing_actions(heading: str, owner_default: str) -> None:
        """Cross-reference the lead table instead of repeating it verbatim."""
        document.add_heading(heading, level=1)
        rows = structured_rows(
            validation.get("next_actions", []),
            [("owner", owner_default), ("action", "Not recorded"), ("output", "Not recorded")],
        )
        if rows == lead_action_rows:
            document.add_paragraph(
                "The owned actions for this product are consolidated in the "
                "Next practical actions table at the start of this document."
            )
        else:
            add_table(document, ["Owner", "Action", "Output or gate"], rows, [1.45, 3.85, 1.6])

    if route == "complete_report" and not complete:
        note = document.add_paragraph()
        note.style = document.styles["Intense Quote"]
        note.add_run("Completion boundary: ").bold = True
        note.add_run("Missing " + ", ".join(missing_classes) + ". This product must remain a desk-research draft.")

    def add_findings_block(empty: str = "No approved finding was recorded.") -> None:
        if not findings:
            document.add_paragraph(empty)
        for finding in findings:
            p = document.add_paragraph(finding.get("text", ""))
            cite_ids(p, finding.get("evidence_ids", []), id_map)

    def add_unknowns() -> None:
        rows = []
        for item in record.get("unresolved_questions", []):
            if isinstance(item, dict):
                rows.append([
                    item.get("id", "U---"),
                    item.get("owner", "Acquisition team"),
                    item.get("question", item.get("text", "")),
                    item.get("gate", "Before the related reserved decision"),
                    item.get("evidence_needed", "Resolve through the approved research plan"),
                ])
            else:
                text = str(item)
                identifier, _, question = text.partition(":")
                rows.append([
                    identifier.strip() if identifier.strip().startswith("U") else "U---",
                    "Acquisition team",
                    question.strip() or text,
                    "Before the related reserved decision",
                    "Resolve through the approved research plan",
                ])
        add_table(
            document,
            ["ID", "Owner", "Unknown", "Decision gate", "Evidence or action needed"],
            rows or [["None", "-", "No unresolved item was recorded.", "-", "-"]],
            [0.55, 1.15, 2.25, 1.35, 1.6],
        )

    def add_decisions_to_close() -> None:
        decisions = record.get("user_decisions", [])
        unknowns = record.get("unresolved_questions", [])
        if not decisions and not unknowns:
            return
        document.add_heading("Decisions to close", level=1)
        if decisions:
            add_bullets(document, decisions)
        if unknowns:
            add_unknowns()

    scope = record.get("scope", {})
    if route == "complete_report":
        # Start the report body on a fresh page without an explicit break
        # paragraph: an empty paragraph that carries only a page-break run can
        # land alone on a page and render an entirely blank page when the
        # first-page content already spills past one page.
        frame_heading = document.add_heading("Acquisition and decision frame", level=1)
        frame_heading.paragraph_format.page_break_before = True
        document.add_paragraph(record.get("question", "Not provided"))
        add_table(document, ["Scope field", "Working value"], [[pretty_label(key), display_value(value)] for key, value in scope.items()], [2.0, 4.9])
        document.add_heading("Context and assumptions", level=2)
        add_bullets(document, record.get("user_context", []) + record.get("assumptions", []))

        document.add_heading("What the evidence establishes", level=1)
        add_findings_block()
        document.add_paragraph(validation.get("small_business_analysis", "No approved small-business or competition analysis was recorded."))
        document.add_paragraph(validation.get("pricing_analysis", "No approved pricing or contract-structure analysis was recorded."))

        document.add_heading("Market capability and packaging", level=1)
        add_table(
            document,
            ["Capability or hypothesis", "Evidence to request", "Failure signal or tradeoff"],
            structured_rows(
                validation.get("capability_model", [])
                + validation.get("packaging_hypotheses", [])
                or fallbacks["capability_model"],
                [("title", "Not recorded"), ("evidence_to_request", "Not recorded"), ("failure_signal", "Not recorded")],
            ),
            [1.8, 3.05, 2.05],
        )

        document.add_heading("Market engagement instrument", level=1)
        add_table(
            document,
            ["Theme", "Evidence-focused prompt", "Decision use"],
            structured_rows(
                validation.get("market_engagement_instrument", []) or fallbacks["market_engagement_instrument"],
                [("theme", "Not recorded"), ("prompt", "No approved instrument was recorded."), ("decision_use", "Not recorded")],
            ),
            [1.55, 3.85, 1.5],
        )

        document.add_heading("Evidence-to-decision gates", level=1)
        add_table(
            document,
            ["Gate", "Owner", "Exit condition", "Evidence of completion"],
            structured_rows(
                validation.get("decision_gates", []) or fallbacks["decision_gates"],
                [("gate", "Not recorded"), ("owner", "Acquisition team"), ("exit_condition", "Not recorded"), ("evidence", "Not recorded")],
            ),
            [0.85, 1.25, 2.3, 2.5],
        )

        add_decisions_to_close()

    elif route == "refresh":
        document.add_heading("Refresh assessment", level=1)
        document.add_paragraph(validation.get("change_assessment", validation.get("executive_summary", "No change assessment was recorded.")))
        document.add_heading("Prior-to-current evidence comparison", level=1)
        refresh_rows = require_route_value(record, "refresh_comparison", "a dated prior/current comparison")
        add_table(
            document,
            ["Decision area", "Prior baseline", "Current evidence", "Material delta", "Acquisition consequence"],
            [[item.get("decision_area", ""), item.get("prior_baseline", ""), with_evidence(item, "current_evidence", id_map), item.get("delta", ""), item.get("decision_impact", "")] for item in refresh_rows],
            [1.15, 1.35, 1.75, 1.25, 1.4],
        )
        document.add_heading("Vendor and market-structure changes", level=1)
        add_bullets(document, validation.get("vendor_landscape_changes", []), "No material vendor-landscape change was established.")
        document.add_heading("Strategy changes to make now", level=1)
        add_bullets(document, validation.get("strategy_changes", []), "No acquisition-strategy change was approved.")
        document.add_heading("What remains usable and what must be rechecked", level=1)
        add_table(
            document,
            ["Carry forward", "Recheck before use"],
            [[item_text(value), item_text(validation.get("recheck_items", [])[index]) if index < len(validation.get("recheck_items", [])) else "No paired recheck item"] for index, value in enumerate(validation.get("remains_usable", []))]
            or [["No prior conclusion was approved for carry-forward.", "Rebuild the baseline before relying on the refresh."]],
            [3.45, 3.45],
        )
        add_closing_actions("Refresh action plan", "Acquisition team")
        add_decisions_to_close()

    elif route == "one_question":
        document.add_heading("Bounded answer", level=1)
        document.add_paragraph(validation.get("executive_summary", "No approved bounded answer was recorded."))
        if validation.get("analysis_focus") == "small_business":
            document.add_heading("Candidate small-business market", level=1)
            candidates = require_route_value(record, "small_business_candidates", "a named candidate or documented search-result population")
            add_table(
                document,
                ["Concern", "Status / vehicles", "Relevant capability", "Recent federal evidence", "Gap to close"],
                [[item.get("name", ""), item.get("status_and_vehicles", ""), item.get("capability", ""), with_evidence(item, "recent_award_evidence", id_map), item.get("gap", "")] for item in candidates],
                [1.3, 1.35, 1.65, 1.6, 1.0],
            )
            document.add_heading("Rule of Two evidence assessment", level=1)
            document.add_paragraph(validation.get("rule_of_two_assessment", ""))
            document.add_heading("Evidence supporting and cutting against a small-business strategy", level=1)
            add_table(
                document,
                ["Supports", "Cuts against / remains unknown"],
                [[item_text(value), item_text(validation.get("contrary_evidence", [])[index]) if index < len(validation.get("contrary_evidence", [])) else "No paired contrary item"] for index, value in enumerate(validation.get("supporting_evidence", []))]
                or [["No approved supporting evidence was recorded.", "The market conclusion remains open."]],
                [3.45, 3.45],
            )
            document.add_heading("Targeted outreach plan", level=1)
            add_table(
                document,
                ["Owner", "Outreach action", "Proof requested", "Decision use"],
                structured_rows(validation.get("outreach_plan", []), [("owner", "Market research lead"), ("action", "Not recorded"), ("proof_requested", "Not recorded"), ("decision_use", "Not recorded")]),
                [1.2, 2.1, 2.2, 1.4],
            )
        else:
            document.add_heading("Evidence for and against", level=1)
            add_findings_block()
            add_bullets(document, record.get("conflicts", []), "No contrary evidence or conflict was recorded.")
        document.add_heading("Decision implications", level=1)
        add_bullets(document, validation.get("decision_implications", []))
        document.add_heading("Further research options", level=1)
        if validation.get("next_actions"):
            document.add_paragraph(
                "Further research actions are consolidated in the Next practical actions table at the start of this document."
            )
        else:
            add_bullets(document, [])
        add_decisions_to_close()

    elif route == "pre_award_handoff":
        document.add_heading("Handoff summary", level=1)
        document.add_paragraph(validation.get("executive_summary", "No approved handoff summary was recorded."))
        document.add_heading("Approved market observations", level=1)
        add_findings_block()
        document.add_heading("Market findings translated into acquisition inputs", level=1)
        add_table(
            document,
            ["Pre-Award area", "Approved implication", "Source boundary", "Owner / decision gate"],
            [
                [label, item.get("implication", ""), with_evidence(item, "source_boundary", id_map), item.get("owner_gate", "")]
                for label, field in (
                    ("Scope", "scope_implications"),
                    ("Packaging", "packaging_implications"),
                    ("Performance", "performance_implications"),
                    ("Competition", "competition_implications"),
                )
                for item in route_payload(record, field)
            ],
            [1.0, 2.5, 1.8, 1.6],
        )
        document.add_heading("Pricing inputs and boundaries", level=1)
        add_table(
            document,
            ["Input", "Usable evidence", "Do not infer", "Next owner"],
            [[item.get("input", ""), with_evidence(item, "usable_evidence", id_map), item.get("boundary", ""), item.get("owner", "")] for item in require_route_value(record, "pricing_inputs", "pricing inputs and boundaries")],
            [1.2, 2.4, 2.2, 1.1],
        )
        document.add_heading("Pre-Award risk register", level=1)
        add_table(
            document,
            ["Risk", "Why it matters", "Mitigation / evidence gate", "Owner"],
            [[item.get("risk", ""), item.get("why", ""), with_evidence(item, "mitigation", id_map), item.get("owner", "")] for item in require_route_value(record, "handoff_risks", "a risk register")],
            [1.3, 2.0, 2.6, 1.0],
        )
        add_closing_actions("Pre-Award intake and next actions", "Pre-Award lead")
        add_decisions_to_close()

    else:
        raise ValueError(f"unsupported workflow_mode: {route}")

    methodology = str(validation.get("methodology") or "").strip()
    conflicts = record.get("conflicts", [])
    if methodology or conflicts:
        document.add_heading("Method and material limitations", level=1)
        if methodology:
            document.add_paragraph(methodology)
        if conflicts:
            add_bullets(document, conflicts)

    document.add_heading("Source Register", level=1)
    for row in source_register_rows(record, id_map):
        item = next(
            (
                value
                for value in record.get("evidence", [])
                if isinstance(value, dict) and id_map.get(value.get("id")) == row[0]
            ),
            {},
        )
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(3)
        paragraph.add_run(f"[{row[0]}] ").bold = True
        paragraph.add_run(f"{row[1]}. {item.get('title', '')}. ")
        locator = str(item.get("locator", ""))
        url = str(item.get("canonical_url") or (locator if locator.startswith("http") else ""))
        if url:
            add_hyperlink(paragraph, url, url)
        elif locator:
            paragraph.add_run(locator)
        date = item.get("as_of_date") or item.get("retrieved_at")
        if date:
            paragraph.add_run(f". {date}")

    numeric_checks = record.get("validation", {}).get("numeric_checks", []) if route == "complete_report" else []
    for index, check in enumerate(numeric_checks):
        components = [float(value) for value in check.get("components", [])]
        locator = f"validation.numeric_checks[{index}]"
        calculation_ids = [
            item.get("id")
            for item in record.get("evidence", [])
            if isinstance(item, dict)
            and item.get("source_class") == "calculation"
            and item.get("locator") == locator
        ]
        if len(calculation_ids) != 1:
            raise ValueError(
                f"numeric check {index} requires exactly one calculation evidence item whose locator is {locator}"
            )
        paragraph = document.add_paragraph(format_calculated_total(check.get("label", "Calculated total"), sum(components)))
        cite_ids(paragraph, calculation_ids, id_map)

    report_inferences = record.get("inferences", []) if route == "complete_report" else validation.get("inferences", [])
    for item in report_inferences:
        p = document.add_paragraph("Inference: " + item.get("text", item.get("reasoning", "")), style="List Bullet")
        cite_ids(p, item.get("evidence_ids", []), id_map)

    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("record", type=Path)
    parser.add_argument("output", type=Path)
    args = parser.parse_args()
    validator = Path(__file__).with_name("validate_research_record.py")
    result = subprocess.run([sys.executable, str(validator), str(args.record)], check=False)
    if result.returncode:
        return result.returncode
    record = json.loads(args.record.read_text(encoding="utf-8"))
    build(record, args.output)
    print(args.output)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
