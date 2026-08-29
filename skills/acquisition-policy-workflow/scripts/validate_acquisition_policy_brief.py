#!/usr/bin/env python3
"""Validate a generated Acquisition Policy Impact Brief."""

from __future__ import annotations

import argparse
import importlib.util
import json
import re
import sys
import zipfile
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


REQUIRED_HEADINGS = [
    "Planning Posture and Implications",
    "Owners and Decision Gates",
    "Question and Scope",
    "Documented Current Status",
    "Source Hierarchy and Authorities",
    "Planning Scenarios",
    "Change Timeline",
    "Government and Industry Impacts",
    "Open Issues and Comment Deadlines",
    "Operational Considerations",
    "Evidence Register",
    "Limitations and Reserved Determinations",
]
# Route-native payload Heading 1 sections for each focused product; see
# references/report-specification.md ("Per-route required payload and the
# shared-framing cap"). The full twelve-section REQUIRED_HEADINGS structure is
# reserved for the Acquisition Policy Impact Brief.
FOCUSED_HEADINGS = {
    "current_rule": ["Current Rule Card"],
    "agency_status": ["Agency Adoption Status"],
    "three_layer": ["Three-Layer Comparison and Adoption Test"],
    "change_brief": ["Before/After Change Map"],
    "rulemaking": ["Rulemaking Milestones and Next Trigger"],
    "watchlist": ["Open Rulemaking Watchlist"],
    "comments": ["Comment Sample and Theme Coverage", "Coded themes and acquisition implications"],
    "refresh": ["Refresh Change Register"],
}
FORBIDDEN = [
    re.compile(r"\bmcp__|/mnt/|/Users/|[A-Za-z]:\\", re.I),
    re.compile(r"\bgh[pousr]_[A-Za-z0-9]{20,}\b|\b(?:sk|cfat|SAM)-[A-Za-z0-9_-]{16,}\b", re.I),
    re.compile(r"\b(?:this model deviation|model deviation text)\s+(?:is|was)\s+(?:legally )?(?:operative|applicable)\b", re.I),
]
BOUNDARY_PHRASES = (
    "authorized agency official must determine procurement-specific applicability",
    "does not provide legal advice",
)
# Documented-status vocabulary for status cells; see references/report-specification.md.
CODIFIED_BASELINE_LABELS = {"Government-wide baseline", "Codified current baseline"}
STATUS_CELL_VOCABULARY = CODIFIED_BASELINE_LABELS | {
    "Published model text; not agency-operative",
    "Published model text; not operative alone",
    "Pending rulemaking; not current policy",
    "Final rule pending effective date; not current policy",
    "Withdrawn; not current policy",
    "Superseded; not current policy",
    "Guidance; not regulation",
    "Named-agency evidence",
    "Comparator only; does not establish adoption for the named agency",
    "Agency class deviation",
}
STATUS_COLUMN_HEADERS = {"Status for this question", "Documented status"}
NONBASELINE_LAYER = re.compile(r"\b(?:proposed|deviation|model)\b", re.I)


def load_record_validator(path: Path):
    spec = importlib.util.spec_from_file_location("policy_record_validator", path)
    module = importlib.util.module_from_spec(spec)
    assert spec.loader is not None
    spec.loader.exec_module(module)
    return module


def all_text(document: Document) -> str:
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


def external_hyperlinks(document: Document) -> set[str]:
    urls: set[str] = set()
    for relationship in document.part.rels.values():
        if relationship.reltype.endswith("/hyperlink") and relationship.is_external:
            urls.add(relationship.target_ref)
    return urls


def collect_evidence_ids(value: object) -> set[str]:
    """Collect evidence IDs from the portion of the approved record rendered by a focused product."""
    ids: set[str] = set()
    if isinstance(value, dict):
        for key, item in value.items():
            if key == "evidence_ids" and isinstance(item, list):
                ids.update(str(evidence_id) for evidence_id in item if evidence_id)
            else:
                ids.update(collect_evidence_ids(item))
    elif isinstance(value, list):
        for item in value:
            ids.update(collect_evidence_ids(item))
    return ids


def focused_evidence_ids(record: dict) -> set[str]:
    mode = record.get("workflow_mode", "")
    validation = record.get("validation", {})
    route_fields = {
        "change_brief": ("focused_findings", "focused_impacts", "planning_posture", "decision_gates", "change_map", "implementation_actions"),
        "watchlist": ("focused_findings", "focused_impacts", "planning_posture", "decision_gates", "rulemaking_watchlist", "watch_priorities"),
        "comments": ("focused_findings", "focused_impacts", "planning_posture", "decision_gates", "comment_themes"),
        "refresh": ("focused_findings", "focused_impacts", "planning_posture", "decision_gates", "refresh_changes", "carry_forward_decisions"),
    }
    ids = collect_evidence_ids({field: validation.get(field) for field in route_fields.get(mode, ())})
    # Record-native routes render evidence from the approved record itself, not
    # from validation payload rows; require that cited evidence in the product.
    if mode == "current_rule":
        ids |= collect_evidence_ids([
            item
            for item in record.get("policy_items", [])
            if item.get("status") in {"codified_current", "model_deviation", "agency_class_deviation"}
        ])
    elif mode in {"agency_status", "three_layer"}:
        ids |= collect_evidence_ids(record.get("policy_items", []))
    elif mode == "rulemaking":
        ids |= collect_evidence_ids(record.get("timeline", []))
    if mode in {"current_rule", "agency_status", "three_layer", "rulemaking"}:
        ids |= collect_evidence_ids({field: validation.get(field) for field in ("planning_posture", "decision_gates")})
    return ids


def status_vocabulary_failures(document: Document) -> list[str]:
    failures: list[str] = []
    for table_index, table in enumerate(document.tables):
        if not table.rows:
            continue
        headers = [cell.text.strip() for cell in table.rows[0].cells]
        for column_index, header in enumerate(headers):
            if header not in STATUS_COLUMN_HEADERS:
                continue
            for row_index, row in enumerate(table.rows[1:], start=2):
                value = row.cells[column_index].text.strip()
                if value not in STATUS_CELL_VOCABULARY:
                    failures.append(
                        f"table {table_index + 1} row {row_index} status cell is outside the "
                        f"documented-status vocabulary: {value!r}"
                    )
                layer = row.cells[0].text.strip()
                if value in CODIFIED_BASELINE_LABELS and NONBASELINE_LAYER.search(layer):
                    failures.append(
                        f"table {table_index + 1} row {row_index} labels a non-baseline layer "
                        f"({layer!r}) with the codified-baseline status"
                    )
    return failures


def table_geometry_failures(document: Document) -> list[str]:
    failures: list[str] = []
    for table_index, table in enumerate(document.tables):
        tbl_pr = table._tbl.tblPr
        tbl_w = tbl_pr.find(qn("w:tblW"))
        tbl_ind = tbl_pr.find(qn("w:tblInd"))
        layout = tbl_pr.find(qn("w:tblLayout"))
        if tbl_w is None or tbl_w.get(qn("w:type")) != "dxa" or tbl_w.get(qn("w:w")) != "9360":
            failures.append(f"table {table_index + 1} does not have fixed 9360-DXA width")
        if tbl_ind is None or tbl_ind.get(qn("w:w")) != "120":
            failures.append(f"table {table_index + 1} does not have 120-DXA indent")
        if layout is None or layout.get(qn("w:type")) != "fixed":
            failures.append(f"table {table_index + 1} does not use fixed layout")
        grid_widths = [int(col.get(qn("w:w"), "0")) for col in table._tbl.tblGrid]
        if not grid_widths or sum(grid_widths) != 9360:
            failures.append(f"table {table_index + 1} grid widths do not total 9360 DXA")
        header_props = table.rows[0]._tr.get_or_add_trPr()
        if header_props.find(qn("w:tblHeader")) is None:
            failures.append(f"table {table_index + 1} does not repeat its header row")
        for row_index, row in enumerate(table.rows):
            widths = []
            for cell in row.cells:
                tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
                widths.append(int(tc_w.get(qn("w:w"), "0")) if tc_w is not None else 0)
            if grid_widths and widths != grid_widths:
                failures.append(f"table {table_index + 1} row {row_index + 1} cell widths differ from grid")
                break
    return failures


def validate(document_path: Path, record_path: Path) -> dict:
    failures: list[str] = []
    if not zipfile.is_zipfile(document_path):
        return {"status": "fail", "failures": ["file is not a valid DOCX ZIP"]}

    record = json.loads(record_path.read_text(encoding="utf-8"))
    record_validator = load_record_validator(Path(__file__).with_name("validate_policy_research_record.py"))
    record_result = record_validator.validate_record(record)
    if record_result["status"] != "pass":
        failures.extend(f"record: {failure}" for failure in record_result["failures"])

    document = Document(document_path)
    text = all_text(document)
    headings = [
        paragraph.text.strip()
        for paragraph in document.paragraphs
        if getattr(paragraph.style, "name", "") == "Heading 1"
    ]
    mode = record.get("workflow_mode", "")
    required_headings = (
        [
            "Planning Posture and Implications",
            "Owners and Decision Gates",
            *FOCUSED_HEADINGS[mode],
            "Management Actions",
            "Evidence and Source Notes",
            "Limitations and Reserved Determinations",
        ]
        if mode in FOCUSED_HEADINGS
        else REQUIRED_HEADINGS
    )
    for heading in required_headings:
        if heading not in headings:
            failures.append(f"missing Heading 1 section: {heading}")
    if [heading for heading in headings if heading in required_headings] != required_headings:
        failures.append("required Heading 1 sections are out of order")

    for pattern in FORBIDDEN:
        if pattern.search(text):
            failures.append(f"forbidden content matched: {pattern.pattern}")
    lowered = text.lower()
    for phrase in BOUNDARY_PHRASES:
        if phrase not in lowered:
            failures.append(f"required decision-boundary language is missing: {phrase}")
    product_elements = ("Decision-ready evidence", "Owner", "Timing") if mode in FOCUSED_HEADINGS else ("Decision-ready evidence", "Owner", "Timing", "Scenario", "Planning treatment")
    for product_element in product_elements:
        if product_element not in text:
            failures.append(f"required reader-facing product element is missing: {product_element}")
    if not headings or headings[0] != "Planning Posture and Implications":
        failures.append("the first Heading 1 must state the planning posture and implications")
    as_of = record.get("scope", {}).get("as_of_date", "")
    if as_of and as_of not in text:
        failures.append("record as-of date is missing from the brief")
    scope = record.get("scope", {})
    for field, label in (("customer_organization", "customer organization"), ("decision_date", "decision date")):
        value = scope.get(field, "")
        if value and str(value) not in text:
            failures.append(f"record {label} is missing from the brief scope header")
    if mode == "refresh":
        prior = scope.get("prior_analysis") or {}
        for field, label in (("title", "prior-analysis title"), ("date", "prior-analysis date")):
            value = prior.get(field, "")
            if not value or str(value) not in text:
                failures.append(f"refresh product must identify the {label} in the scope header")

    validation = record.get("validation", {})
    focused = mode in FOCUSED_HEADINGS
    findings = validation.get("focused_findings", []) if focused else record.get("findings", [])
    relevant_evidence_ids = (
        focused_evidence_ids(record)
        if focused
        else {item.get("id") for item in record.get("evidence", []) if isinstance(item, dict)}
    )
    for finding in findings:
        for evidence_id in finding.get("evidence_ids", []):
            if evidence_id not in text:
                failures.append(f"finding evidence ID not present in brief: {evidence_id}")
    if mode not in FOCUSED_HEADINGS:
        for policy in record.get("policy_items", []):
            policy_id = policy.get("id")
            if policy_id and policy_id not in text:
                failures.append(f"policy item ID not present in brief: {policy_id}")
    for evidence_id in sorted(relevant_evidence_ids):
        if evidence_id not in text:
            failures.append(f"evidence register is missing ID: {evidence_id}")

    urls = external_hyperlinks(document)
    for item in record.get("evidence", []):
        if focused and item.get("id") not in relevant_evidence_ids:
            continue
        url = item.get("canonical_url", "")
        if url and url not in urls:
            failures.append(f"evidence URL is not a live DOCX hyperlink: {url}")

    failures.extend(status_vocabulary_failures(document))
    failures.extend(table_geometry_failures(document))
    if len(document.tables) < 3:
        failures.append("brief must contain at least three structured evidence tables")

    return {
        "status": "pass" if not failures else "fail",
        "heading_count": len(headings),
        "table_count": len(document.tables),
        "hyperlink_count": len(urls),
        "failures": failures,
    }


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("document", type=Path)
    parser.add_argument("--record", required=True, type=Path)
    parser.add_argument("--json", action="store_true")
    args = parser.parse_args()
    try:
        result = validate(args.document, args.record)
    except (OSError, ValueError, json.JSONDecodeError, zipfile.BadZipFile) as exc:
        print(f"ERROR: {exc}", file=sys.stderr)
        return 2
    if args.json:
        print(json.dumps(result, indent=2, sort_keys=True))
    elif result["status"] == "pass":
        print("Acquisition Policy DOCX validation passed.")
    else:
        print("VALIDATION FAILED")
        for failure in result["failures"]:
            print(f"- {failure}")
    return 0 if result["status"] == "pass" else 1


if __name__ == "__main__":
    raise SystemExit(main())
