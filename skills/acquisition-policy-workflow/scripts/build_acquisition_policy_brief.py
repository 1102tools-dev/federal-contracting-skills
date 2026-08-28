#!/usr/bin/env python3
"""Build a validated Acquisition Policy Impact Brief from a policy record."""

from __future__ import annotations

import argparse
import json
import subprocess
import sys
from pathlib import Path
from urllib.parse import urlparse

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Inches, Pt, RGBColor


BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
PALE_BLUE = "E8EEF5"
PALE_TEAL = "E8F3F1"
PALE_GOLD = "FFF4D6"
PALE_RED = "FCE8E6"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "5B6573"
BLACK = "202124"
WHITE = "FFFFFF"
TEAL = "2F6F75"
GOLD = "B7791F"
RED = "A33A2B"
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGINS_DXA = {"top": 80, "bottom": 80, "start": 120, "end": 120}

PRODUCT_TITLES = {
    "current_rule": "Current Rule Explanation",
    "agency_status": "Agency Policy Status Matrix",
    "three_layer": "Three-Layer Policy Comparison",
    "change_brief": "Regulatory Change Briefing",
    "rulemaking": "Rulemaking Timeline",
    "watchlist": "Open Rulemaking Watchlist",
    "comments": "Public Comment Position Analysis",
    "refresh": "Policy Analysis Refresh",
    "impact_brief": "Acquisition Policy Impact Brief",
}

POLICY_ITEM_STATUSES = {
    "codified_current",
    "model_deviation",
    "agency_class_deviation",
    "proposed_rule",
    "final_rule_pending_effective",
    "final_rule_effective",
    "withdrawn",
    "superseded",
    "nonregulatory_guidance",
}

# Reader-facing documented-status vocabulary. The Agency Policy Status Matrix
# and Three-Layer Comparison status cells must use only these values; see
# references/report-specification.md.
STATUS_QUESTION_LABELS = {
    "codified_current": "Government-wide baseline",
    "final_rule_effective": "Government-wide baseline",
    "model_deviation": "Published model text; not agency-operative",
    "proposed_rule": "Pending rulemaking; not current policy",
    "final_rule_pending_effective": "Final rule pending effective date; not current policy",
    "withdrawn": "Withdrawn; not current policy",
    "superseded": "Superseded; not current policy",
    "nonregulatory_guidance": "Guidance; not regulation",
}

FOCUSED_PRODUCTS = {
    "current_rule",
    "agency_status",
    "three_layer",
    "change_brief",
    "rulemaking",
    "watchlist",
    "comments",
    "refresh",
}


def set_run_font(run, *, name: str = "Calibri", size: float | None = None, color: str | None = None,
                 bold: bool | None = None, italic: bool | None = None) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_cell_margins(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in CELL_MARGINS_DXA.items():
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int]) -> None:
    if sum(widths_dxa) != CONTENT_WIDTH_DXA:
        raise ValueError("table widths must total 9360 DXA")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    indent = tbl_pr.find(qn("w:tblInd"))
    if indent is None:
        indent = OxmlElement("w:tblInd")
        tbl_pr.append(indent)
    indent.set(qn("w:w"), str(TABLE_INDENT_DXA))
    indent.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        cant_split = OxmlElement("w:cantSplit")
        row._tr.get_or_add_trPr().append(cant_split)
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[index])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_hyperlink(paragraph, text: str, url: str, *, font_size: float | None = None) -> None:
    if not url:
        return
    relationship_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_properties.extend([color, underline])
    if font_size is not None:
        size = OxmlElement("w:sz")
        size.set(qn("w:val"), str(int(font_size * 2)))
        run_properties.append(size)
    run.append(run_properties)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def configure_styles(document: Document) -> None:
    document.settings.odd_and_even_pages_header_footer = False
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = document.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = document.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167

    if "Policy Title" not in document.styles:
        title = document.styles.add_style("Policy Title", WD_STYLE_TYPE.PARAGRAPH)
    else:
        title = document.styles["Policy Title"]
    title.font.name = "Calibri"
    title.font.size = Pt(24)
    title.font.bold = True
    title.font.color.rgb = RGBColor.from_string(BLACK)
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(4)
    title.paragraph_format.keep_with_next = True
    # Keep the display title inside the page text frame in every renderer.
    # Some DOCX-to-PDF engines otherwise preserve an inherited hanging indent
    # or refuse to wrap the longest route names cleanly.
    title.paragraph_format.left_indent = Inches(0)
    title.paragraph_format.right_indent = Inches(0)
    title.paragraph_format.first_line_indent = Inches(0)

    if "Evidence ID" not in document.styles:
        evidence_style = document.styles.add_style("Evidence ID", WD_STYLE_TYPE.CHARACTER)
    else:
        evidence_style = document.styles["Evidence ID"]
    evidence_style.font.name = "Calibri"
    evidence_style.font.size = Pt(9)
    evidence_style.font.bold = True
    evidence_style.font.color.rgb = RGBColor.from_string(DARK_BLUE)

    header = section.header.paragraphs[0]
    header.text = "1102tools  |  Acquisition Policy Impact Brief"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run_font(header.runs[0], size=8, color=MID_GRAY)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    prefix = footer.add_run("Prepared as of the date shown  |  Page ")
    set_run_font(prefix, size=8, color=MID_GRAY)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    footer._p.append(field)


def add_metadata(document: Document, label: str, value: object) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(2)
    label_run = paragraph.add_run(f"{label}: ")
    set_run_font(label_run, bold=True)
    value_run = paragraph.add_run(str(value if value not in (None, "") else "Not provided"))
    set_run_font(value_run)


def set_cell_text(
    cell, value: object, *, bold: bool = False, color: str | None = None, font_size: float = 9.5
) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.05
    run = paragraph.add_run(str(value if value not in (None, "") else "Not stated"))
    set_run_font(run, size=font_size, color=color, bold=bold)


def add_table(
    document: Document,
    headers: list[str],
    rows: list[list[object]],
    widths_dxa: list[int],
    *,
    font_size: float = 9.5,
    add_spacer: bool = True,
):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[index], header, bold=True, color=WHITE, font_size=font_size)
        shade(table.rows[0].cells[index], DARK_BLUE)
    set_repeat_table_header(table.rows[0])
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        for index, value in enumerate(values):
            set_cell_text(cells[index], value, font_size=font_size)
            if row_index % 2:
                shade(cells[index], LIGHT_GRAY)
    set_table_geometry(table, widths_dxa)
    if add_spacer:
        spacer = document.add_paragraph()
        spacer.paragraph_format.space_after = Pt(2)
    return table


def add_bullets(document: Document, items: list[object], empty_text: str) -> None:
    if not items:
        document.add_paragraph(empty_text)
        return
    for item in items:
        if isinstance(item, dict):
            text = item.get("text") or item.get("question") or item.get("description") or json.dumps(item, sort_keys=True)
            evidence_ids = item.get("evidence_ids", [])
        else:
            text = str(item)
            evidence_ids = []
        paragraph = document.add_paragraph(str(text), style="List Bullet")
        cite_ids(paragraph, evidence_ids)


def cite_ids(paragraph, ids: list[str]) -> None:
    if not ids:
        return
    paragraph.add_run(" [")
    paragraph.add_run(", ".join(ids), style="Evidence ID")
    paragraph.add_run("]")


def display_value(value: object) -> str:
    if value in (None, "", [], {}):
        return "Not stated"
    if isinstance(value, dict):
        value = {key: item for key, item in value.items() if item not in (None, "", [], {})}
        if not value:
            return "Not stated"
        return "; ".join(f"{pretty_label(key)}: {display_value(item)}" for key, item in value.items())
    if isinstance(value, list):
        return ", ".join(display_value(item) for item in value)
    return str(value)


def pretty_label(key: object) -> str:
    """Format metadata labels for a reader instead of exposing field names."""
    label = str(key).replace("_", " ").title()
    replacements = {
        "As Of Date": "As of date",
        "As Of": "As of",
        "Naics": "NAICS",
        "Psc": "PSC",
        "Far Parts": "FAR parts",
        "Docket Id": "Docket ID",
        "Id": "ID",
    }
    for source, target in replacements.items():
        label = label.replace(source, target)
    return label


def policy_rows(record: dict) -> list[list[object]]:
    return [
        [
            item.get("id", ""),
            item.get("status", "").replace("_", " ").title(),
            item.get("citation", ""),
            item.get("agency", "") or "Government-wide or not stated",
            item.get("applicability_summary", ""),
            ", ".join(item.get("evidence_ids", [])),
        ]
        for item in record.get("policy_items", [])
    ]


def impact_items(record: dict, lens: str) -> list[dict]:
    impacts = record.get("validation", {}).get("impacts", {})
    if not isinstance(impacts, dict):
        return []
    if lens == "neutral":
        return list(impacts.get("government", [])) + list(impacts.get("industry", []))
    return list(impacts.get(lens, []))


def evidence_ids_for_statuses(record: dict, statuses: set[str]) -> list[str]:
    ids: list[str] = []
    for item in record.get("policy_items", []):
        if item.get("status") not in statuses:
            continue
        for evidence_id in item.get("evidence_ids", []):
            if evidence_id and evidence_id not in ids:
                ids.append(evidence_id)
    return ids


def normalized_agency(value: object) -> str:
    text = str(value or "").lower()
    for suffix in ("(fictional test data)", "(fictional)"):
        text = text.replace(suffix, "")
    return " ".join(text.split())


def agency_item_matches_scope(record: dict, item: dict) -> bool:
    scope_agency = normalized_agency(record.get("scope", {}).get("agency"))
    item_agency = normalized_agency(item.get("agency"))
    return bool(scope_agency and item_agency and scope_agency == item_agency)


def derive_planning_posture(record: dict) -> dict:
    supplied = record.get("validation", {}).get("planning_posture")
    if isinstance(supplied, dict) and supplied.get("headline") and supplied.get("rationale"):
        return {
            "label": supplied.get("label") or "Planning posture",
            "headline": supplied["headline"],
            "rationale": supplied["rationale"],
            "evidence_ids": list(supplied.get("evidence_ids", [])),
            "fill": supplied.get("fill") or PALE_GOLD,
            "accent": supplied.get("accent") or GOLD,
        }

    agency = record.get("scope", {}).get("agency") or "the agency in scope"
    matched_deviation = next(
        (
            item
            for item in record.get("policy_items", [])
            if item.get("status") == "agency_class_deviation"
            and agency_item_matches_scope(record, item)
        ),
        None,
    )
    current_rule_rationale = (
        f"The codified rule remains the government-wide baseline. This record also identifies a {agency} "
        "deviation, which applies only within its documented scope, effective date, and transition terms."
        if matched_deviation
        else (
            "The model text and proposed rule do not replace the codified rule, and this record does not "
            f"establish a {agency} deviation. Confirm agency adoption before departing from the codified baseline."
        )
    )
    agency_status_headline = (
        f"{agency} adoption is documented; confirm scope and transition timing"
        if matched_deviation
        else f"No {agency} adoption is established"
    )
    agency_status_rationale = (
        "The named-agency deviation is adoption evidence for this status question. The contracting team must "
        "still map the contemplated procurement to its covered scope, effective date, and transition terms."
        if matched_deviation
        else (
            "The approved record includes comparator evidence, not a deviation issued for the agency in scope. "
            "The policy office should locate and authenticate any agency-specific issuance before the contracting "
            "team treats the model text as operative."
        )
    )
    recorded_deviation = item_with_status(record, "agency_class_deviation")
    deviation_issuer = recorded_deviation.get("agency") or "issuing-agency"
    three_layer_rationale = (
        "The codified FAR is the government-wide baseline; the FAR Council text is a non-operative model; and "
        f"the documented {agency} deviation is the agency-operative layer only within its recorded scope and timing."
        if matched_deviation
        else (
            "The codified FAR is current; the FAR Council text is a non-operative model; and the documented "
            f"{deviation_issuer} deviation applies only within its own scope. No approved evidence shows adoption by {agency}."
        )
    )
    baseline_ids = evidence_ids_for_statuses(record, {"codified_current", "final_rule_effective"})
    adoption_ids = evidence_ids_for_statuses(record, {"model_deviation", "agency_class_deviation"})
    rulemaking_ids = evidence_ids_for_statuses(
        record, {"proposed_rule", "final_rule_pending_effective", "final_rule_effective"}
    )
    all_ids = evidence_ids_for_statuses(record, POLICY_ITEM_STATUSES)
    route_postures = {
        "current_rule": {
            "label": "Current rule",
            "headline": "Use the codified FAR Part 10 text as the current baseline",
            "rationale": current_rule_rationale,
            "evidence_ids": baseline_ids + [i for i in adoption_ids if i not in baseline_ids],
            "fill": PALE_BLUE,
            "accent": BLUE,
        },
        "agency_status": {
            "label": "Agency status",
            "headline": agency_status_headline,
            "rationale": agency_status_rationale,
            "evidence_ids": adoption_ids,
            "fill": PALE_GOLD,
            "accent": GOLD,
        },
        "three_layer": {
            "label": "Comparison answer",
            "headline": "The three layers do not produce one common operative rule",
            "rationale": three_layer_rationale,
            "evidence_ids": baseline_ids + [i for i in adoption_ids if i not in baseline_ids],
            "fill": PALE_TEAL,
            "accent": TEAL,
        },
        "change_brief": {
            "label": "Change finding",
            "headline": "No defensible before-and-after policy delta can be stated",
            "rationale": (
                "The approved record identifies codified, model, deviation, and proposed-rule layers but supplies "
                "no matched section-level text. Obtain the old and new provisions before describing changed duties, "
                "thresholds, or procedures."
            ),
            "evidence_ids": baseline_ids + [i for i in adoption_ids + rulemaking_ids if i not in baseline_ids],
            "fill": PALE_GOLD,
            "accent": GOLD,
        },
        "rulemaking": {
            "label": "Rulemaking status",
            "headline": "The record reaches a proposed rule, not a final effective rule",
            "rationale": (
                "The documented sequence of model text, any recorded deviation, and related rulemaking reaches a proposed rule. "
                "No final-rule or effective-date event is approved here; monitor the docket for the next formal trigger."
            ),
            "evidence_ids": adoption_ids + [i for i in rulemaking_ids if i not in adoption_ids],
            "fill": PALE_BLUE,
            "accent": BLUE,
        },
        "watchlist": {
            "label": "Deadline status",
            "headline": "No verified open comment deadline is established",
            "rationale": (
                "The record identifies a proposed-rule docket but does not supply a verified open period or closing "
                "date. Check the live docket and Federal Register notice before assigning a response deadline."
            ),
            "evidence_ids": rulemaking_ids,
            "fill": PALE_GOLD,
            "accent": GOLD,
        },
        "comments": {
            "label": "Analysis status",
            "headline": "No comment position is supportable from the approved record",
            "rationale": (
                "No bounded public-comment sample was supplied, so stakeholder themes, prevalence, and positions "
                "cannot be responsibly characterized. Define the sampling frame, retrieve the comments, and code "
                "the sample before drawing conclusions."
            ),
            "evidence_ids": rulemaking_ids,
            "fill": PALE_RED,
            "accent": RED,
        },
        "refresh": {
            "label": "Refresh result",
            "headline": "No material refresh delta can be established",
            "rationale": (
                "The approved record does not provide a prior analysis snapshot paired with newly retrieved sources. "
                "Preserve the prior source register, retrieve current versions, and compare status, text, scope, and "
                "dates before reporting a change."
            ),
            "evidence_ids": all_ids,
            "fill": PALE_GOLD,
            "accent": GOLD,
        },
    }
    route_posture = route_postures.get(record.get("workflow_mode", ""))
    if route_posture:
        return route_posture

    unresolved_conflicts = [
        item for item in record.get("conflicts", []) if item.get("status") == "unresolved"
    ]
    statuses = {item.get("status") for item in record.get("policy_items", [])}
    agency_items = [
        item
        for item in record.get("policy_items", [])
        if item.get("status") == "agency_class_deviation"
        and item.get("operative_for_agency")
        and agency_item_matches_scope(record, item)
    ]
    if unresolved_conflicts:
        return {
            "label": "Hold point",
            "headline": "Do not operationalize the disputed policy value",
            "rationale": (
                "The approved record contains an unresolved material conflict. Preserve both source positions "
                "and obtain an authorized resolution before using either value in acquisition execution."
            ),
            "evidence_ids": sorted(
                {
                    evidence_id
                    for item in unresolved_conflicts
                    for evidence_id in item.get("evidence_ids", [])
                }
            ),
            "fill": PALE_RED,
            "accent": RED,
        }
    if agency_items:
        return {
            "label": "Conditional",
            "headline": "Implement only within the documented agency scope",
            "rationale": (
                "An agency-issued deviation is represented in the approved record. Confirm that the actual "
                "procurement and its timing fall within the issuing document's scope and transition terms."
            ),
            "evidence_ids": evidence_ids_for_statuses(record, {"agency_class_deviation"}),
            "fill": PALE_TEAL,
            "accent": TEAL,
        }
    if statuses & {"model_deviation", "proposed_rule", "final_rule_pending_effective"}:
        return {
            "label": "Planning baseline",
            "headline": "Plan from the codified baseline and monitor the non-operative layers",
            "rationale": (
                "Model text, proposed rules, and future-effective material do not replace the current baseline "
                "by themselves. Require agency adoption or effective-rule evidence before operationalizing them."
            ),
            "evidence_ids": evidence_ids_for_statuses(
                record,
                {"codified_current", "model_deviation", "proposed_rule", "final_rule_pending_effective"},
            ),
            "fill": PALE_GOLD,
            "accent": GOLD,
        }
    return {
        "label": "Planning baseline",
        "headline": "Proceed from the documented current policy baseline",
        "rationale": (
            "Use the approved current-status finding for planning, then refresh when a material source, "
            "effective date, agency instruction, or procurement date changes."
        ),
        "evidence_ids": evidence_ids_for_statuses(record, {"codified_current", "final_rule_effective"}),
        "fill": PALE_BLUE,
        "accent": BLUE,
    }


def derive_front_page_interpretation(record: dict) -> str:
    approved = str(record.get("validation", {}).get("reader_bottom_line", "")).strip()
    if approved:
        return approved
    agency = record.get("scope", {}).get("agency") or "the agency in scope"
    matched_deviation = any(
        item.get("status") == "agency_class_deviation" and agency_item_matches_scope(record, item)
        for item in record.get("policy_items", [])
    )
    interpretations = {
        "current_rule": (
            "Bottom line: codified FAR Part 10 is the government-wide baseline, and the approved record also "
            f"documents a {agency} deviation for use only within its stated scope and timing. The model text and "
            "proposed rule remain separate planning layers."
            if matched_deviation
            else (
                "Bottom line: the codified FAR Part 10 text is the only government-wide current rule established "
                "by this record. The model text and proposed rule are planning signals, while the comparator "
                "deviation does not prove adoption by the agency in scope."
            )
        ),
        "agency_status": (
            f"Bottom line: {agency} adoption is documented in the approved record. Confirm that the contemplated "
            "procurement falls within the deviation's scope, effective date, and transition terms before use."
            if matched_deviation
            else (
                "Bottom line: agency adoption remains unverified. The documented comparator demonstrates how "
                f"adoption can occur, but it cannot be used as evidence that {agency} adopted the same text."
            )
        ),
        "three_layer": (
            "Bottom line: the layers have different legal and operational roles. Use the codified rule as the "
            "baseline, treat the model as non-operative, and apply an agency deviation only to procurements within "
            "the issuing agency's documented scope."
        ),
        "change_brief": (
            "Bottom line: this evidence set supports a status comparison, not a textual change analysis. Until "
            "matched before-and-after provisions are obtained, the briefing should not claim that any requirement, "
            "threshold, or workflow changed."
        ),
        "rulemaking": (
            "Bottom line: the timeline shows policy development still in progress. The proposed-rule event is the "
            "latest formal rulemaking milestone in the approved record, so the next decision trigger is publication "
            "of a later docket event."
        ),
        "watchlist": (
            "Bottom line: there is a docket to monitor, but no approved evidence of a currently open comment window "
            "or due date. Assign an owner to verify the live notice before calendaring or mobilizing a response."
        ),
        "comments": (
            "Bottom line: a public-comment position analysis has not yet been earned by the evidence. The immediate "
            "work is to approve a bounded sample, retrieve it reproducibly, and document a coding method."
        ),
        "refresh": (
            "Bottom line: this record cannot distinguish what is new from what was previously known. A credible "
            "refresh requires a dated prior baseline and a current retrieval set before any changed/unchanged finding."
        ),
    }
    return interpretations.get(
        record.get("workflow_mode", ""),
        record.get("validation", {}).get("executive_summary", "No approved executive summary was supplied."),
    )


def collect_evidence_ids(value: object) -> set[str]:
    ids: set[str] = set()
    if isinstance(value, dict):
        for key, item in value.items():
            if key == "evidence_ids" and isinstance(item, list):
                ids.update(str(candidate) for candidate in item)
            else:
                ids.update(collect_evidence_ids(item))
    elif isinstance(value, list):
        for item in value:
            ids.update(collect_evidence_ids(item))
    return ids


def focused_evidence_ids(record: dict) -> set[str]:
    """Return evidence cited by the selected focused product, excluding unused fixture fields."""
    mode = record.get("workflow_mode", "")
    validation = record.get("validation", {})
    route_fields = {
        "change_brief": ("focused_findings", "focused_impacts", "planning_posture", "decision_gates", "change_map", "implementation_actions"),
        "watchlist": ("focused_findings", "focused_impacts", "planning_posture", "decision_gates", "rulemaking_watchlist", "watch_priorities"),
        "comments": ("focused_findings", "focused_impacts", "planning_posture", "decision_gates", "comment_themes"),
        "refresh": ("focused_findings", "focused_impacts", "planning_posture", "decision_gates", "refresh_changes", "carry_forward_decisions"),
    }
    if mode == "current_rule":
        return collect_evidence_ids(
            {
                "policy_items": [
                    item
                    for item in record.get("policy_items", [])
                    if item.get("status")
                    in {"codified_current", "model_deviation", "agency_class_deviation"}
                ],
                "planning_posture": derive_planning_posture(record),
                "decision_gates": derive_decision_gates(record),
            }
        )
    if mode in {"agency_status", "three_layer"}:
        return collect_evidence_ids(
            {
                "policy_items": record.get("policy_items", []),
                "planning_posture": derive_planning_posture(record),
                "decision_gates": derive_decision_gates(record),
            }
        )
    if mode == "rulemaking":
        return collect_evidence_ids(
            {
                "timeline": record.get("timeline", []),
                "planning_posture": derive_planning_posture(record),
                "decision_gates": derive_decision_gates(record),
            }
        )
    return collect_evidence_ids({field: validation.get(field) for field in route_fields.get(mode, ())})


def route_owner_labels(mode: str) -> tuple[str, str, str]:
    if mode in {"rulemaking", "watchlist", "comments", "refresh"}:
        return "Policy analyst", "Policy office", "Contracting team"
    if mode in {"current_rule", "change_brief"}:
        return "Policy analyst", "Contracting officer", "Policy and counsel"
    return "Policy office", "Contracting officer", "Policy and counsel"


def derive_decision_gates(record: dict) -> list[dict]:
    supplied = record.get("validation", {}).get("decision_gates")
    if isinstance(supplied, list) and supplied:
        return [item for item in supplied if isinstance(item, dict)]
    first_owner, second_owner, third_owner = route_owner_labels(record.get("workflow_mode", ""))
    unresolved = bool(record.get("unresolved_questions") or record.get("conflicts"))
    first_evidence = evidence_ids_for_statuses(
        record, {"agency_class_deviation", "model_deviation", "codified_current"}
    )
    rulemaking_evidence = evidence_ids_for_statuses(
        record, {"proposed_rule", "final_rule_pending_effective", "final_rule_effective"}
    )
    return [
        {
            "gate": "A",
            "evidence": "Confirm the current source layer, status, scope, and effective timing.",
            "owner": first_owner,
            "timing": "Before the analysis is used to draft or approve acquisition language",
            "evidence_ids": first_evidence,
        },
        {
            "gate": "B",
            "evidence": (
                "Resolve the recorded conflict or open applicability question."
                if unresolved
                else "Map the documented policy to the actual solicitation, award, option, or modification date."
            ),
            "owner": second_owner,
            "timing": "Before release or the next material procurement decision",
            "evidence_ids": first_evidence,
        },
        {
            "gate": "C",
            "evidence": "Refresh agency and rulemaking status when a material source or date changes.",
            "owner": third_owner,
            "timing": "At the stated refresh trigger and before relying on prior conclusions",
            "evidence_ids": rulemaking_evidence or first_evidence,
        },
    ]


def derive_scenarios(record: dict) -> list[dict]:
    supplied = record.get("validation", {}).get("planning_scenarios")
    if isinstance(supplied, list) and supplied:
        return [item for item in supplied if isinstance(item, dict)]
    statuses = {item.get("status") for item in record.get("policy_items", [])}
    baseline_ids = evidence_ids_for_statuses(record, {"codified_current", "final_rule_effective"})
    adoption_ids = evidence_ids_for_statuses(record, {"model_deviation", "agency_class_deviation"})
    rulemaking_ids = evidence_ids_for_statuses(
        record, {"proposed_rule", "final_rule_pending_effective", "final_rule_effective"}
    )
    scenarios = [
        {
            "scenario": "Baseline holds",
            "trigger": "No new agency adoption or effective rule changes the documented status.",
            "treatment": "Continue from the documented current baseline and retain the cited source set.",
            "evidence_ids": baseline_ids,
        }
    ]
    if statuses & {"model_deviation", "agency_class_deviation"}:
        scenarios.append(
            {
                "scenario": "Agency adoption is confirmed",
                "trigger": "An agency-issued source supplies applicable scope, text, and transition timing.",
                "treatment": "Apply only within that documented scope; reconcile acquisition language and file support.",
                "evidence_ids": adoption_ids,
            }
        )
    if statuses & {"proposed_rule", "final_rule_pending_effective", "final_rule_effective"}:
        scenarios.append(
            {
                "scenario": "Rulemaking changes status",
                "trigger": "A final rule becomes effective, is corrected, withdrawn, or is reflected in the codified text.",
                "treatment": "Refresh the Federal Register and codified baseline separately, then reassess transition treatment.",
                "evidence_ids": rulemaking_ids,
            }
        )
    if record.get("conflicts"):
        scenarios.append(
            {
                "scenario": "Conflict remains unresolved",
                "trigger": "Cited sources continue to disagree about a material value or timing term.",
                "treatment": "Preserve both positions and hold the disputed implementation point for an authorized official.",
                "evidence_ids": sorted(
                    {
                        evidence_id
                        for item in record.get("conflicts", [])
                        for evidence_id in item.get("evidence_ids", [])
                    }
                ),
            }
        )
    return scenarios[:3]


def add_posture_banner(document: Document, posture: dict) -> None:
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    set_repeat_table_header(table.rows[0])
    set_cell_text(table.cell(0, 0), str(posture["label"]).upper(), bold=True, color=WHITE)
    table.cell(0, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    shade(table.cell(0, 0), posture["accent"])
    shade(table.cell(0, 1), posture["fill"])
    table.cell(0, 1).text = ""
    paragraph = table.cell(0, 1).paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    headline = paragraph.add_run(posture["headline"])
    set_run_font(headline, size=12, color=DARK_BLUE, bold=True)
    paragraph.add_run("\n")
    rationale = paragraph.add_run(posture["rationale"])
    set_run_font(rationale, size=10, color=BLACK)
    cite_ids(paragraph, posture.get("evidence_ids", []))
    set_table_geometry(table, [1800, 7560])
    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def item_with_status(record: dict, status: str) -> dict:
    return next((item for item in record.get("policy_items", []) if item.get("status") == status), {})


def evidence_suffix(item: dict) -> str:
    ids = item.get("evidence_ids", [])
    return " [" + ", ".join(ids) + "]" if ids else ""


def validation_rows(record: dict, field: str) -> list[dict]:
    value = record.get("validation", {}).get(field, [])
    return [item for item in value if isinstance(item, dict)] if isinstance(value, list) else []


def require_validation_rows(record: dict, field: str, purpose: str) -> list[dict]:
    rows = validation_rows(record, field)
    if not rows:
        raise ValueError(
            f"{record.get('workflow_mode')} cannot produce a paid-value DOCX without {purpose}; "
            "return a concise evidence-acquisition note instead"
        )
    return rows


def value_with_evidence(item: dict, field: str) -> str:
    text = str(item.get(field, ""))
    ids = item.get("evidence_ids", [])
    if ids:
        text += " [" + ", ".join(str(value) for value in ids) + "]"
    return text


def enforce_route_content(record: dict) -> None:
    mode = record.get("workflow_mode")
    statuses = {item.get("status") for item in record.get("policy_items", [])}
    if mode == "current_rule":
        if "codified_current" not in statuses:
            raise ValueError(
                "current_rule cannot produce a paid-value DOCX without an approved codified-current source; "
                "return a concise evidence-acquisition note instead"
            )
    elif mode == "agency_status":
        if not record.get("scope", {}).get("agency") or not record.get("policy_items"):
            raise ValueError(
                "agency_status cannot produce a paid-value DOCX without a named agency and approved policy layers; "
                "return a concise evidence-acquisition note instead"
            )
    elif mode == "three_layer":
        required = {"codified_current", "model_deviation", "agency_class_deviation"}
        if not required.issubset(statuses):
            raise ValueError(
                "three_layer cannot produce a paid-value DOCX without codified, model, and agency-deviation layers; "
                "return a concise evidence-acquisition note instead"
            )
    elif mode == "change_brief":
        require_validation_rows(record, "change_map", "matched before-and-after provisions")
    elif mode == "rulemaking":
        if not record.get("timeline"):
            raise ValueError(
                "rulemaking cannot produce a paid-value DOCX without a dated milestone sequence; "
                "return a concise evidence-acquisition note instead"
            )
    elif mode == "watchlist":
        require_validation_rows(record, "rulemaking_watchlist", "live rulemaking matters, timing, and owner actions")
    elif mode == "comments":
        require_validation_rows(record, "comment_themes", "an approved comment sample and coded themes")
        if not record.get("stakeholder_positions"):
            raise ValueError("comments cannot produce a paid-value DOCX without stakeholder-position coverage")
    elif mode == "refresh":
        require_validation_rows(record, "refresh_changes", "a dated prior/current comparison")


def add_route_native_analysis(document: Document, record: dict) -> None:
    mode = record.get("workflow_mode", "")
    agency = record.get("scope", {}).get("agency") or "the named agency"
    codified = item_with_status(record, "codified_current")
    model = item_with_status(record, "model_deviation")
    deviation = item_with_status(record, "agency_class_deviation")
    proposed = item_with_status(record, "proposed_rule")

    if mode == "current_rule":
        document.add_heading("Current Rule Card", level=1)
        agency_check = (
            "Named-agency deviation is documented; apply only after confirming the procurement falls within its "
            "scope, effective date, and transition terms."
            if agency_item_matches_scope(record, deviation)
            else (
                "No agency-issued deviation for the named agency is established by this approved record; the "
                "recorded deviation is comparator evidence only."
            )
        )
        rows = [
            ["Documented baseline", codified.get("citation", "Not recorded"), codified.get("applicability_summary", "Not recorded") + evidence_suffix(codified)],
            ["Published comparison layer", model.get("citation", "Not recorded"), "Model text is informative but not agency-operative by itself." + evidence_suffix(model)],
            ["Agency check", agency, agency_check + evidence_suffix(deviation)],
        ]
        add_table(document, ["Rule-card field", "Recorded value", "Planning meaning"], rows, [1800, 2500, 5060])
    elif mode == "agency_status":
        document.add_heading("Agency Adoption Status", level=1)
        rows = []
        for item in record.get("policy_items", []):
            status = item.get("status", "")
            if status == "agency_class_deviation":
                relevance = (
                    "Named-agency evidence"
                    if agency_item_matches_scope(record, item)
                    else "Comparator only; does not establish adoption for the named agency"
                )
            else:
                relevance = STATUS_QUESTION_LABELS.get(status, "Documented status not classified")
            rows.append([status.replace("_", " ").title(), item.get("agency") or "Government-wide", relevance, ", ".join(item.get("evidence_ids", []))])
        add_table(document, ["Layer", "Issuer/agency", "Status for this question", "Evidence"], rows, [1800, 1900, 4200, 1460])
    elif mode == "three_layer":
        document.add_heading("Three-Layer Comparison and Adoption Test", level=1)
        deviation_test = (
            "Named-agency adoption evidence; confirm procurement scope and timing before use"
            if agency_item_matches_scope(record, deviation)
            else (
                "Comparator only unless the issuing agency matches the procurement agency and scope/timing are confirmed"
            )
        )
        deviation_source = deviation.get("citation", "Not recorded")
        if deviation.get("agency"):
            deviation_source += f" ({deviation['agency']})"
        rows = [
            ["Codified baseline", codified.get("citation", "Not recorded"), "Codified current baseline", "Use as the planning baseline" + evidence_suffix(codified)],
            ["Model text", model.get("citation", "Not recorded"), "Published model text; not operative alone", "Use only to identify possible deltas" + evidence_suffix(model)],
            ["Agency deviation", deviation_source, "Agency class deviation", deviation_test + evidence_suffix(deviation)],
        ]
        add_table(document, ["Layer", "Source", "Documented status", "Adoption test"], rows, [1600, 2100, 2500, 3160])
    elif mode == "change_brief":
        document.add_heading("Before/After Change Map", level=1)
        rows = require_validation_rows(record, "change_map", "matched before-and-after provisions")
        add_table(
            document,
            ["Provision", "Before", "After", "Substantive delta", "Operational consequence"],
            [[item.get("provision", ""), item.get("before", ""), item.get("after", ""), value_with_evidence(item, "delta"), item.get("operational_consequence", "")] for item in rows],
            [1250, 1900, 1900, 2100, 2210],
            font_size=8.5,
        )
        document.add_heading("Implementation decisions", level=2)
        add_bullets(document, record.get("validation", {}).get("implementation_actions", []), "No implementation action was approved.")
    elif mode == "rulemaking":
        document.add_heading("Rulemaking Milestones and Next Trigger", level=1)
        rows = [[item.get("date", ""), item.get("status", ""), item.get("event", ""), ", ".join(item.get("evidence_ids", []))] for item in record.get("timeline", [])]
        add_table(document, ["Date", "Status", "Milestone", "Evidence"], rows, [1300, 1700, 4900, 1460])
        document.add_paragraph("Next status-changing trigger: an effective final rule, correction, withdrawal, codification update, or applicable agency instruction. Until then, the proposed-rule layer remains non-operative.")
    elif mode == "watchlist":
        document.add_heading("Open Rulemaking Watchlist", level=1)
        rows = require_validation_rows(record, "rulemaking_watchlist", "live rulemaking matters, timing, and owner actions")
        add_table(
            document,
            ["Matter / docket", "Stage", "Deadline", "Next event", "Owner", "Recommended action"],
            [[item.get("matter", ""), item.get("stage", ""), value_with_evidence(item, "deadline"), item.get("next_event", ""), item.get("owner", ""), item.get("recommended_action", "")] for item in rows],
            [1800, 1050, 1250, 1700, 1200, 2360],
            font_size=8.5,
        )
        document.add_heading("Response priorities", level=2)
        add_bullets(document, record.get("validation", {}).get("watch_priorities", []), "No response priority was approved.")
    elif mode == "comments":
        document.add_heading("Comment Sample and Theme Coverage", level=1)
        positions = record.get("stakeholder_positions", [])
        rows = [[item.get("submitter_type", ""), item.get("position", ""), f"{item.get('reviewed_count', 0)} of {item.get('returned_count', 0)}", item.get("sample_method", ""), item.get("limitations", "")] for item in positions]
        add_table(document, ["Stakeholder segment", "Observed position", "Coverage", "Sample method", "Limits"], rows, [1500, 2500, 1050, 1900, 2410], font_size=7.5)
        document.add_heading("Coded themes and acquisition implications", level=1)
        themes = require_validation_rows(record, "comment_themes", "an approved comment sample and coded themes")
        add_table(
            document,
            ["Theme", "Observed pattern", "Segments", "Contrary view", "Acquisition implication"],
            [[item.get("theme", ""), value_with_evidence(item, "observed_pattern"), item.get("segments", ""), item.get("contrary_view", ""), item.get("acquisition_implication", "")] for item in themes],
            [1400, 2100, 1450, 1800, 2610],
            font_size=7.5,
        )
    elif mode == "refresh":
        document.add_heading("Refresh Change Register", level=1)
        rows = require_validation_rows(record, "refresh_changes", "a dated prior/current comparison")
        add_table(
            document,
            ["Policy issue", "Prior conclusion", "Current evidence", "Changed / unchanged", "Planning consequence"],
            [[item.get("issue", ""), item.get("prior_conclusion", ""), value_with_evidence(item, "current_evidence"), item.get("delta", ""), item.get("planning_consequence", "")] for item in rows],
            [1450, 2000, 2250, 1500, 2160],
            font_size=8.5,
        )
        document.add_heading("Carry-forward decisions", level=2)
        add_bullets(document, record.get("validation", {}).get("carry_forward_decisions", []), "No prior conclusion was approved for carry-forward.")


def build(record: dict, output: Path) -> None:
    validation = record.get("validation", {})
    if not validation.get("findings_approved") or not validation.get("brief_approved"):
        raise ValueError("findings and brief generation must be approved before building the DOCX")
    enforce_route_content(record)

    document = Document()
    configure_styles(document)
    request = record["request"]
    scope = record["scope"]
    as_of = scope["as_of_date"]
    agency = scope.get("agency") or "Published federal acquisition policy"
    lens = request["audience_lens"]

    kicker = document.add_paragraph()
    kicker.paragraph_format.space_before = Pt(18)
    kicker.paragraph_format.space_after = Pt(8)
    run = kicker.add_run("ACQUISITION POLICY")
    set_run_font(run, size=10, color=BLUE, bold=True)
    product_title = validation.get("report_title") or PRODUCT_TITLES.get(
        record.get("workflow_mode", ""), "Acquisition Policy Impact Brief"
    )
    for section in document.sections:
        header = section.header.paragraphs[0]
        header.text = f"1102tools  |  {product_title}"
        header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_run_font(header.runs[0], size=8, color=MID_GRAY)
    title = document.add_paragraph(style="Policy Title")
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_run = title.add_run(product_title)
    set_run_font(
        title_run,
        size=20 if len(product_title) >= 26 else 24,
        color=BLACK,
        bold=True,
    )
    subtitle = document.add_paragraph(request.get("question", ""))
    subtitle.paragraph_format.space_after = Pt(14)
    if subtitle.runs:
        set_run_font(subtitle.runs[0], size=13, color=MID_GRAY)
    add_metadata(document, "Prepared for", scope.get("customer_organization"))
    add_metadata(document, "Decision date", scope.get("decision_date"))
    add_metadata(document, "Agency or scope", agency)
    add_metadata(document, "As of", as_of)
    if record.get("workflow_mode") == "refresh":
        prior = scope.get("prior_analysis") or {}
        prior_label = " ".join(
            part for part in (str(prior.get("title", "")).strip(), f"({prior.get('date', '')})" if prior.get("date") else "")
            if part
        )
        add_metadata(document, "Prior analysis", prior_label)
    add_metadata(document, "Audience lens", lens.title())
    add_metadata(document, "Status", "Documented published-source analysis; not a legal opinion or procurement-specific determination")
    rule = document.add_paragraph()
    rule.paragraph_format.space_before = Pt(8)
    rule.paragraph_format.space_after = Pt(12)
    p_pr = rule._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), BLUE)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)

    document.add_heading("Planning Posture and Implications", level=1)
    add_posture_banner(document, derive_planning_posture(record))
    document.add_paragraph(derive_front_page_interpretation(record))
    document.add_heading("What is established", level=2)
    approved_findings = validation.get("focused_findings") or record.get("findings", [])
    for finding in approved_findings[:3]:
        paragraph = document.add_paragraph(finding.get("text", ""), style="List Bullet")
        cite_ids(paragraph, finding.get("evidence_ids", []))
    if not approved_findings:
        document.add_paragraph("No approved finding was recorded.")
    document.add_heading("Immediate implications", level=2)
    leading_impacts = (validation.get("focused_impacts") or impact_items(record, lens))[:3]
    add_bullets(
        document,
        leading_impacts,
        "Use the documented status to frame acquisition planning, timing, monitoring, and file support.",
    )
    document.add_heading("Owners and Decision Gates", level=1)
    decision_gates = derive_decision_gates(record)
    gate_rows = []
    for gate in decision_gates:
        evidence_text = gate.get("evidence", "")
        evidence_ids = gate.get("evidence_ids", [])
        if evidence_ids:
            evidence_text += " [" + ", ".join(evidence_ids) + "]"
        gate_rows.append(
            [
                gate.get("gate", ""),
                evidence_text,
                gate.get("owner", ""),
                gate.get("timing", ""),
            ]
        )
    # The first column must stay wide enough for multi-word gate labels such as
    # "Adoption confirmation"; narrower widths force mid-word breaks in fixed layout.
    add_table(
        document,
        ["Gate", "Decision-ready evidence", "Owner", "Timing"],
        gate_rows,
        [1800, 3160, 1700, 2700],
    )
    boundary = document.add_paragraph()
    boundary.paragraph_format.space_before = Pt(4)
    boundary.paragraph_format.space_after = Pt(6)
    boundary_run = boundary.add_run("Reserved decision: ")
    set_run_font(boundary_run, size=9.5, color=DARK_BLUE, bold=True)
    boundary_text = boundary.add_run(
        "This brief states what cited published sources indicate as of the date shown. "
        "An authorized agency official must determine procurement-specific applicability."
    )
    set_run_font(boundary_text, size=9.5, color=BLACK)

    add_route_native_analysis(document, record)

    if record.get("workflow_mode") in FOCUSED_PRODUCTS:
        document.add_heading("Management Actions", level=1)
        supplied_actions = record.get("validation", {}).get("management_actions")
        management_actions = (
            [item for item in supplied_actions if isinstance(item, dict)]
            if isinstance(supplied_actions, list)
            else []
        )
        if not management_actions or management_actions == decision_gates:
            # Do not repeat the decision-gate rows as a second table.
            document.add_paragraph(
                "The owners, timing, and decision-ready evidence for each management action are stated once in "
                "the Owners and Decision Gates table above."
            )
        else:
            action_rows = []
            for action in management_actions:
                evidence_text = action.get("evidence", "")
                if action.get("evidence_ids"):
                    evidence_text += " [" + ", ".join(action["evidence_ids"]) + "]"
                action_rows.append([action.get("owner", ""), action.get("timing", ""), action.get("gate", ""), evidence_text])
            add_table(document, ["Owner", "Timing", "Action / decision gate", "Evidence needed"], action_rows, [1700, 1900, 2200, 3560], font_size=9)

        document.add_heading("Evidence and Source Notes", level=1)
        focused_ids = focused_evidence_ids(record)
        evidence = [item for item in record.get("evidence", []) if item.get("id") in focused_ids]
        if evidence:
            table = add_table(
                document,
                ["ID", "Source", "Decision-useful fact", "Limit"],
                [[item.get("id", ""), f"{item.get('title', '')}\n{item.get('locator', '')}", item.get("fact", ""), item.get("limitations", "")] for item in evidence],
                [700, 2600, 3900, 2160],
                font_size=8.5,
                add_spacer=False,
            )
            for row, item in zip(table.rows[1:], evidence):
                url = item.get("canonical_url", "")
                if url:
                    paragraph = row.cells[1].add_paragraph()
                    paragraph.paragraph_format.space_after = Pt(0)
                    add_hyperlink(paragraph, urlparse(url).netloc or "Official source", url, font_size=8.5)
        else:
            document.add_paragraph("No approved evidence item was recorded.")

        document.add_heading("Limitations and Reserved Determinations", level=1)
        limitations = list(record.get("limitations", [])) or ["No additional limitation was recorded."]
        limitations[-1] += (
            " This product does not provide legal advice. It states what cited published sources indicate as of the date shown. "
            "An authorized agency official must determine procurement-specific applicability."
        )
        limitation_start = len(document.paragraphs)
        add_bullets(document, limitations, "No additional limitation was recorded.")
        for paragraph in document.paragraphs[limitation_start:]:
            paragraph.paragraph_format.space_after = Pt(2)
            paragraph.paragraph_format.line_spacing = 1.0
            for run in paragraph.runs:
                set_run_font(run, size=8.5, color=BLACK)
        output.parent.mkdir(parents=True, exist_ok=True)
        document.save(output)
        return

    document.add_heading("Question and Scope", level=1)
    document.add_paragraph(request.get("question", "Not stated"))
    scope_rows = [[pretty_label(key), display_value(value)] for key, value in scope.items()]
    add_table(document, ["Scope field", "Approved value"], scope_rows, [2700, 6660])

    document.add_heading("Documented Current Status", level=1)
    if record.get("policy_items"):
        add_table(
            document,
            ["ID", "Status", "Citation", "Agency", "Documented treatment", "Evidence"],
            policy_rows(record),
            [850, 1450, 1200, 1450, 3060, 1350],
        )
    else:
        document.add_paragraph("No approved policy item was recorded.")
    document.add_heading("Approved findings", level=2)
    for finding in record.get("findings", []):
        paragraph = document.add_paragraph(finding.get("text", ""), style="List Bullet")
        cite_ids(paragraph, finding.get("evidence_ids", []))
    if not record.get("findings"):
        document.add_paragraph("No approved finding was recorded.")

    document.add_heading("Source Hierarchy and Authorities", level=1)
    hierarchy = validation.get("source_hierarchy", [])
    add_bullets(
        document,
        hierarchy,
        "The analysis distinguishes codified text, agency deviations, model text, rulemaking, guidance, comments, and supplied documents.",
    )

    document.add_heading("Planning Scenarios", level=1)
    scenario_rows = []
    for scenario in derive_scenarios(record):
        treatment = scenario.get("treatment", "")
        evidence_ids = scenario.get("evidence_ids", [])
        if evidence_ids:
            treatment += " [" + ", ".join(evidence_ids) + "]"
        scenario_rows.append(
            [scenario.get("scenario", ""), scenario.get("trigger", ""), treatment]
        )
    add_table(
        document,
        ["Scenario", "Trigger", "Planning treatment"],
        scenario_rows,
        [1900, 3300, 4160],
    )

    document.add_heading("Change Timeline", level=1)
    timeline = record.get("timeline", [])
    if timeline:
        add_table(
            document,
            ["Date", "Event", "Status", "Evidence"],
            [[item.get("date", ""), item.get("event", ""), item.get("status", ""), ", ".join(item.get("evidence_ids", []))] for item in timeline],
            [1250, 4700, 1850, 1560],
        )
    else:
        document.add_paragraph("No change event was required for the approved scope.")

    document.add_heading("Government and Industry Impacts", level=1)
    if lens == "neutral":
        document.add_heading("Government lens", level=2)
        add_bullets(document, validation.get("impacts", {}).get("government", []), "No approved government impact was recorded.")
        document.add_heading("Industry lens", level=2)
        add_bullets(document, validation.get("impacts", {}).get("industry", []), "No approved industry impact was recorded.")
    else:
        document.add_heading(f"{lens.title()} lens", level=2)
        add_bullets(document, impact_items(record, lens), f"No approved {lens} impact was recorded.")

    document.add_heading("Open Issues and Comment Deadlines", level=1)
    deadlines = validation.get("open_issues", [])
    add_bullets(document, deadlines, "No open issue or comment deadline was identified within the approved scope.")
    if record.get("stakeholder_positions"):
        document.add_heading("Observed stakeholder positions", level=2)
        add_table(
            document,
            ["Submitter type", "Observed position", "Sample", "Coverage and limitations", "Evidence"],
            [
                [
                    item.get("submitter_type", ""),
                    item.get("position", ""),
                    f"{item.get('reviewed_count', 0)} of {item.get('returned_count', 0)}; {item.get('sample_method', '')}",
                    item.get("limitations", ""),
                    ", ".join(item.get("evidence_ids", [])),
                ]
                for item in record["stakeholder_positions"]
            ],
            [1250, 3000, 1600, 2350, 1160],
        )

    document.add_heading("Operational Considerations", level=1)
    add_bullets(document, validation.get("operational_considerations", []), "No operational consideration was approved.")
    if record.get("conflicts"):
        document.add_heading("Conflicts", level=2)
        add_table(
            document,
            ["ID", "Issue", "Status", "Resolution and source"],
            [
                [
                    item.get("id", ""),
                    item.get("issue", ""),
                    item.get("status", "").replace("_", " ").title(),
                    (
                        f"{item.get('resolution', '')} ({item.get('resolved_by', '')}; {item.get('resolved_at', '')})"
                        if item.get("resolution")
                        else "Reserved to an authorized official"
                    ),
                ]
                for item in record["conflicts"]
            ],
            [850, 4200, 1700, 2610],
        )
    if record.get("unresolved_questions"):
        document.add_heading("Unresolved questions", level=2)
        add_bullets(document, record["unresolved_questions"], "")

    document.add_heading("Evidence Register", level=1)
    evidence = record.get("evidence", [])
    if evidence:
        table = add_table(
            document,
            ["ID", "Type", "Source and locator", "Supported fact", "Limits"],
            [
                [
                    item.get("id", ""),
                    item.get("source_type", "").replace("_", " ").title(),
                    f"{item.get('title', '')}\n{item.get('locator', '')}",
                    item.get("fact", ""),
                    item.get("limitations", ""),
                ]
                for item in evidence
            ],
            [850, 1400, 2350, 3000, 1760],
            font_size=8.5,
            add_spacer=False,
        )
        for row, item in zip(table.rows[1:], evidence):
            url = item.get("canonical_url", "")
            if url:
                paragraph = row.cells[2].add_paragraph()
                paragraph.paragraph_format.space_before = Pt(2)
                paragraph.paragraph_format.space_after = Pt(0)
                add_hyperlink(paragraph, urlparse(url).netloc or "Official source", url, font_size=8.5)
    else:
        document.add_paragraph("No evidence item was recorded.")

    document.add_section(WD_SECTION.NEW_PAGE)
    document.add_heading("Limitations and Reserved Determinations", level=1)
    add_bullets(document, record.get("limitations", []), "No additional limitation was recorded.")
    document.add_paragraph(
        "This brief does not provide legal advice, approve policy, select clauses, or determine which rule legally governs a specific procurement. "
        "Confirm transaction-specific treatment with the responsible contracting, policy, and legal officials."
    )

    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("record", type=Path)
    parser.add_argument("output", type=Path)
    args = parser.parse_args()
    validator = Path(__file__).with_name("validate_policy_research_record.py")
    result = subprocess.run([sys.executable, str(validator), str(args.record)], check=False)
    if result.returncode:
        return result.returncode
    try:
        build(json.loads(args.record.read_text(encoding="utf-8")), args.output)
    except (OSError, ValueError, json.JSONDecodeError) as exc:
        print(f"ERROR: {exc}", file=sys.stderr)
        return 2
    print(args.output)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
